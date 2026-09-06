from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader, select_autoescape
from playwright.async_api import async_playwright

from modules.exams.rich_text import parse_inline, parse_rich_text, render_rich_html

TEMPLATE_DIR = Path(__file__).parent / "templates"

_env = Environment(
    loader=FileSystemLoader(str(TEMPLATE_DIR)),
    autoescape=select_autoescape(["html"]),
)

VALID_EXPORT_TYPES = {"de", "dapan", "de_dapan"}


def _split_answer_keys(correct_answer: Any) -> list[str]:
    if not correct_answer:
        return []
    if isinstance(correct_answer, list):
        return [str(item).strip() for item in correct_answer]
    return [part.strip() for part in str(correct_answer).split(",") if part.strip()]


def _format_answer(snapshot: dict, correct_answer: Any) -> str:
    keys = _split_answer_keys(correct_answer)
    question_type = str(
        (snapshot.get("classification") or {}).get("assessment_type") or ""
    ).upper()
    if question_type in {"SAP_XEP", "ORDERING", "GHEP_COT", "MATCHING"}:
        return ", ".join(keys)
    return ", ".join(sorted(set(keys))) if keys else str(correct_answer or "")


def _build_context(
    header: dict,
    exam_code: str,
    questions: list[dict],
    export_type: str,
) -> dict:
    show_questions = export_type in {"de", "de_dapan"}
    show_answers = export_type == "de_dapan"
    show_answer_table = export_type in {"dapan", "de_dapan"}

    rendered_questions = []
    answer_rows = []
    for entry in sorted(questions, key=lambda item: item["order"]):
        snapshot = entry["content_snapshot"]
        question_data = snapshot.get("question_data") or {}
        options = question_data.get("options") or {}
        correct_answer = question_data.get("correct_answer")
        correct_keys = set(_split_answer_keys(correct_answer))
        rendered_options = [
            {
                "label": key,
                "text": value,
                "html": render_rich_html(value),
                "correct": key in correct_keys,
            }
            for key, value in options.items()
        ]
        rendered_questions.append(
            {
                "number": entry["order"],
                "content": snapshot.get("content", ""),
                "content_html": render_rich_html(snapshot.get("content", "")),
                "options": rendered_options,
            }
        )
        answer_rows.append(
            {
                "number": entry["order"],
                "answer": _format_answer(snapshot, correct_answer),
            }
        )

    return {
        "header": header,
        "exam_code": exam_code,
        "questions": rendered_questions,
        "answer_rows": answer_rows,
        "show_questions": show_questions,
        "show_answers": show_answers,
        "show_answer_table": show_answer_table,
    }


def render_exam_html(
    header: dict,
    exam_code: str,
    questions: list[dict],
    export_type: str,
) -> str:
    if export_type not in VALID_EXPORT_TYPES:
        raise ValueError(f"export_type không hợp lệ: {export_type}")
    template = _env.get_template("exam_pdf.html")
    context = _build_context(header, exam_code, questions, export_type)
    return template.render(**context)


async def render_exam_pdf(
    header: dict,
    exam_code: str,
    questions: list[dict],
    export_type: str,
) -> bytes:
    if not questions:
        raise ValueError("Đề thi chưa có câu hỏi, không thể xuất PDF")
    html = render_exam_html(header, exam_code, questions, export_type)
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch()
        try:
            page = await browser.new_page()
            await page.set_content(html, wait_until="load")
            pdf_bytes = await page.pdf(format="A4", print_background=True)
        finally:
            await browser.close()
    return pdf_bytes


def _add_labeled_line(document: Document, label: str, value: Any) -> None:
    if value in (None, ""):
        return
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value))


def _set_cell_text(cell, value: str, *, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold


def _append_math(paragraph, text: str) -> None:
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = text
    math_run.append(math_text)
    math.append(math_run)
    paragraph._p.append(math)


def _append_inlines(paragraph, inlines: list[dict], *, bold: bool = False) -> None:
    for item in inlines:
        kind = item["kind"]
        if kind == "image":
            run = paragraph.add_run()
            run.add_picture(BytesIO(item["image"]), width=Inches(2.5))
            continue
        if kind == "math":
            _append_math(paragraph, str(item.get("text") or ""))
            continue
        run = paragraph.add_run(str(item.get("text") or ""))
        run.bold = bold or kind == "strong"
        if kind == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEEEEE")
            run._r.get_or_add_rPr().append(shading)


def _add_rich_content(
    document: Document,
    value: Any,
    *,
    prefix: str = "",
    left_indent: Pt | None = None,
    bold: bool = False,
) -> None:
    blocks = parse_rich_text(value)
    prefix_pending = prefix
    for block in blocks:
        kind = block["kind"]
        if kind == "table":
            if prefix_pending:
                paragraph = document.add_paragraph()
                paragraph.add_run(prefix_pending).bold = True
                prefix_pending = ""
            table = document.add_table(rows=1, cols=max(1, len(block["header"])))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, block["header"]):
                _set_cell_text(cell, str(text), bold=True)
            for values in block["rows"]:
                for cell, text in zip(table.add_row().cells, values):
                    _set_cell_text(cell, str(text))
            continue
        paragraph = document.add_paragraph()
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = left_indent
        if prefix_pending:
            paragraph.add_run(prefix_pending).bold = True
            prefix_pending = ""
        if kind == "paragraph":
            _append_inlines(paragraph, block["inlines"], bold=bold)
        elif kind == "math":
            _append_math(paragraph, str(block.get("text") or ""))
        elif kind == "code":
            run = paragraph.add_run(str(block.get("text") or ""))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEEEEE")
            run._r.get_or_add_rPr().append(shading)
    if prefix_pending:
        paragraph = document.add_paragraph()
        paragraph.add_run(prefix_pending).bold = True


def render_exam_docx(
    header: dict,
    exam_code: str,
    questions: list[dict],
    export_type: str,
) -> bytes:
    if export_type not in VALID_EXPORT_TYPES:
        raise ValueError(f"export_type không hợp lệ: {export_type}")
    if not questions:
        raise ValueError("Đề thi chưa có câu hỏi, không thể xuất DOCX")

    context = _build_context(header, exam_code, questions, export_type)
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)

    section = document.sections[0]
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    if header.get("school_name"):
        document.add_paragraph(str(header["school_name"]))
    if header.get("faculty_name"):
        document.add_paragraph(str(header["faculty_name"]))

    title = document.add_heading(header.get("exam_name") or "Đề thi", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_labeled_line(document, "Mã đề", exam_code)
    _add_labeled_line(document, "Môn học", header.get("subject_name"))
    if header.get("duration_minutes"):
        _add_labeled_line(document, "Thời gian", f"{header.get('duration_minutes')} phút")
    _add_labeled_line(document, "Lớp", header.get("class_name"))
    _add_labeled_line(document, "Phòng", header.get("room"))
    _add_labeled_line(document, "Ngày thi", header.get("exam_date"))

    if context["show_questions"]:
        document.add_paragraph()
        for question in context["questions"]:
            paragraph_start = len(document.paragraphs)
            _add_rich_content(
                document,
                question["content"],
                prefix=f"Câu {question['number']}. ",
                bold=True,
            )
            for option in question["options"]:
                _add_rich_content(
                    document,
                    option["text"],
                    prefix=f"{option['label']}. ",
                    left_indent=Pt(18),
                    bold=context["show_answers"] and option["correct"],
                )
            question_paragraphs = document.paragraphs[paragraph_start:]
            for paragraph in question_paragraphs:
                paragraph.paragraph_format.keep_together = True
            for paragraph in question_paragraphs[:-1]:
                paragraph.paragraph_format.keep_with_next = True

    if context["show_answer_table"]:
        document.add_paragraph()
        document.add_heading("Đáp án", level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        _set_cell_text(header_cells[0], "Câu", bold=True)
        _set_cell_text(header_cells[1], "Đáp án", bold=True)
        for row in context["answer_rows"]:
            cells = table.add_row().cells
            cells[0].text = str(row["number"])
            cells[1].text = str(row["answer"])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
