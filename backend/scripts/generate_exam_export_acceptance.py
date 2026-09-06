"""Generate and structurally verify four PDF/DOCX exam variants for visual QA."""

from __future__ import annotations

import argparse
import asyncio
import base64
import hashlib
import json
import re
from io import BytesIO
from pathlib import Path

import fitz
from bson import ObjectId
from docx import Document
from PIL import Image, ImageDraw

from core.dependencies import CurrentUser
from modules.exams.pdf_service import _format_answer, render_exam_docx, render_exam_pdf
from modules.exams.schemas import ExamVariantCreateRequest
from modules.exams.service import ExamVariantService


def _diagram_data_uri() -> str:
    buffer = BytesIO()
    image = Image.new("RGB", (360, 100), "white")
    draw = ImageDraw.Draw(image)
    for index, label in enumerate(("front", "A", "B", "rear")):
        left = 10 + index * 85
        draw.rectangle((left, 20, left + 70, 75), outline="black", width=2)
        draw.text((left + 12, 40), label, fill="black")
    image.save(buffer, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buffer.getvalue()).decode("ascii")


DIAGRAM_DATA_URI = _diagram_data_uri()


QUESTION_DEFINITIONS = (
    (
        "TRAC_NGHIEM",
        "Với đoạn mã C++ `int a[] = {4, 2, 7};`, độ phức tạp truy cập a[2] là gì?\n\n"
        "```cpp\nint value = a[2];\n```",
        {"A": "O(n)", "B": "O(1)", "C": "O(log n)", "D": "O(n log n)"},
        "B",
    ),
    (
        "NHIEU_LUA_CHON",
        "Chọn mọi cấu trúc hỗ trợ thao tác chèn và xóa ở hai đầu trong O(1).",
        {"A": "Deque", "B": "Mảng tĩnh không dịch phần tử", "C": "Danh sách liên kết đôi", "D": "Heap nhị phân"},
        "A,C",
    ),
    (
        "DUNG_SAI",
        "Cây tìm kiếm nhị phân cân bằng có chiều cao O(log n).\n\n"
        f"![Sơ đồ hàng đợi minh họa]({DIAGRAM_DATA_URI})",
        {"A": "Đúng", "B": "Sai"},
        "A",
    ),
    (
        "DIEN_KHUYET",
        "Ngăn xếp tuân theo nguyên tắc ____.",
        {},
        "LIFO",
    ),
    (
        "GHEP_COT",
        "Ghép cấu trúc ở cột số với đặc trưng phù hợp ở cột chữ.",
        {
            "1": "Stack",
            "2": "Queue",
            "3": "Heap",
            "A": "LIFO",
            "B": "FIFO",
            "C": "Lấy phần tử ưu tiên",
            "D": "Duyệt theo khóa băm",
        },
        "1-A,2-B,3-C",
    ),
    (
        "SAP_XEP",
        "Sắp xếp các bước duyệt BFS đúng thứ tự.",
        {
            "A": "Đánh dấu đỉnh đã thăm",
            "B": "Đưa đỉnh bắt đầu vào hàng đợi",
            "C": "Lặp cho đến khi hàng đợi rỗng",
            "D": "Lấy một đỉnh khỏi hàng đợi và thêm các đỉnh kề chưa thăm",
        },
        "B,A,D,C",
    ),
    (
        "TINH_HUONG",
        "Một hệ thống cần tìm kiếm khóa trung bình O(1), không yêu cầu thứ tự. Chọn cấu trúc phù hợp nhất.\n\n"
        "Công thức tải:\n\n$$\\alpha = \\frac{n}{m}$$\n\n"
        "| n | bucket | α |\n|---:|---:|---:|\n| 8 | 16 | 0.5 |",
        {
            "A": "Danh sách liên kết đơn",
            "B": "Cây tìm kiếm nhị phân không cân bằng",
            "C": "Bảng băm với chính sách xử lý va chạm",
            "D": "Ngăn xếp",
        },
        "C",
    ),
)


class _ExamRepository:
    def __init__(self, exam: dict):
        self.exam = exam

    def find(self, _exam_id):
        return self.exam

    def count_variants(self, _exam_id):
        return 0


class _VariantRepository:
    def __init__(self):
        self.rows: list[dict] = []

    def create(self, value):
        self.rows.append(value)
        return value


def _normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _question_rows() -> list[dict]:
    rows = []
    for qtype, content, options, answer in QUESTION_DEFINITIONS:
        question_id = ObjectId()
        rows.append(
            {
                "question_id": question_id,
                "version_id": ObjectId(),
                "content_snapshot": {
                    "content": content,
                    "classification": {"assessment_type": qtype},
                    "question_data": {
                        "options": options,
                        "correct_answer": answer,
                        "explanation": "SECRET_EXPLANATION_MUST_NOT_LEAK",
                    },
                },
            }
        )
    return rows


def _user() -> CurrentUser:
    user_id = ObjectId()
    return CurrentUser(
        id=user_id,
        firebase_uid=str(user_id),
        email="acceptance.teacher@example.edu",
        role="Teacher",
        is_active=True,
        permissions=(),
        display_name="Acceptance Teacher",
    )


def _extract_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def _docx_rich_checks(data: bytes) -> dict[str, bool]:
    document = Document(BytesIO(data))
    xml = document._element.xml
    return {
        "docx_code_marker_absent": "```" not in _extract_docx(data),
        "docx_rich_table_present": len(document.tables) >= 1,
        "docx_math_present": "<m:oMath>" in xml,
        "docx_math_text_present": "α = (n)/(m)" in xml,
        "docx_image_present": len(document.inline_shapes) >= 1,
    }


def _extract_pdf(data: bytes) -> tuple[str, int]:
    document = fitz.open(stream=data, filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in document), document.page_count
    finally:
        document.close()


def _validate_export(
    *,
    variant: dict,
    export_type: str,
    file_format: str,
    payload: bytes,
) -> dict:
    if file_format == "pdf":
        extracted, page_count = _extract_pdf(payload)
    else:
        extracted, page_count = _extract_docx(payload), None
    normalized = _normalize(extracted)
    checks = {
        "non_empty": len(payload) > 1000,
        "secret_explanation_absent": "SECRET_EXPLANATION_MUST_NOT_LEAK" not in extracted,
        "student_answer_heading_absent": export_type != "de" or "Đáp án" not in extracted,
        "answer_heading_present": export_type != "dapan" or "Đáp án" in extracted,
    }
    if export_type != "dapan":
        checks.update(
            {
                "raw_fence_absent": "```" not in extracted,
                "raw_table_separator_absent": "|---" not in extracted,
                "raw_math_delimiter_absent": "$$" not in extracted,
                "rich_code_text_present": "int value = a[2];" in extracted,
                "rich_math_text_present": (
                    "α = (n)/(m)" in _normalize(extracted)
                    if file_format == "pdf"
                    else True
                ),
            }
        )
        if file_format == "docx":
            checks.update(_docx_rich_checks(payload))
    expected_answers = []
    for entry in variant["questions"]:
        snapshot = entry["content_snapshot"]
        correct_answer = (snapshot.get("question_data") or {}).get("correct_answer")
        question_id = str(entry["question_id"])
        checks[f"answer_key_matches_snapshot:{question_id}"] = (
            variant["answer_key"].get(question_id) == correct_answer
        )
        if export_type == "dapan":
            formatted = _format_answer(snapshot, correct_answer)
            expected_answers.append(formatted)
            checks[f"answer_rendered:{question_id}"] = _normalize(formatted) in normalized
    return {
        "checks": checks,
        "passed": all(checks.values()),
        "expected_answers": expected_answers,
        "page_count": page_count,
        "extracted_text_sha256": _sha256(extracted.encode("utf-8")),
    }


async def generate(output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    user = _user()
    questions = _question_rows()
    exam_id = ObjectId()
    exam = {
        "_id": exam_id,
        "created_by_user_id": user.id,
        "subject_id": ObjectId(),
        "status": "FINALIZED",
        "question_count": len(questions),
        "questions": questions,
        "finalized_snapshot": {"questions": questions, "revision": 1},
    }
    header = {
        "school_name": "TRƯỜNG ĐẠI HỌC CẦN THƠ",
        "faculty_name": "KHOA CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG",
        "exam_name": "KIỂM TRA NGHIỆM THU XUẤT ĐỀ",
        "subject_name": "Cấu trúc dữ liệu",
        "duration_minutes": 60,
        "class_name": "CTDL-ACCEPTANCE",
        "room": "LAB",
        "exam_date": "06/09/2026",
    }
    variants = _VariantRepository()
    service = ExamVariantService(_ExamRepository(exam), variants)
    files = []
    variant_rows = []
    for exam_code in ("101", "102", "103", "104"):
        variant = service.create_variant(
            str(exam_id),
            ExamVariantCreateRequest(exam_code=exam_code, shuffle=True),
            user,
        )
        variant_rows.append(
            {
                "exam_code": exam_code,
                "seed": variant["seed"],
                "permutation": variant["permutation"],
                "answer_key": variant["answer_key"],
                "question_types": [
                    entry["content_snapshot"]["classification"]["assessment_type"]
                    for entry in variant["questions"]
                ],
            }
        )
        for export_type in ("de", "dapan"):
            pdf = await render_exam_pdf(header, exam_code, variant["questions"], export_type)
            docx = render_exam_docx(header, exam_code, variant["questions"], export_type)
            for file_format, payload in (("pdf", pdf), ("docx", docx)):
                path = output_dir / f"ctdl-{exam_code}-{export_type}.{file_format}"
                path.write_bytes(payload)
                validation = _validate_export(
                    variant=variant,
                    export_type=export_type,
                    file_format=file_format,
                    payload=payload,
                )
                files.append(
                    {
                        "path": path.name,
                        "format": file_format,
                        "export_type": export_type,
                        "exam_code": exam_code,
                        "bytes": len(payload),
                        "sha256": _sha256(payload),
                        **validation,
                    }
                )
    manifest = {
        "scope": "Synthetic CTDL export fixture; technical QA evidence, not licensed-corpus or human UAT evidence.",
        "question_types": [item[0] for item in QUESTION_DEFINITIONS],
        "variants": variant_rows,
        "files": files,
        "passed": all(row["passed"] for row in files),
    }
    manifest_path = output_dir / "export-manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    manifest = asyncio.run(generate(args.output.resolve()))
    print(
        json.dumps(
            {
                "passed": manifest["passed"],
                "files": len(manifest["files"]),
                "question_types": len(manifest["question_types"]),
            },
            ensure_ascii=False,
        )
    )
    return 0 if manifest["passed"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
