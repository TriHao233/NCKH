from __future__ import annotations

import argparse
import json
import random
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFilter, ImageFont
from pypdf import PdfReader, PdfWriter


BACKEND_DIR = Path(__file__).resolve().parents[1]
CORPUS_DIR = BACKEND_DIR / "tests" / "golden_corpus" / "v1"


def _font(size: int) -> ImageFont.FreeTypeFont:
    candidates = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
    )
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    raise RuntimeError("Không tìm thấy font Unicode để render golden corpus")


def _render_scan(text: str, *, degraded: bool = False) -> Image.Image:
    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    font = _font(34)
    y = 100
    for source_line in text.splitlines():
        words = source_line.split()
        line = ""
        for word in words or [""]:
            candidate = f"{line} {word}".strip()
            if draw.textlength(candidate, font=font) > 1040 and line:
                draw.text((100, y), line, fill="black", font=font)
                y += 52
                line = word
            else:
                line = candidate
        draw.text((100, y), line, fill="black", font=font)
        y += 58
    if degraded:
        image = image.rotate(1.4, resample=Image.Resampling.BICUBIC, expand=False, fillcolor="white")
        image = image.filter(ImageFilter.GaussianBlur(radius=1.1))
        pixels = image.load()
        random.seed(20260901)
        for _ in range(9000):
            x = random.randrange(image.width)
            y = random.randrange(image.height)
            shade = random.randrange(150, 245)
            pixels[x, y] = (shade, shade, shade)
    return image


def _save_images_pdf(images: list[Image.Image], destination: Path) -> None:
    images[0].save(destination, "PDF", resolution=150, save_all=True, append_images=images[1:])


def _add_graph_diagram(image: Image.Image, *, caption: str) -> Image.Image:
    draw = ImageDraw.Draw(image)
    font = _font(30)
    nodes = ((360, 700, "u"), (850, 700, "v"), (605, 1030, "w"))
    for left, top, _label in nodes:
        draw.ellipse((left - 55, top - 55, left + 55, top + 55), outline="black", width=5)
    draw.line((415, 700, 795, 700), fill="black", width=5)
    draw.line((395, 745, 565, 985), fill="black", width=5)
    draw.line((815, 745, 645, 985), fill="black", width=5)
    for left, top, label in nodes:
        draw.text((left - 10, top - 22), label, fill="black", font=font)
    draw.text((420, 1180), caption, fill="black", font=font)
    return image


def _illustration_image() -> Image.Image:
    image = Image.new("RGB", (640, 360), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((80, 70, 560, 290), outline="black", width=5)
    draw.ellipse((245, 115, 395, 265), outline="black", width=5)
    return image


def _write_docx(destination: Path, pages: list[dict], *, legacy_fixture: bool = False) -> None:
    document = Document()
    if "Code" not in document.styles:
        document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    for index, page in enumerate(pages):
        table_text_lines = {
            cell
            for table_truth in page.get("tables") or []
            for row in table_truth.get("rows") or []
            for cell in row
        }
        for line in page["text"].splitlines():
            if "|" in line or line.strip() in table_text_lines or not line.strip():
                continue
            paragraph = document.add_paragraph(line)
            if line.startswith(("if ", "    ", "for ")):
                paragraph.style = "Code"
        for table_truth in page.get("tables") or []:
            rows = table_truth["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    table.cell(row_index, column_index).text = value
            for merge in table_truth.get("merged_cells") or []:
                start = table.cell(int(merge["row"]), int(merge["column"]))
                end = table.cell(
                    int(merge["row"]) + int(merge.get("row_span", 1)) - 1,
                    int(merge["column"]) + int(merge.get("column_span", 1)) - 1,
                )
                start.merge(end)
                start.text = rows[int(merge["row"])][int(merge["column"])]
        if legacy_fixture:
            caption = next(
                (
                    str(asset.get("caption") or "").strip()
                    for asset in page.get("assets") or []
                    if str(asset.get("caption") or "").strip()
                ),
                "Illustration",
            )
            image = _illustration_image()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as handle:
                temp_image = Path(handle.name)
            try:
                image.resize((320, 180)).save(temp_image)
                drawing = document.add_picture(str(temp_image))
                drawing._inline.docPr.set("descr", caption)
                document.add_paragraph(caption, style="Caption")
            finally:
                temp_image.unlink(missing_ok=True)
        if index + 1 < len(pages):
            document.add_page_break()
    document.save(destination)


def _convert(source: Path, destination_dir: Path, extension: str, filter_name: str | None = None) -> Path:
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        raise RuntimeError("LibreOffice/soffice là bắt buộc để build PDF/DOC golden fixtures")
    destination_dir.mkdir(parents=True, exist_ok=True)
    convert_to = extension if not filter_name else f"{extension}:{filter_name}"
    completed = subprocess.run(
        [converter, "--headless", "--convert-to", convert_to, "--outdir", str(destination_dir), str(source)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    output = destination_dir / f"{source.stem}.{extension}"
    if completed.returncode or not output.exists():
        raise RuntimeError((completed.stderr or completed.stdout)[-1000:])
    return output


def build(output: Path) -> None:
    resolved = output.resolve()
    forbidden = [(BACKEND_DIR / "data").resolve(), (BACKEND_DIR / "tests" / "golden_corpus" / "v1" / "truth").resolve()]
    if any(resolved == root or root in resolved.parents for root in forbidden):
        raise ValueError("Output benchmark không được nằm trong production data hoặc ground truth")
    output.mkdir(parents=True, exist_ok=True)
    truth = json.loads((CORPUS_DIR / "truth" / "pages.json").read_text(encoding="utf-8"))["pages"]
    by_number = {page["page_number"]: page for page in truth}

    born_docx = output / "born-digital-vi.docx"
    _write_docx(born_docx, [by_number[1], by_number[2]])
    born_pdf = _convert(born_docx, output, "pdf")
    born_pdf.replace(output / "born-digital-vi.pdf") if born_pdf.name != "born-digital-vi.pdf" else None
    _save_images_pdf(
        [
            _render_scan(by_number[3]["text"]),
            _add_graph_diagram(_render_scan(by_number[4]["text"]), caption="Sơ đồ đồ thị"),
        ],
        output / "scan-standard-vi.pdf",
    )
    _save_images_pdf(
        [_render_scan(by_number[5]["text"], degraded=True), _render_scan(by_number[6]["text"], degraded=True)],
        output / "scan-degraded-vi.pdf",
    )
    mixed_digital_docx = output / "mixed-digital.docx"
    _write_docx(mixed_digital_docx, [by_number[7]], legacy_fixture=True)
    mixed_digital_pdf = _convert(mixed_digital_docx, output, "pdf")
    mixed_scan_pdf = output / "mixed-scan.pdf"
    _save_images_pdf(
        [_add_graph_diagram(_render_scan(by_number[8]["text"]), caption="Sơ đồ duyệt BFS")],
        mixed_scan_pdf,
    )
    writer = PdfWriter()
    writer.add_page(PdfReader(str(mixed_digital_pdf)).pages[0])
    writer.add_page(PdfReader(str(mixed_scan_pdf)).pages[0])
    with (output / "mixed-pdf-vi.pdf").open("wb") as destination:
        writer.write(destination)

    _write_docx(output / "docx-vi.docx", [by_number[9]])
    legacy_docx = output / "doc-vi.docx"
    _write_docx(legacy_docx, [by_number[10]], legacy_fixture=True)
    _convert(legacy_docx, output, "doc", "MS Word 97")


def main() -> int:
    parser = argparse.ArgumentParser(description="Build independent golden corpus inputs")
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.output)
    print(json.dumps({"status": "built", "output": str(args.output.resolve())}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
