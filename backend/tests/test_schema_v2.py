import hashlib
import inspect
import re
import unittest
from contextlib import nullcontext
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

from bson import ObjectId
from docx import Document
from fastapi import HTTPException
from pydantic import ValidationError

from main import app
from core.bootstrap import SCHEMA_VERSION, VALIDATORS
from core.config import settings
from core.dependencies import CurrentUser
from core import job_recovery
from modules.admin import jobs_service as admin_jobs_module
from modules.admin.jobs_service import (
    AdminJobService,
    _generation_status_filter,
    _parse_object_id,
    _uppercase_status_filter,
)
from modules.admin.audit_service import AdminAuditService
from modules.admin.moodle_service import MoodleTargetService, _safe_publication_item
from modules.admin.moodle_schemas import MoodleTargetPayload
from modules.admin.overview_service import AdminOverviewService
from modules.catalog.schemas import (
    AiModelActivationPayload,
    ChapterPayload,
    ChapterUpdatePayload,
    EvaluationPolicyActivationPayload,
    LearningOutcomePayload,
    LearningOutcomeUpdatePayload,
    PromptTemplateActivationPayload,
    SubjectUpdatePayload,
)
from modules.catalog.service import CatalogService
from modules.notifications.service import NotificationService
from modules.auth import login as auth_login
from modules.documents.repository import MongoDocumentRepository
from modules.documents.schemas import DocumentPageUpdateRequest, DocumentStatus
from modules.documents.service import DocumentService
from modules.exams.service import ExamService, ExamVariantService
from modules.exams.pdf_service import render_exam_docx
from modules.ocr.ocr import extract_docx_pages
from modules.exams.schemas import (
    AddQuestionsManualRequest,
    ExamMatrixRequest,
    ExamStatusUpdateRequest,
    ExamUpdateRequest,
    ExamVariantCreateRequest,
)
from modules.generation.schemas import (
    BloomLevel,
    GeneratedQuestion,
    GenerationPlanSummary,
    QuestionGenerateRequest,
    QuestionType,
)
from modules.generation.llm.deepseek import DeepseekProvider
from modules.generation.llm.factory import get_llm_service
from modules.generation.prompt_builder import PromptBuilder
from modules.generation.question import _build_retry_prompt, _check_type_format
from modules.questions.schemas import (
    QuestionCreateRequest,
    QuestionResponse,
    QuestionSourceViewerResponse,
    QuestionUpdateRequest,
)
from modules.questions import repository as question_repository_module
from modules.questions.repository import MongoQuestionRepository
from modules.questions.service import QuestionService, stable_hash
from modules.questions.workflow_schemas import (
    AutoEvaluationRequest,
    MoodlePublicationRequest,
    ReviewAssignmentRequest,
    ReviewCreateRequest,
    ReviewOverride,
)
from modules.questions import workflow_service as question_workflow_module
from modules.questions.workflow_service import QuestionWorkflowService
from modules.rag.search import _heading_matches_target, _normalize_heading_text
from modules.users.schemas import (
    GenerationPresetPayload,
    PublicRegisterRequest,
    RoleEnum,
    UserAdminUpdateRequest,
)
from modules.users.service import UserService


def _current_user(role="Teacher", user_id=None):
    oid = user_id or ObjectId()
    return CurrentUser(
        id=oid,
        firebase_uid=f"firebase-{oid}",
        email=f"{role.lower()}@example.com",
        role=role,
        is_active=True,
    )


def _exam_doc(owner_id):
    now = datetime.now(timezone.utc)
    return {
        "_id": ObjectId(),
        "name": "Midterm",
        "exam_title": "Data Structures Midterm",
        "subject_id": ObjectId(),
        "question_count": 1,
        "header": {"school": "CTU"},
        "matrix": [],
        "questions": [],
        "status": "draft",
        "created_by_user_id": owner_id,
        "created_at": now,
        "updated_at": now,
    }


class FakeExamRepository:
    def __init__(self, exams=None, variants=None):
        self.exams = {str(exam["_id"]): exam for exam in (exams or [])}
        self.variants = variants or []

    def create(self, exam):
        self.exams[str(exam["_id"])] = exam
        return exam

    def find(self, exam_id):
        return self.exams.get(str(exam_id))

    def list(self, page, page_size, created_by_user_id):
        items = list(self.exams.values())
        if created_by_user_id is not None:
            items = [
                exam for exam in items
                if str(exam.get("created_by_user_id")) == str(created_by_user_id)
            ]
        return items, len(items)

    def update(self, exam_id, updates):
        exam = self.find(exam_id)
        if not exam:
            return None
        exam.update(updates)
        exam["updated_at"] = datetime.now(timezone.utc)
        return exam

    def delete(self, exam_id):
        return self.exams.pop(str(exam_id), None) is not None

    def count_variants(self, exam_id):
        return len([v for v in self.variants if str(v["exam_id"]) == str(exam_id)])


class FakeExamVariantRepository:
    def __init__(self, variants=None):
        self.variants = {str(variant["_id"]): variant for variant in (variants or [])}

    def create(self, variant):
        self.variants[str(variant["_id"])] = variant
        return variant

    def find(self, variant_id):
        return self.variants.get(str(variant_id))

    def list_by_exam(self, exam_id):
        return [
            variant for variant in self.variants.values()
            if str(variant["exam_id"]) == str(exam_id)
        ]

    def delete(self, variant_id):
        return self.variants.pop(str(variant_id), None) is not None


class FakeExamQuestionRepository:
    def __init__(self, pairs=None):
        self.pairs = {str(question["_id"]): (question, version) for question, version in (pairs or [])}

    def find_pair(self, question_id):
        return self.pairs.get(str(question_id))

    def list(self, page, page_size, review_status, search, **filters):
        pairs = list(self.pairs.values())
        if review_status:
            pairs = [
                pair for pair in pairs
                if pair[0].get("review_status") == review_status
            ]
        subject_id = filters.get("subject_id")
        if subject_id:
            pairs = [
                pair for pair in pairs
                if str((pair[1].get("classification") or {}).get("subject", {}).get("id")) == str(subject_id)
            ]
        if filters.get("approved_current_only"):
            pairs = [
                pair for pair in pairs
                if str(pair[0].get("approved_version_id")) == str(pair[0].get("current_version_id"))
            ]
        total = len(pairs)
        start = (page - 1) * page_size
        return pairs[start:start + page_size], total


def _user_doc(role="Admin", is_active=True):
    now = datetime.now(timezone.utc)
    oid = ObjectId()
    return {
        "_id": oid,
        "firebase_uid": f"firebase-{oid}",
        "email": f"{role.lower()}-{oid}@example.com",
        "display_name": f"{role} User",
        "role": role,
        "profile": {"school": "", "address": "", "avatar": ""},
        "is_active": is_active,
        "created_at": now,
        "updated_at": now,
    }


class FakeUserRepository:
    def __init__(self, users):
        self.users = {str(user["_id"]): dict(user) for user in users}

    def find_by_id(self, user_id):
        return self.users.get(str(user_id))

    def list(self, page, page_size, role, search):
        records = list(self.users.values())
        if role:
            records = [user for user in records if user.get("role") == role]
        if search:
            search_text = search.lower()
            records = [
                user for user in records
                if search_text in user.get("display_name", "").lower()
                or search_text in user.get("email", "").lower()
            ]
        total = len(records)
        start = (page - 1) * page_size
        return records[start:start + page_size], total

    def update(self, user_id, fields):
        user = self.find_by_id(user_id)
        if not user:
            return None
        user.update(fields)
        user["updated_at"] = datetime.now(timezone.utc)
        return user

    def count_active_admins(self):
        return len([
            user for user in self.users.values()
            if user.get("role") == "Admin" and user.get("is_active", True)
        ])


class FakeIdentityGateway:
    def __init__(self):
        self.disabled_calls = []
        self.update_calls = []

    def create_user(self, *, email, password, display_name):
        raise NotImplementedError

    def update_user(self, firebase_uid, **fields):
        self.update_calls.append((firebase_uid, fields))

    def set_user_disabled(self, firebase_uid, disabled):
        self.disabled_calls.append((firebase_uid, disabled))

    def delete_user(self, firebase_uid):
        raise NotImplementedError


class FakeSessions:
    def __init__(self):
        self.upserts = []

    def upsert(self, firebase_uid, id_token):
        self.upserts.append((firebase_uid, id_token))


def _path_values(record, path):
    parts = path.split(".")

    def walk(value, remaining):
        if not remaining:
            return [value]
        if isinstance(value, list):
            results = []
            for item in value:
                results.extend(walk(item, remaining))
            return results
        if not isinstance(value, dict):
            return []
        key = remaining[0]
        if key not in value:
            return []
        return walk(value[key], remaining[1:])

    return walk(record, parts)


def _matches_query(record, query):
    for key, expected in (query or {}).items():
        if key == "$or":
            if not any(_matches_query(record, item) for item in expected):
                return False
            continue
        if key == "$and":
            if not all(_matches_query(record, item) for item in expected):
                return False
            continue
        values = _path_values(record, key)
        if isinstance(expected, dict):
            for operator, operand in expected.items():
                if operator == "$in":
                    if not any(value in operand for value in values):
                        return False
                elif operator == "$lt":
                    if not any(value is not None and value < operand for value in values):
                        return False
                elif operator == "$lte":
                    if not any(value is not None and value <= operand for value in values):
                        return False
                elif operator == "$gt":
                    if not any(value is not None and value > operand for value in values):
                        return False
                elif operator == "$gte":
                    if not any(value is not None and value >= operand for value in values):
                        return False
                elif operator == "$exists":
                    if bool(values) is not bool(operand):
                        return False
                elif operator == "$regex":
                    flags = re.IGNORECASE if expected.get("$options") == "i" else 0
                    if not any(re.search(operand, str(value), flags) for value in values if value is not None):
                        return False
                elif operator == "$options":
                    continue
                else:
                    return False
        elif expected is None:
            if values and not any(value is None for value in values):
                return False
        elif not any(value == expected for value in values):
            return False
    return True


class InMemoryCursor(list):
    def sort(self, key_or_list, direction=1):
        if isinstance(key_or_list, list):
            sort_keys = key_or_list
        else:
            sort_keys = [(key_or_list, direction)]
        for key, sort_direction in reversed(sort_keys):
            self[:] = sorted(
                self,
                key=lambda item: (_path_values(item, key) or [None])[0],
                reverse=sort_direction == -1,
            )
        return self

    def skip(self, count):
        return InMemoryCursor(self[count:])

    def limit(self, count):
        return InMemoryCursor(self[:count])


class InMemoryCollection:
    def __init__(self, records=None):
        self.records = [dict(record) for record in (records or [])]

    def count_documents(self, query):
        return len([record for record in self.records if _matches_query(record, query)])

    def find(self, query=None, *_args, **_kwargs):
        return InMemoryCursor([record for record in self.records if _matches_query(record, query)])

    def find_one(self, query=None, *_args, **_kwargs):
        return next((record for record in self.records if _matches_query(record, query)), None)

    @staticmethod
    def _set_path(record, path, value, filter_query):
        if ".$." not in path:
            current = record
            parts = path.split(".")
            for part in parts[:-1]:
                current = current.setdefault(part, {})
            current[parts[-1]] = value
            return
        collection_key, field_key = path.split(".$.", 1)
        target_id = filter_query.get(f"{collection_key}._id")
        for item in record.get(collection_key, []):
            if item.get("_id") == target_id:
                item[field_key] = value
                return

    def insert_one(self, record, *_args, **_kwargs):
        item = dict(record)
        self.records.append(item)
        return type("Result", (), {"inserted_id": item.get("_id")})()

    def find_one_and_update(
        self,
        filter_query,
        update,
        *,
        upsert=False,
        return_document=None,
        **_kwargs,
    ):
        record = self.find_one(filter_query)
        if record is None and upsert:
            record = dict((update.get("$setOnInsert") or {}))
            record.update(update.get("$set") or {})
            self.records.append(record)
        if record is None:
            return None
        for path, value in (update.get("$set") or {}).items():
            self._set_path(record, path, value, filter_query)
        for path, value in (update.get("$push") or {}).items():
            record.setdefault(path, []).append(value)
        return record

    def update_one(self, filter_query, update, *, upsert=False, **_kwargs):
        matched = 0
        modified = 0
        record = self.find_one(filter_query)
        if record is None and upsert:
            record = dict(update.get("$setOnInsert") or {})
            self.records.append(record)
            matched = 1
            modified = 1
        elif record is not None:
            matched = 1

        if record is not None:
            changed = False
            for path, value in (update.get("$set") or {}).items():
                self._set_path(record, path, value, filter_query)
                changed = True
            for path, value in (update.get("$push") or {}).items():
                record.setdefault(path, []).append(value)
                changed = True
            if changed:
                modified = 1

        return type(
            "Result",
            (),
            {"matched_count": matched, "modified_count": modified},
        )()

    def update_many(self, filter_query, update, **_kwargs):
        matched = 0
        modified = 0
        for record in self.records:
            if not _matches_query(record, filter_query):
                continue
            matched += 1
            changed = False
            for path, value in (update.get("$set") or {}).items():
                self._set_path(record, path, value, filter_query)
                changed = True
            for path, value in (update.get("$push") or {}).items():
                record.setdefault(path, []).append(value)
                changed = True
            if changed:
                modified += 1

        return type(
            "Result",
            (),
            {"matched_count": matched, "modified_count": modified},
        )()


class FakeCatalogDatabase:
    def __init__(
        self,
        *,
        subjects=None,
        documents=None,
        question_versions=None,
        questions=None,
        exams=None,
        ai_models=None,
        prompt_templates=None,
        evaluation_policies=None,
    ):
        self.subjects = InMemoryCollection(subjects)
        self.documents = InMemoryCollection(documents)
        self.question_versions = InMemoryCollection(question_versions)
        self.questions = InMemoryCollection(questions)
        self.exams = InMemoryCollection(exams)
        self.ai_models = InMemoryCollection(ai_models)
        self.prompt_templates = InMemoryCollection(prompt_templates)
        self.evaluation_policies = InMemoryCollection(evaluation_policies)


class SchemaV2Tests(unittest.TestCase):
    def test_only_admin_teacher_and_reviewer_roles_exist(self):
        self.assertEqual({role.value for role in RoleEnum}, {"Admin", "Teacher", "Reviewer"})

    def test_public_register_cannot_choose_role(self):
        with self.assertRaises(ValidationError):
            PublicRegisterRequest(
                email="teacher@example.com",
                password="secret1",
                full_name="Teacher",
                role="Admin",
            )

    def test_question_update_requires_optimistic_version(self):
        with self.assertRaises(ValidationError):
            QuestionUpdateRequest(content="Updated")

    def test_question_can_reference_multiple_chunks(self):
        payload = QuestionCreateRequest(
            content="Question",
            source_chunk_ids=["507f1f77bcf86cd799439011", "507f1f77bcf86cd799439012"],
        )
        self.assertEqual(len(payload.source_chunk_ids), 2)

    def test_question_restore_creates_new_version_and_resets_workflow(self):
        owner = _current_user("Teacher")
        question_id = ObjectId()
        version1_id = ObjectId()
        version2_id = ObjectId()
        review_id = ObjectId()
        reviewer_id = ObjectId()
        now = datetime.now(timezone.utc)
        classification = {
            "subject": {"id": None},
            "chapter": {"id": None},
            "assessment_type": "TRAC_NGHIEM",
            "bloom": {"level": 2, "code": "UNDERSTAND", "name": "Understand"},
            "difficulty": None,
        }
        version1_data = {
            "content": "Original stack question",
            "question_data": {
                "options": {"A": "LIFO", "B": "FIFO"},
                "correct_answer": "A",
                "explanation": "Stack pops the newest item first.",
            },
        }
        version2_data = {
            "content": "Edited queue question",
            "question_data": {
                "options": {"A": "LIFO", "B": "FIFO"},
                "correct_answer": "B",
                "explanation": "Queue removes the oldest item first.",
            },
        }
        versions = [
            {
                "_id": version1_id,
                "schema_version": SCHEMA_VERSION,
                "question_id": question_id,
                "version": 1,
                "origin": "MANUAL",
                "generation_run_id": None,
                "document_id": None,
                "created_by_user_id": owner.id,
                "generated_by_model_id": None,
                "classification": classification,
                "clos": [],
                "sources": [],
                "keywords": [],
                "content_hash": stable_hash(
                    {**version1_data, "classification": classification, "clos": [], "sources": []}
                ),
                "change_note": "Initial version",
                "created_at": now - timedelta(minutes=2),
                **version1_data,
            },
            {
                "_id": version2_id,
                "schema_version": SCHEMA_VERSION,
                "question_id": question_id,
                "version": 2,
                "origin": "MANUAL",
                "generation_run_id": None,
                "document_id": None,
                "created_by_user_id": owner.id,
                "generated_by_model_id": None,
                "classification": classification,
                "clos": [],
                "sources": [],
                "keywords": [],
                "content_hash": stable_hash(
                    {**version2_data, "classification": classification, "clos": [], "sources": []}
                ),
                "change_note": "Teacher edit",
                "created_at": now - timedelta(minutes=1),
                **version2_data,
            },
        ]
        question = {
            "_id": question_id,
            "schema_version": SCHEMA_VERSION,
            "question_code": "Q-RESTORE",
            "current_version": 2,
            "current_version_id": version2_id,
            "approved_version_id": version2_id,
            "lifecycle_status": "ACTIVE",
            "evaluation_status": "PASSED",
            "review_status": "APPROVED",
            "publication_status": "PUBLISHED",
            "quality_summary": {"overall_score": 0.92, "color": "GREEN"},
            "review_assignment": {
                "status": "CLAIMED",
                "reviewer_user_id": reviewer_id,
                "assigned_by_user_id": ObjectId(),
                "assigned_at": now - timedelta(minutes=5),
                "claimed_at": now - timedelta(minutes=4),
                "lock_expires_at": now + timedelta(minutes=30),
                "last_released_at": None,
                "release_reason": None,
            },
            "latest_review_id": review_id,
            "created_by_user_id": owner.id,
            "created_at": now - timedelta(minutes=3),
            "updated_at": now - timedelta(minutes=1),
            "archived_at": None,
        }
        db = FakeCatalogDatabase(questions=[question], question_versions=versions)
        original_transaction = question_repository_module.mongo_transaction
        try:
            question_repository_module.mongo_transaction = lambda: nullcontext(None)
            service = QuestionService(MongoQuestionRepository(db), references=None)

            history = service.versions(str(question_id), owner)
            restored = service.update(
                str(question_id),
                QuestionUpdateRequest(
                    expected_version=2,
                    content=version1_data["content"],
                    question_type=classification["assessment_type"],
                    bloom_level=classification["bloom"]["level"],
                    question_data=version1_data["question_data"],
                    source_chunk_ids=[],
                    clo_ids=[],
                    change_note="Restore from version 1",
                ),
                owner.id,
                actor_role=owner.role,
            )
            restored_history = service.versions(str(question_id), owner)
        finally:
            question_repository_module.mongo_transaction = original_transaction

        self.assertEqual([item["version"] for item in history], [2, 1])
        self.assertEqual(restored["current_version"], 3)
        self.assertEqual(restored["content"], version1_data["content"])
        self.assertEqual(restored["question_data"], version1_data["question_data"])
        self.assertEqual(restored["evaluation_status"], "NOT_STARTED")
        self.assertEqual(restored["review_status"], "PENDING")
        self.assertEqual(restored["publication_status"], "STALE")
        self.assertEqual(restored["quality_summary"], {})
        self.assertEqual(restored["review_assignment"]["status"], "UNASSIGNED")
        self.assertEqual([item["version"] for item in restored_history], [3, 2, 1])
        self.assertEqual(restored_history[0]["change_note"], "Restore from version 1")
        self.assertEqual(restored_history[0]["content"], version1_data["content"])

    def test_review_override_requires_reason(self):
        with self.assertRaises(ValidationError):
            ReviewOverride(applied=True, score=0.9)

    def test_structured_review_requires_reject_reason_and_revision_issue(self):
        with self.assertRaises(ValidationError):
            ReviewCreateRequest(expected_version=1, decision="REJECTED")

        with self.assertRaises(ValidationError):
            ReviewCreateRequest(
                expected_version=1,
                decision="NEEDS_REVISION",
                review_form={"overall_note": "Cần sửa đáp án."},
            )

        payload = ReviewCreateRequest(
            expected_version=1,
            decision="NEEDS_REVISION",
            review_form={
                "overall_note": "Cần sửa trước khi duyệt.",
                "revision_issues": [
                    {
                        "title": "Đáp án chưa khớp nguồn",
                        "severity": "HIGH",
                        "detail": "Nguồn nói FIFO nhưng đáp án chọn Stack.",
                        "page_number": 2,
                    }
                ],
            },
        )
        self.assertEqual(payload.review_form.revision_issues[0].severity, "HIGH")

    def test_review_assignment_request_allows_admin_unassign(self):
        payload = ReviewAssignmentRequest(reviewer_user_id=None, note="unassign")
        self.assertIsNone(payload.reviewer_user_id)
        self.assertEqual(payload.note, "unassign")

    def test_question_response_exposes_review_assignment(self):
        now = datetime.now(timezone.utc)
        response = QuestionResponse(
            id=str(ObjectId()),
            question_code="Q-1",
            current_version=1,
            current_version_id=str(ObjectId()),
            approved_version_id=None,
            document_id=None,
            lifecycle_status="ACTIVE",
            evaluation_status="PASSED",
            review_status="PENDING",
            publication_status="NOT_PUBLISHED",
            content="Question",
            question_data={},
            classification={},
            clos=[],
            sources=[],
            content_hash="hash",
            quality_summary={},
            review_assignment={
                "status": "IN_REVIEW",
                "reviewer_user_id": str(ObjectId()),
            },
            latest_review_id=None,
            created_at=now,
            updated_at=now,
        )
        self.assertEqual(response.review_assignment["status"], "IN_REVIEW")

    def test_question_source_viewer_schema_accepts_chunk_pages(self):
        response = QuestionSourceViewerResponse(
            question_id=str(ObjectId()),
            question_code="Q-1",
            version_id=str(ObjectId()),
            version=1,
            document={
                "id": str(ObjectId()),
                "title": "Document",
                "original_filename": "source.pdf",
                "page_count": 4,
                "current_ocr_job_id": str(ObjectId()),
                "current_chunk_set_id": str(ObjectId()),
                "pdf_available": True,
                "pdf_url": "/questions/q/source-pdf",
            },
            items=[
                {
                    "citation_order": 1,
                    "source_type": "CHUNK",
                    "is_primary": True,
                    "chunk_id": str(ObjectId()),
                    "page_range": {"start": 2, "end": 3, "pages": [2, 3]},
                    "context_excerpt": "Queue follows FIFO.",
                    "pages": [{"page_number": 2, "text": "Queue follows FIFO."}],
                }
            ],
        )
        self.assertEqual(response.items[0].pages[0].page_number, 2)

    def test_question_source_viewer_reports_stale_chunk_set_and_pages(self):
        question_id = ObjectId()
        version_id = ObjectId()
        document_id = ObjectId()
        old_chunk_set_id = ObjectId()
        current_chunk_set_id = ObjectId()
        ocr_job_id = ObjectId()
        chunk_id = ObjectId()
        now = datetime.now(timezone.utc)
        question = {
            "_id": question_id,
            "question_code": "Q-1",
            "current_version_id": version_id,
            "current_version": 1,
            "created_by_user_id": ObjectId(),
        }
        version = {
            "_id": version_id,
            "version": 1,
            "document_id": document_id,
            "created_by_user_id": question["created_by_user_id"],
            "sources": [
                {
                    "source_type": "CHUNK",
                    "chunk_id": chunk_id,
                    "chunk_set_id": old_chunk_set_id,
                    "chunk_content_hash": "hash-a",
                    "citation_order": 1,
                    "is_primary": True,
                    "context_excerpt": "Queue follows FIFO.",
                }
            ],
        }
        document = {
            "_id": document_id,
            "title": "Data Structures",
            "original_filename": "ctdl.pdf",
            "page_count": 5,
            "current_processing": {
                "ocr_job_id": ocr_job_id,
                "chunk_set_id": current_chunk_set_id,
            },
            "artifacts": [
                {
                    "type": "ORIGINAL_PDF",
                    "is_current": True,
                    "storage": {"provider": "LOCAL", "uri": "data/uploads/source.pdf"},
                    "mime_type": "application/pdf",
                }
            ],
        }
        chunk = {
            "_id": chunk_id,
            "document_id": document_id,
            "chunk_set_id": old_chunk_set_id,
            "chunk_no": 7,
            "content": "Queue follows FIFO.",
            "content_hash": "hash-a",
            "page_range": {"start": 2, "end": 3, "pages": [2, 3]},
            "heading": {"title": "Queue"},
        }

        class PairRepository:
            def find_pair(self, _question_id):
                return question, version

        class References:
            def find_chunk(self, requested_chunk_id):
                return chunk if requested_chunk_id == chunk_id else None

            def find_document(self, requested_document_id):
                return document if requested_document_id == document_id else None

            def find_pages(self, requested_document_id, requested_ocr_job_id, page_numbers):
                self.page_query = (requested_document_id, requested_ocr_job_id, page_numbers)
                return [
                    {
                        "page_number": 2,
                        "cleaned_text": "Page 2 Queue follows FIFO.",
                        "formula_blocks": [],
                    }
                ]

            def find_subject(self, _subject_id):
                return None

        references = References()
        service = QuestionService(repository=PairRepository(), references=references)
        result = service.source_viewer(str(question_id), _current_user("Reviewer"))

        self.assertEqual(result["document"]["title"], "Data Structures")
        self.assertTrue(result["document"]["pdf_available"])
        self.assertEqual(result["items"][0]["chunk_no"], 7)
        self.assertFalse(result["items"][0]["is_current_chunk_set"])
        self.assertIn("Nguồn không còn thuộc chunk set hiện hành", result["items"][0]["warnings"])
        self.assertEqual(result["items"][0]["pages"][0]["page_number"], 2)
        self.assertEqual(references.page_query, (document_id, ocr_job_id, [2, 3]))

    def test_submit_for_review_moves_draft_or_revision_to_pending(self):
        teacher = _current_user("Teacher")

        class SubmitRepository:
            def __init__(self, review_status):
                question_id = ObjectId()
                version_id = ObjectId()
                now = datetime.now(timezone.utc)
                self.update_calls = 0
                self.question = {
                    "_id": question_id,
                    "schema_version": SCHEMA_VERSION,
                    "question_code": "Q-SUBMIT",
                    "current_version": 1,
                    "current_version_id": version_id,
                    "approved_version_id": None,
                    "lifecycle_status": "ACTIVE",
                    "evaluation_status": "NOT_STARTED",
                    "review_status": review_status,
                    "publication_status": "NOT_PUBLISHED",
                    "quality_summary": {},
                    "review_assignment": {"status": "UNASSIGNED"},
                    "latest_review_id": None,
                    "created_by_user_id": teacher.id,
                    "created_at": now,
                    "updated_at": now,
                    "archived_at": None,
                }
                self.version = {
                    "_id": version_id,
                    "schema_version": SCHEMA_VERSION,
                    "question_id": question_id,
                    "version": 1,
                    "origin": "MANUAL",
                    "generation_run_id": None,
                    "document_id": None,
                    "created_by_user_id": teacher.id,
                    "generated_by_model_id": None,
                    "classification": {
                        "subject": {"id": None},
                        "chapter": {"id": None},
                        "assessment_type": "TRAC_NGHIEM",
                        "bloom": {"level": 2},
                        "difficulty": None,
                    },
                    "clos": [],
                    "content": "Queue follows FIFO.",
                    "question_data": {"options": {"A": "FIFO", "B": "LIFO"}, "correct_answer": "A"},
                    "sources": [],
                    "keywords": [],
                    "content_hash": "hash-submit",
                    "change_note": "Initial version",
                    "created_at": now,
                }

            def find_pair(self, _question_id):
                return self.question, self.version

            def update_review_status(self, _question_id, allowed_statuses, review_status):
                self.update_calls += 1
                if self.question["review_status"] not in allowed_statuses:
                    return None
                self.question["review_status"] = review_status
                self.question["updated_at"] = datetime.now(timezone.utc)
                return self.question, self.version

        for initial_status in ("DRAFT", "NEEDS_REVISION"):
            repository = SubmitRepository(initial_status)
            service = QuestionService(repository=repository, references=object())

            submitted = service.submit_for_review(str(repository.question["_id"]), teacher)
            submitted_again = service.submit_for_review(str(repository.question["_id"]), teacher)

            self.assertEqual(submitted["review_status"], "PENDING")
            self.assertEqual(submitted_again["review_status"], "PENDING")
            self.assertEqual(repository.update_calls, 1)

        repository = SubmitRepository("APPROVED")
        service = QuestionService(repository=repository, references=object())
        with self.assertRaises(ValueError):
            service.submit_for_review(str(repository.question["_id"]), teacher)
        self.assertEqual(repository.update_calls, 0)

    def test_duplicate_question_creates_draft_copy_without_review_state(self):
        teacher = _current_user("Teacher")
        question_id = ObjectId()
        version_id = ObjectId()
        now = datetime.now(timezone.utc)
        source_question = {
            "_id": question_id,
            "schema_version": SCHEMA_VERSION,
            "question_code": "Q-SOURCE",
            "current_version": 3,
            "current_version_id": version_id,
            "approved_version_id": version_id,
            "lifecycle_status": "ACTIVE",
            "evaluation_status": "PASSED",
            "review_status": "APPROVED",
            "publication_status": "PUBLISHED",
            "quality_summary": {"color": "GREEN"},
            "review_assignment": {"status": "ASSIGNED"},
            "latest_review_id": ObjectId(),
            "created_by_user_id": teacher.id,
            "created_at": now,
            "updated_at": now,
            "archived_at": None,
        }
        source_version = {
            "_id": version_id,
            "schema_version": SCHEMA_VERSION,
            "question_id": question_id,
            "version": 3,
            "origin": "AI",
            "generation_run_id": ObjectId(),
            "document_id": None,
            "created_by_user_id": teacher.id,
            "generated_by_model_id": ObjectId(),
            "classification": {
                "subject": {"id": None},
                "chapter": {"id": None},
                "assessment_type": "TRAC_NGHIEM",
                "bloom": {"level": 2},
                "difficulty": "de",
            },
            "clos": [{"code": "CLO1"}],
            "content": "Queue follows FIFO.",
            "question_data": {"options": {"A": "FIFO", "B": "LIFO"}, "correct_answer": "A"},
            "sources": [],
            "keywords": ["queue"],
            "content_hash": "source-hash",
            "change_note": "Approved version",
            "created_at": now,
        }

        class DuplicateRepository:
            def __init__(self):
                self.created_question = None
                self.created_version = None

            def find_pair(self, _question_id):
                return source_question, source_version

            def create(self, aggregate, version):
                self.created_question = aggregate
                self.created_version = version
                return aggregate, version

        repository = DuplicateRepository()
        service = QuestionService(repository=repository, references=object())

        duplicated = service.duplicate(str(question_id), teacher)

        self.assertNotEqual(duplicated["id"], str(question_id))
        self.assertEqual(duplicated["review_status"], "DRAFT")
        self.assertEqual(duplicated["evaluation_status"], "NOT_STARTED")
        self.assertEqual(duplicated["publication_status"], "NOT_PUBLISHED")
        self.assertIsNone(duplicated["approved_version_id"])
        self.assertEqual(duplicated["content"], source_version["content"])
        self.assertEqual(repository.created_question["created_by_user_id"], teacher.id)
        self.assertEqual(repository.created_version["origin"], "MANUAL")
        self.assertEqual(source_question["review_status"], "APPROVED")

    def test_reviewer_must_hold_active_review_lock(self):
        service = QuestionWorkflowService(database=None)
        reviewer_id = ObjectId()
        reviewer = _current_user("Reviewer", reviewer_id)
        now = datetime.now(timezone.utc)

        with self.assertRaises(PermissionError):
            service._ensure_review_lock(
                {"review_assignment": {"status": "ASSIGNED", "reviewer_user_id": reviewer_id}},
                reviewer,
                now,
            )

        service._ensure_review_lock(
            {
                "review_assignment": {
                    "status": "IN_REVIEW",
                    "reviewer_user_id": reviewer_id,
                    "lock_expires_at": (now + timedelta(minutes=1)).replace(tzinfo=None),
                }
            },
            reviewer,
            now,
        )

        with self.assertRaises(ValueError):
            service._ensure_review_lock(
                {
                    "review_assignment": {
                        "status": "IN_REVIEW",
                        "reviewer_user_id": reviewer_id,
                        "lock_expires_at": now - timedelta(minutes=1),
                    }
                },
                reviewer,
                now,
            )

    def test_review_assignment_claim_and_review_enforce_lock_owner(self):
        admin = _current_user("Admin")
        teacher = _current_user("Teacher")
        reviewer = _current_user("Reviewer")
        other_reviewer = _current_user("Reviewer")
        question_id = ObjectId()
        version_id = ObjectId()
        now = datetime.now(timezone.utc)
        version = {
            "_id": version_id,
            "schema_version": SCHEMA_VERSION,
            "question_id": question_id,
            "version": 1,
            "origin": "MANUAL",
            "generation_run_id": None,
            "document_id": None,
            "created_by_user_id": teacher.id,
            "generated_by_model_id": None,
            "classification": {
                "subject": {"id": None},
                "chapter": {"id": None},
                "assessment_type": "TRAC_NGHIEM",
                "bloom": {"level": 2},
                "difficulty": None,
            },
            "clos": [],
            "content": "Stack follows which rule?",
            "question_data": {
                "options": {"A": "LIFO", "B": "FIFO"},
                "correct_answer": "A",
                "explanation": "Stack is last-in-first-out.",
            },
            "sources": [],
            "keywords": [],
            "content_hash": "hash-review-lock",
            "change_note": "Initial version",
            "created_at": now,
        }
        question = {
            "_id": question_id,
            "schema_version": SCHEMA_VERSION,
            "question_code": "Q-LOCK",
            "current_version": 1,
            "current_version_id": version_id,
            "approved_version_id": None,
            "lifecycle_status": "ACTIVE",
            "evaluation_status": "PASSED",
            "review_status": "PENDING",
            "publication_status": "NOT_PUBLISHED",
            "quality_summary": {"overall_score": 0.91, "color": "GREEN"},
            "review_assignment": {"status": "UNASSIGNED"},
            "latest_review_id": None,
            "created_by_user_id": teacher.id,
            "created_at": now,
            "updated_at": now,
            "archived_at": None,
        }

        class FakeWorkflowDatabase:
            def __init__(self):
                self.questions = InMemoryCollection([question])
                self.question_versions = InMemoryCollection([version])
                self.users = InMemoryCollection(
                    [
                        {"_id": reviewer.id, "role": "Reviewer", "is_active": True},
                        {"_id": other_reviewer.id, "role": "Reviewer", "is_active": True},
                    ]
                )
                self.question_reviews = InMemoryCollection()
                self.audit_logs = InMemoryCollection()
                self.notifications = InMemoryCollection()

        db = FakeWorkflowDatabase()
        original_transaction = question_workflow_module.mongo_transaction
        try:
            question_workflow_module.mongo_transaction = lambda: nullcontext(None)
            service = QuestionWorkflowService(db)

            assigned = service.assign_review(
                str(question_id),
                ReviewAssignmentRequest(
                    reviewer_user_id=str(reviewer.id),
                    note="Assign to reviewer",
                ),
                admin,
            )
            with self.assertRaises(PermissionError):
                service.claim_review(str(question_id), other_reviewer)

            claimed = service.claim_review(str(question_id), reviewer)
            with self.assertRaises(PermissionError):
                service.release_review(str(question_id), other_reviewer)
            with self.assertRaises(PermissionError):
                service.review(
                    str(question_id),
                    ReviewCreateRequest(expected_version=1, decision="APPROVED"),
                    other_reviewer,
                )

            review = service.review(
                str(question_id),
                ReviewCreateRequest(expected_version=1, decision="APPROVED"),
                reviewer,
            )
        finally:
            question_workflow_module.mongo_transaction = original_transaction

        self.assertEqual(assigned["review_assignment"]["status"], "ASSIGNED")
        self.assertEqual(assigned["review_assignment"]["reviewer_user_id"], str(reviewer.id))
        self.assertEqual(claimed["review_assignment"]["status"], "IN_REVIEW")
        self.assertEqual(claimed["review_assignment"]["reviewer_user_id"], str(reviewer.id))
        self.assertEqual(review["decision"], "APPROVED")
        updated_question = db.questions.find_one({"_id": question_id})
        self.assertEqual(updated_question["review_status"], "APPROVED")
        self.assertEqual(updated_question["approved_version_id"], version_id)
        self.assertEqual(updated_question["review_assignment"]["status"], "UNASSIGNED")
        self.assertEqual(len(db.question_reviews.records), 1)
        self.assertEqual(len(db.notifications.records), 2)
        self.assertEqual(db.notifications.records[0]["recipient_user_id"], reviewer.id)
        self.assertEqual(db.notifications.records[0]["type"], "QUESTION_REVIEW_ASSIGNED")
        self.assertEqual(db.notifications.records[0]["link"], f"/kiem-duyet?questionId={question_id}")
        self.assertEqual(db.notifications.records[1]["recipient_user_id"], teacher.id)
        self.assertEqual(db.notifications.records[1]["type"], "QUESTION_APPROVED")
        self.assertEqual(
            [event["action"] for event in db.audit_logs.records],
            ["QUESTION_REVIEW_ASSIGNED", "QUESTION_REVIEW_CLAIMED", "QUESTION_APPROVED"],
        )

    def test_auto_evaluation_requires_expected_version(self):
        with self.assertRaises(ValidationError):
            AutoEvaluationRequest()

        payload = AutoEvaluationRequest(expected_version=1)
        self.assertEqual(payload.evaluator_model_code, settings.evaluation_model_provider)
        self.assertFalse(payload.fallback_to_heuristic)

    def test_llm_factory_accepts_ollama_model_alias(self):
        provider = get_llm_service("ollama:qwen2.5:7b")
        self.assertIsInstance(provider, DeepseekProvider)
        self.assertEqual(provider.model_name, "qwen2.5:7b")

    def test_moodle_publication_request_has_demo_defaults(self):
        payload = MoodlePublicationRequest(expected_version=1)
        self.assertEqual(payload.moodle_site_id, "demo-moodle")
        self.assertIsNone(payload.target_id)
        self.assertTrue(payload.mock)

    def test_moodle_publication_item_marks_mock_record(self):
        record = {
            "_id": ObjectId(),
            "question_id": ObjectId(),
            "question_version_id": ObjectId(),
            "question_version": 2,
            "publisher_user_id": ObjectId(),
            "target": {
                "moodle_site_id": "demo-moodle",
                "site_name": "Demo Moodle",
                "mode": "REST_API",
                "configured_mode": "REST_API",
            },
            "status": "PUBLISHED",
            "moodle_question_ref_id": "mock-demo-version",
            "request_payload": {
                "question_code": "Q-001",
                "export_format": "BOTH",
                "mock": True,
            },
            "response_payload": {
                "message": "Mô phỏng Moodle: chỉ ghi nhận publication cục bộ, chưa gửi dữ liệu sang Moodle thật.",
                "export_formats": ["gift", "xml"],
            },
        }

        item = _safe_publication_item(record)

        self.assertEqual(item["publication_mode"], "MOCK")
        self.assertEqual(item["configured_mode"], "REST_API")
        self.assertFalse(item["external_sync"])
        self.assertEqual(item["status_detail"], "SIMULATED_LOCAL_RECORD")
        self.assertEqual(item["status_label"], "Đã ghi mô phỏng")
        self.assertIn("chưa gửi", item["message"])

    def test_moodle_publish_rejects_real_sync_until_configured(self):
        service = QuestionWorkflowService(object())
        payload = MoodlePublicationRequest(expected_version=1, mock=False)

        with self.assertRaises(ValueError) as ctx:
            service.publish_to_moodle(str(ObjectId()), payload, ObjectId())

        self.assertIn("chưa được cấu hình", str(ctx.exception))

    def test_moodle_mock_publish_is_disabled_in_production_non_demo(self):
        original_app_env = settings.app_env
        original_demo_mode = settings.demo_mode
        try:
            settings.app_env = "production"
            settings.demo_mode = False
            service = QuestionWorkflowService(object())
            payload = MoodlePublicationRequest(expected_version=1, mock=True)

            with self.assertRaises(ValueError) as ctx:
                service.publish_to_moodle(str(ObjectId()), payload, ObjectId())
        finally:
            settings.app_env = original_app_env
            settings.demo_mode = original_demo_mode

        self.assertIn("Mô phỏng Moodle bị tắt trong production", str(ctx.exception))

    def _moodle_retry_database(self, *, current_version=1, failed_version=1):
        question_id = ObjectId()
        failed_version_id = ObjectId()
        current_version_id = failed_version_id if current_version == failed_version else ObjectId()
        target_id = ObjectId()
        failed_id = ObjectId()
        now = datetime.now(timezone.utc)
        moodle_site_id = "demo-moodle"
        course_id = "ctdl"
        category_id = "qbank"
        failed_hash = "hash-v1"
        idempotency_key = hashlib.sha256(
            "|".join(
                [
                    moodle_site_id,
                    course_id,
                    category_id,
                    str(failed_version_id),
                    failed_hash,
                ]
            ).encode("utf-8")
        ).hexdigest()
        versions = [
            {
                "_id": failed_version_id,
                "question_id": question_id,
                "version": failed_version,
                "document_id": None,
                "content": "Queue xử lý phần tử theo nguyên tắc nào?",
                "question_data": {
                    "options": {"A": "FIFO", "B": "LIFO"},
                    "correct_answer": "A",
                    "explanation": "Queue là hàng đợi FIFO.",
                },
                "classification": {"assessment_type": "TRAC_NGHIEM"},
                "clos": [],
                "sources": [],
                "content_hash": failed_hash,
            }
        ]
        if current_version_id != failed_version_id:
            versions.append(
                {
                    **versions[0],
                    "_id": current_version_id,
                    "version": current_version,
                    "content_hash": "hash-current",
                }
            )

        class FakeRetryDatabase:
            def __init__(self):
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "schema_version": SCHEMA_VERSION,
                            "question_code": "Q-001",
                            "current_version": current_version,
                            "current_version_id": current_version_id,
                            "approved_version_id": current_version_id,
                            "lifecycle_status": "ACTIVE",
                            "evaluation_status": "PASSED",
                            "review_status": "APPROVED",
                            "publication_status": "FAILED",
                            "quality_summary": {},
                            "review_assignment": {"status": "UNASSIGNED"},
                            "latest_review_id": None,
                            "created_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.question_versions = InMemoryCollection(versions)
                self.moodle_targets = InMemoryCollection(
                    [
                        {
                            "_id": target_id,
                            "site_key": moodle_site_id,
                            "site_name": "Demo Moodle",
                            "mode": "REST_API",
                            "default_course_id": course_id,
                            "default_category_id": category_id,
                            "is_active": True,
                        }
                    ]
                )
                self.moodle_publications = InMemoryCollection(
                    [
                        {
                            "_id": failed_id,
                            "schema_version": SCHEMA_VERSION,
                            "question_id": question_id,
                            "question_version_id": failed_version_id,
                            "question_version": failed_version,
                            "publisher_user_id": ObjectId(),
                            "target": {
                                "target_id": target_id,
                                "moodle_site_id": moodle_site_id,
                                "site_name": "Demo Moodle",
                                "mode": "MOCK",
                                "configured_mode": "REST_API",
                                "course_id": course_id,
                                "category_id": category_id,
                            },
                            "publication_mode": "MOCK",
                            "external_sync": False,
                            "status_detail": None,
                            "published_content_hash": failed_hash,
                            "idempotency_key": idempotency_key,
                            "status": "FAILED",
                            "attempt_no": 1,
                            "moodle_question_ref_id": None,
                            "request_payload": {
                                "question_code": "Q-001",
                                "export_format": "BOTH",
                                "mock": True,
                            },
                            "response_payload": {},
                            "error": {"message": "Moodle timeout"},
                            "created_at": now,
                            "updated_at": now,
                            "published_at": None,
                        }
                    ]
                )

        return FakeRetryDatabase(), failed_id, question_id

    def test_moodle_publish_is_idempotent_for_same_version_and_target(self):
        question_id = ObjectId()
        version_id = ObjectId()
        target_id = ObjectId()
        publisher_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakePublicationDatabase:
            def __init__(self):
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "schema_version": SCHEMA_VERSION,
                            "question_code": "Q-001",
                            "current_version": 1,
                            "current_version_id": version_id,
                            "approved_version_id": version_id,
                            "lifecycle_status": "ACTIVE",
                            "evaluation_status": "PASSED",
                            "review_status": "APPROVED",
                            "publication_status": "NOT_PUBLISHED",
                            "quality_summary": {},
                            "review_assignment": {"status": "UNASSIGNED"},
                            "latest_review_id": None,
                            "created_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.question_versions = InMemoryCollection(
                    [
                        {
                            "_id": version_id,
                            "question_id": question_id,
                            "version": 1,
                            "document_id": None,
                            "content": "Queue xử lý phần tử theo nguyên tắc nào?",
                            "question_data": {
                                "options": {"A": "FIFO", "B": "LIFO"},
                                "correct_answer": "A",
                                "explanation": "Queue là hàng đợi FIFO.",
                            },
                            "classification": {"assessment_type": "TRAC_NGHIEM"},
                            "clos": [],
                            "sources": [],
                            "content_hash": "hash-v1",
                        }
                    ]
                )
                self.moodle_targets = InMemoryCollection(
                    [
                        {
                            "_id": target_id,
                            "site_key": "demo-moodle",
                            "site_name": "Demo Moodle",
                            "mode": "REST_API",
                            "default_course_id": "ctdl",
                            "default_category_id": "qbank",
                            "is_active": True,
                        }
                    ]
                )
                self.moodle_publications = InMemoryCollection()

        db = FakePublicationDatabase()
        original_transaction = question_workflow_module.mongo_transaction
        original_app_env = settings.app_env
        original_demo_mode = settings.demo_mode
        try:
            question_workflow_module.mongo_transaction = lambda: nullcontext(None)
            settings.app_env = "demo"
            settings.demo_mode = True
            service = QuestionWorkflowService(db)
            payload = MoodlePublicationRequest(expected_version=1)

            first = service.publish_to_moodle(str(question_id), payload, publisher_id, "Reviewer")
            second = service.publish_to_moodle(str(question_id), payload, publisher_id, "Reviewer")
        finally:
            question_workflow_module.mongo_transaction = original_transaction
            settings.app_env = original_app_env
            settings.demo_mode = original_demo_mode

        self.assertEqual(len(db.moodle_publications.records), 1)
        self.assertEqual(first["_id"], second["_id"])
        self.assertEqual(first["idempotency_key"], second["idempotency_key"])
        self.assertEqual(db.questions.find_one({"_id": question_id})["publication_status"], "PUBLISHED")
        publication = db.moodle_publications.records[0]
        self.assertEqual(publication["publication_mode"], "MOCK")
        self.assertFalse(publication["external_sync"])
        self.assertEqual(publication["status_detail"], "SIMULATED_LOCAL_RECORD")
        self.assertEqual(publication["target"]["configured_mode"], "REST_API")
        self.assertEqual(publication["target"]["course_id"], "ctdl")
        self.assertEqual(publication["target"]["category_id"], "qbank")
        self.assertEqual(publication["target"]["allowed_roles"], ["Admin", "Reviewer"])
        self.assertIn("gift", publication["request_payload"]["exports"])
        self.assertIn("xml", publication["request_payload"]["exports"])
        self.assertIn("chưa gửi dữ liệu sang Moodle thật", publication["response_payload"]["message"])

    def test_moodle_publish_enforces_target_allowed_roles(self):
        question_id = ObjectId()
        version_id = ObjectId()
        target_id = ObjectId()
        reviewer_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakePublicationDatabase:
            def __init__(self):
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "schema_version": SCHEMA_VERSION,
                            "question_code": "Q-001",
                            "current_version": 1,
                            "current_version_id": version_id,
                            "approved_version_id": version_id,
                            "lifecycle_status": "ACTIVE",
                            "evaluation_status": "PASSED",
                            "review_status": "APPROVED",
                            "publication_status": "NOT_PUBLISHED",
                            "quality_summary": {},
                            "review_assignment": {"status": "UNASSIGNED"},
                            "latest_review_id": None,
                            "created_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.question_versions = InMemoryCollection(
                    [
                        {
                            "_id": version_id,
                            "question_id": question_id,
                            "version": 1,
                            "document_id": None,
                            "content": "Queue xử lý phần tử theo nguyên tắc nào?",
                            "question_data": {
                                "options": {"A": "FIFO", "B": "LIFO"},
                                "correct_answer": "A",
                                "explanation": "Queue là hàng đợi FIFO.",
                            },
                            "classification": {"assessment_type": "TRAC_NGHIEM"},
                            "clos": [],
                            "sources": [],
                            "content_hash": "hash-v1",
                        }
                    ]
                )
                self.moodle_targets = InMemoryCollection(
                    [
                        {
                            "_id": target_id,
                            "site_key": "admin-only",
                            "site_name": "Admin Moodle",
                            "mode": "MOCK",
                            "default_course_id": "ctdl",
                            "default_category_id": "qbank",
                            "allowed_roles": ["Admin"],
                            "is_active": True,
                        }
                    ]
                )
                self.moodle_publications = InMemoryCollection()

        service = QuestionWorkflowService(FakePublicationDatabase())
        payload = MoodlePublicationRequest(expected_version=1, target_id=str(target_id))
        original_app_env = settings.app_env
        original_demo_mode = settings.demo_mode
        try:
            settings.app_env = "demo"
            settings.demo_mode = True

            with self.assertRaises(PermissionError):
                service.publish_to_moodle(str(question_id), payload, reviewer_id, "Reviewer")
        finally:
            settings.app_env = original_app_env
            settings.demo_mode = original_demo_mode

    def test_admin_moodle_retry_failed_publication_updates_existing_attempt(self):
        db, failed_id, question_id = self._moodle_retry_database()
        admin = _current_user("Admin")
        original_transaction = question_workflow_module.mongo_transaction
        original_app_env = settings.app_env
        original_demo_mode = settings.demo_mode
        original_audit = MoodleTargetService._audit
        try:
            question_workflow_module.mongo_transaction = lambda: nullcontext(None)
            settings.app_env = "demo"
            settings.demo_mode = True
            MoodleTargetService._audit = staticmethod(lambda *_args, **_kwargs: None)

            result = MoodleTargetService(db).retry_publication(str(failed_id), admin)
        finally:
            question_workflow_module.mongo_transaction = original_transaction
            settings.app_env = original_app_env
            settings.demo_mode = original_demo_mode
            MoodleTargetService._audit = staticmethod(original_audit)

        self.assertEqual(result["id"], str(failed_id))
        self.assertEqual(result["status"], "PUBLISHED")
        self.assertEqual(result["attempt_no"], 2)
        self.assertEqual(len(db.moodle_publications.records), 1)
        self.assertEqual(db.moodle_publications.records[0]["status"], "PUBLISHED")
        self.assertIsNone(db.moodle_publications.records[0]["error"])
        self.assertEqual(
            db.questions.find_one({"_id": question_id})["publication_status"],
            "PUBLISHED",
        )

    def test_admin_moodle_retry_rejects_changed_question_version(self):
        db, failed_id, question_id = self._moodle_retry_database(current_version=2)
        admin = _current_user("Admin")
        original_transaction = question_workflow_module.mongo_transaction
        original_app_env = settings.app_env
        original_demo_mode = settings.demo_mode
        original_audit = MoodleTargetService._audit
        try:
            question_workflow_module.mongo_transaction = lambda: nullcontext(None)
            settings.app_env = "demo"
            settings.demo_mode = True
            MoodleTargetService._audit = staticmethod(lambda *_args, **_kwargs: None)

            with self.assertRaises(RuntimeError) as ctx:
                MoodleTargetService(db).retry_publication(str(failed_id), admin)
        finally:
            question_workflow_module.mongo_transaction = original_transaction
            settings.app_env = original_app_env
            settings.demo_mode = original_demo_mode
            MoodleTargetService._audit = staticmethod(original_audit)

        self.assertEqual(str(ctx.exception), "VERSION_CONFLICT")
        self.assertEqual(db.moodle_publications.records[0]["status"], "FAILED")
        self.assertEqual(
            db.questions.find_one({"_id": question_id})["publication_status"],
            "FAILED",
        )

    def test_moodle_target_payload_requires_real_api_config(self):
        with self.assertRaises(ValidationError):
            MoodleTargetPayload(
                site_key="ctu",
                site_name="CTU Moodle",
                mode="REST_API",
                default_course_id="ctdl",
                default_category_id="qbank",
            )

        payload = MoodleTargetPayload(
            site_key="ctu",
            site_name="CTU Moodle",
            mode="REST_API",
            base_url="https://moodle.example.edu/",
            token_env_var="MOODLE_TOKEN",
            default_course_id="ctdl",
            default_category_id="qbank",
        )
        self.assertEqual(payload.base_url, "https://moodle.example.edu")

        with self.assertRaises(ValidationError):
            MoodleTargetPayload(
                site_key="ctu",
                site_name="CTU Moodle",
                default_course_id="ctdl",
                default_category_id="qbank",
                allowed_roles=[],
            )

    def test_demo_login_route_registration_follows_demo_mode(self):
        route_paths = {route.path for route in auth_login.router.routes}
        self.assertEqual("/demo-login" in route_paths, settings.demo_mode)

    def test_review_dashboard_route_precedes_question_id_route(self):
        question_route_paths = [
            route.path
            for route in app.routes
            if getattr(route, "path", "").startswith(f"{settings.api_prefix}/questions")
            and "GET" in getattr(route, "methods", set())
        ]

        self.assertLess(
            question_route_paths.index(f"{settings.api_prefix}/questions/review-dashboard"),
            question_route_paths.index(f"{settings.api_prefix}/questions/{{question_id}}"),
        )

    def test_demo_login_does_not_reenable_disabled_firebase_user(self):
        class FakeFirebaseUser:
            uid = "demo-admin-uid"
            disabled = True

        class FakeAuth:
            class UserNotFoundError(Exception):
                pass

            update_calls = []

            @staticmethod
            def get_user_by_email(_email):
                return FakeFirebaseUser()

            @classmethod
            def update_user(cls, *args, **kwargs):
                cls.update_calls.append((args, kwargs))

        original_auth = auth_login.auth
        try:
            auth_login.auth = FakeAuth
            with self.assertRaises(HTTPException) as ctx:
                auth_login._ensure_demo_firebase_user("admin@example.com", "Admin Demo")
        finally:
            auth_login.auth = original_auth

        self.assertEqual(ctx.exception.status_code, 403)
        self.assertEqual(FakeAuth.update_calls, [])

    def test_demo_login_does_not_reactivate_disabled_app_user(self):
        user_id = ObjectId()

        class FakeDemoDatabase:
            def __init__(self):
                self.users = InMemoryCollection(
                    [
                        {
                            "_id": user_id,
                            "schema_version": SCHEMA_VERSION,
                            "firebase_uid": "demo-admin-uid",
                            "email": "old-admin@example.com",
                            "display_name": "Old Admin",
                            "role": "Admin",
                            "profile": {},
                            "is_active": False,
                            "created_at": datetime.now(timezone.utc),
                            "updated_at": datetime.now(timezone.utc),
                        }
                    ]
                )

        fake_db = FakeDemoDatabase()
        original_get_rag_db = auth_login.get_rag_db
        try:
            auth_login.get_rag_db = lambda: fake_db
            user = auth_login._ensure_demo_app_user(
                "demo-admin-uid",
                {
                    "email": "admin@qbankctu.edu.vn",
                    "display_name": "Admin Demo",
                    "role": "Admin",
                },
            )
        finally:
            auth_login.get_rag_db = original_get_rag_db

        self.assertFalse(user["is_active"])
        self.assertEqual(user["email"], "admin@qbankctu.edu.vn")
        self.assertEqual(user["display_name"], "Admin Demo")

    def test_teacher_cannot_access_another_teachers_exam(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        service = ExamService(FakeExamRepository([exam]), question_repository=None)

        with self.assertRaises(PermissionError):
            service.get_exam(str(exam["_id"]), other_teacher)

        result = service.get_exam(str(exam["_id"]), owner)
        self.assertEqual(result["id"], str(exam["_id"]))

    def test_document_job_history_enforces_document_ownership(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        document_id = ObjectId()
        job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeDocumentJobRepository:
            def find_by_id(self, _document_id):
                return {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "schema_version": SCHEMA_VERSION,
                    "archived_at": None,
                }

            def list_jobs(self, _document_id, *, limit=20):
                return [
                    {
                        "_id": job_id,
                        "document_id": document_id,
                        "document_version": 1,
                        "job_type": "OCR",
                        "attempt_no": 2,
                        "status": "FAILED",
                        "progress": 40,
                        "stats": {"pages": 3},
                        "error": {"message": "OCR timeout", "at": now},
                        "queued_at": now,
                        "started_at": now,
                        "finished_at": now,
                    }
                ][:limit]

        service = DocumentService(FakeDocumentJobRepository())

        with self.assertRaises(PermissionError):
            service.list_jobs(str(document_id), other_teacher)

        result = service.list_jobs(str(document_id), owner)
        self.assertEqual(result["items"][0]["id"], str(job_id))
        self.assertEqual(result["items"][0]["document_id"], str(document_id))
        self.assertEqual(result["items"][0]["error"]["message"], "OCR timeout")
        self.assertTrue(result["items"][0]["can_retry"])
        self.assertFalse(result["items"][0]["can_cancel"])

    def test_document_pages_enforce_document_ownership(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        document_id = ObjectId()
        page_id = ObjectId()
        ocr_job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeDocumentPageRepository:
            def find_by_id(self, _document_id):
                return {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "current_version": 2,
                    "schema_version": SCHEMA_VERSION,
                    "archived_at": None,
                }

            def list_pages(self, _document_id, *, document_version=None, limit=100):
                self.requested_version = document_version
                return [
                    {
                        "_id": page_id,
                        "document_id": document_id,
                        "document_version": 2,
                        "ocr_job_id": ocr_job_id,
                        "page_number": 1,
                        "raw_text": "raw OCR",
                        "cleaned_text": "clean OCR",
                        "formula_blocks": [{"latex": "x^2"}],
                        "created_at": now,
                    }
                ][:limit]

        repository = FakeDocumentPageRepository()
        service = DocumentService(repository)

        with self.assertRaises(PermissionError):
            service.list_pages(str(document_id), other_teacher)

        result = service.list_pages(str(document_id), owner)
        self.assertEqual(repository.requested_version, 2)
        self.assertEqual(result["items"][0]["id"], str(page_id))
        self.assertEqual(result["items"][0]["document_id"], str(document_id))
        self.assertEqual(result["items"][0]["cleaned_text"], "clean OCR")

    def test_document_page_update_requires_owner_and_pre_chunk_state(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        document_id = ObjectId()
        page_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeDocumentPageUpdateRepository:
            def __init__(self):
                self.document = {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "current_version": 1,
                    "pipeline_summary": {"chunk_status": "NOT_STARTED", "index_status": "NOT_STARTED"},
                }
                self.page = {
                    "_id": page_id,
                    "document_id": document_id,
                    "document_version": 1,
                    "ocr_job_id": ObjectId(),
                    "page_number": 1,
                    "raw_text": "raw",
                    "cleaned_text": "old",
                    "formula_blocks": [],
                    "created_at": now,
                }

            def find_by_id(self, _document_id):
                return self.document

            def update_page(self, _document_id, _page_id, *, document_version, cleaned_text):
                self.requested_version = document_version
                self.page["cleaned_text"] = cleaned_text
                return self.page

        repository = FakeDocumentPageUpdateRepository()
        service = DocumentService(repository)
        payload = DocumentPageUpdateRequest(cleaned_text="new OCR")

        with self.assertRaises(PermissionError):
            service.update_page(str(document_id), str(page_id), payload, other_teacher)

        result = service.update_page(str(document_id), str(page_id), payload, owner)
        self.assertEqual(repository.requested_version, 1)
        self.assertEqual(result["cleaned_text"], "new OCR")

        repository.document["pipeline_summary"]["chunk_status"] = "COMPLETED"
        with self.assertRaises(ValueError):
            service.update_page(str(document_id), str(page_id), payload, owner)

    def test_document_job_retry_and_cancel_enforce_document_ownership(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        document_id = ObjectId()
        failed_job_id = ObjectId()
        active_job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeBackgroundTasks:
            def __init__(self):
                self.tasks = []

            def add_task(self, func, *args, **kwargs):
                self.tasks.append({"func": func, "args": args, "kwargs": kwargs})

        class FakeDocumentActionRepository:
            def __init__(self):
                self.document = {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "title": "Owner PDF",
                    "original_filename": "owner.pdf",
                    "current_version": 1,
                    "artifacts": [
                        {
                            "type": "ORIGINAL_PDF",
                            "is_current": True,
                            "storage": {"provider": "LOCAL", "uri": "backend/data/uploads/owner.pdf"},
                        }
                    ],
                }
                self.jobs = {
                    str(failed_job_id): {
                        "_id": failed_job_id,
                        "document_id": document_id,
                        "document_version": 1,
                        "job_type": "OCR",
                        "attempt_no": 1,
                        "status": "FAILED",
                        "progress": 20,
                        "stats": None,
                        "error": {"message": "timeout", "at": now},
                        "config": {"lang": "vi"},
                        "queued_at": now,
                        "started_at": now,
                        "finished_at": now,
                    },
                    str(active_job_id): {
                        "_id": active_job_id,
                        "document_id": document_id,
                        "document_version": 1,
                        "job_type": "OCR",
                        "attempt_no": 1,
                        "status": "PROCESSING",
                        "progress": 30,
                        "stats": None,
                        "error": None,
                        "queued_at": now,
                        "started_at": now,
                        "finished_at": None,
                    },
                }

            def find_by_id(self, _document_id):
                return self.document

            def list_jobs(self, _document_id, *, limit=20):
                return list(self.jobs.values())[:limit]

            def find_job(self, job_id):
                return self.jobs.get(str(job_id))

            def create_job(self, document_id_arg, job_type, config=None):
                new_id = ObjectId()
                job = {
                    "_id": new_id,
                    "document_id": document_id_arg,
                    "document_version": 1,
                    "job_type": job_type,
                    "attempt_no": 2,
                    "status": "QUEUED",
                    "progress": 0,
                    "stats": None,
                    "error": None,
                    "config": config or {},
                    "queued_at": now,
                    "started_at": None,
                    "finished_at": None,
                }
                self.jobs[str(new_id)] = job
                return job

            def update_job(self, job_id, status, *, progress=None, stats=None, error_message=None):
                job = self.jobs[str(job_id)]
                job["status"] = status
                job["progress"] = progress if progress is not None else job.get("progress")
                job["stats"] = stats if stats is not None else job.get("stats")
                job["error"] = {"message": error_message, "at": now} if error_message else job.get("error")
                job["finished_at"] = now
                return job

        background_tasks = FakeBackgroundTasks()
        service = DocumentService(FakeDocumentActionRepository())

        with self.assertRaises(PermissionError):
            service.retry_job(str(document_id), str(failed_job_id), background_tasks, other_teacher)

        retried = service.retry_job(str(document_id), str(failed_job_id), background_tasks, owner)
        self.assertEqual(retried["job"]["attempt_no"], 2)
        self.assertEqual(retried["job"]["status"], "QUEUED")
        self.assertEqual(background_tasks.tasks[0]["kwargs"]["document_id"], str(document_id))

        with self.assertRaises(PermissionError):
            service.cancel_job(str(document_id), str(active_job_id), other_teacher)

        cancelled = service.cancel_job(str(document_id), str(active_job_id), owner)
        self.assertEqual(cancelled["job"]["status"], "CANCELLED")
        self.assertIn("Cancelled by Teacher", cancelled["job"]["error"]["message"])

    def test_document_chunk_job_retry_uses_failed_job_config(self):
        owner = _current_user("Teacher")
        document_id = ObjectId()
        failed_job_id = ObjectId()
        new_job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeBackgroundTasks:
            def __init__(self):
                self.tasks = []

            def add_task(self, func, *args, **kwargs):
                self.tasks.append({"func": func, "args": args, "kwargs": kwargs})

        class FakeChunkRetryRepository:
            def __init__(self):
                self.jobs = {
                    str(failed_job_id): {
                        "_id": failed_job_id,
                        "document_id": document_id,
                        "document_version": 1,
                        "job_type": "CHUNK",
                        "attempt_no": 1,
                        "status": "FAILED",
                        "progress": 80,
                        "stats": None,
                        "error": {"message": "chunk failed", "at": now},
                        "config": {"chunk_size": 700, "chunk_overlap": 120, "dry_run": True},
                        "queued_at": now,
                        "started_at": now,
                        "finished_at": now,
                    }
                }

            def find_by_id(self, _document_id):
                return {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "current_version": 1,
                }

            def find_job(self, job_id):
                return self.jobs.get(str(job_id))

        repository = FakeChunkRetryRepository()
        captured = {}

        def fake_queue_chunk_retry(background_tasks, document_id, config):
            captured["document_id"] = document_id
            captured["config"] = config
            repository.jobs[str(new_job_id)] = {
                "_id": new_job_id,
                "document_id": ObjectId(document_id),
                "document_version": 1,
                "job_type": "CHUNK",
                "attempt_no": 2,
                "status": "QUEUED",
                "progress": 0,
                "stats": None,
                "error": None,
                "queued_at": now,
                "started_at": None,
                "finished_at": None,
            }
            background_tasks.add_task(lambda: None)
            return {"chunk_job_id": str(new_job_id), "chunk_set_id": str(ObjectId())}

        import modules.rag.chunking as chunking_module

        original_queue = chunking_module.queue_chunk_retry
        try:
            chunking_module.queue_chunk_retry = fake_queue_chunk_retry
            result = DocumentService(repository).retry_job(
                str(document_id),
                str(failed_job_id),
                FakeBackgroundTasks(),
                owner,
            )
        finally:
            chunking_module.queue_chunk_retry = original_queue

        self.assertEqual(captured["document_id"], str(document_id))
        self.assertEqual(captured["config"]["chunk_size"], 700)
        self.assertEqual(result["job"]["id"], str(new_job_id))
        self.assertTrue(result["job"]["can_cancel"])

    def test_cancel_chunk_job_marks_chunk_set_and_pending_embeddings_cancelled(self):
        document_id = ObjectId()
        chunk_job_id = ObjectId()
        chunk_set_id = ObjectId()
        embedding_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeDocumentDatabase:
            def __init__(self):
                self.documents = InMemoryCollection(
                    [
                        {
                            "_id": document_id,
                            "archived_at": None,
                            "status": "PROCESSING",
                            "pipeline_summary": {"chunk_status": "PROCESSING"},
                            "updated_at": now,
                        }
                    ]
                )
                self.document_jobs = InMemoryCollection(
                    [
                        {
                            "_id": chunk_job_id,
                            "document_id": document_id,
                            "job_type": "CHUNK",
                            "status": "PROCESSING",
                            "queued_at": now,
                            "started_at": now,
                        }
                    ]
                )
                self.chunk_sets = InMemoryCollection(
                    [{"_id": chunk_set_id, "chunk_job_id": chunk_job_id, "status": "PROCESSING"}]
                )
                self.chunk_embeddings = InMemoryCollection(
                    [{"_id": embedding_id, "chunk_set_id": chunk_set_id, "status": "PENDING"}]
                )

        db = FakeDocumentDatabase()
        updated = MongoDocumentRepository(db).update_job(
            str(chunk_job_id),
            "CANCELLED",
            error_message="Cancelled by test",
        )

        self.assertEqual(updated["status"], "CANCELLED")
        self.assertEqual(db.chunk_sets.find_one({"_id": chunk_set_id})["status"], "CANCELLED")
        self.assertEqual(db.chunk_embeddings.find_one({"_id": embedding_id})["status"], "CANCELLED")
        document = db.documents.find_one({"_id": document_id})
        self.assertEqual(document["pipeline_summary"]["chunk_status"], "CANCELLED")
        self.assertEqual(document["status"], "FAILED")

    def test_document_reindex_requires_owner_and_completed_chunks(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        document_id = ObjectId()
        chunk_set_id = ObjectId()
        new_job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeBackgroundTasks:
            def __init__(self):
                self.tasks = []

            def add_task(self, func, *args, **kwargs):
                self.tasks.append({"func": func, "args": args, "kwargs": kwargs})

        class FakeReindexRepository:
            def __init__(self):
                self.document = {
                    "_id": document_id,
                    "uploaded_by_user_id": owner.id,
                    "current_version": 1,
                    "current_processing": {"chunk_set_id": chunk_set_id},
                    "pipeline_summary": {"chunk_status": "COMPLETED"},
                }
                self.jobs = {}

            def find_by_id(self, _document_id):
                return self.document

            def find_job(self, job_id):
                return self.jobs.get(str(job_id))

        repository = FakeReindexRepository()

        def fake_queue_document_reindex(background_tasks, document_id, collection_name=None):
            repository.jobs[str(new_job_id)] = {
                "_id": new_job_id,
                "document_id": ObjectId(document_id),
                "document_version": 1,
                "job_type": "INDEX",
                "attempt_no": 1,
                "status": "QUEUED",
                "progress": 0,
                "stats": None,
                "error": None,
                "queued_at": now,
                "started_at": None,
                "finished_at": None,
            }
            background_tasks.add_task(lambda: None)
            return {"index_job_id": str(new_job_id)}

        import modules.rag.chunking as chunking_module

        original_queue = chunking_module.queue_document_reindex
        try:
            chunking_module.queue_document_reindex = fake_queue_document_reindex
            service = DocumentService(repository)
            with self.assertRaises(PermissionError):
                service.reindex(str(document_id), FakeBackgroundTasks(), other_teacher)

            result = service.reindex(str(document_id), FakeBackgroundTasks(), owner)
        finally:
            chunking_module.queue_document_reindex = original_queue

        self.assertEqual(result["job"]["id"], str(new_job_id))
        self.assertTrue(result["job"]["can_cancel"])

        repository.document["pipeline_summary"]["chunk_status"] = "FAILED"
        with self.assertRaises(ValueError):
            DocumentService(repository).reindex(str(document_id), FakeBackgroundTasks(), owner)

    def test_docx_upload_extractor_builds_pages_from_paragraphs_and_tables(self):
        with TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "source.docx"
            document = Document()
            document.add_paragraph("Chương 1. Tổng quan cấu trúc dữ liệu")
            document.add_paragraph("Hàng đợi sử dụng chính sách FIFO.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Thuật ngữ"
            table.cell(0, 1).text = "Ý nghĩa"
            table.cell(1, 0).text = "Stack"
            table.cell(1, 1).text = "LIFO"
            document.save(docx_path)

            pages, stats = extract_docx_pages(docx_path, page_char_limit=80)

        joined_text = "\n".join(page["text"] for page in pages)
        self.assertGreaterEqual(len(pages), 1)
        self.assertEqual(stats["source_format"], "docx")
        self.assertEqual(stats["paragraph_count"], 2)
        self.assertEqual(stats["table_count"], 1)
        self.assertIn("FIFO", joined_text)
        self.assertIn("Stack | LIFO", joined_text)

    def test_admin_can_access_any_teacher_exam(self):
        owner = _current_user("Teacher")
        admin = _current_user("Admin")
        exam = _exam_doc(owner.id)
        service = ExamService(FakeExamRepository([exam]), question_repository=None)

        result = service.get_exam(str(exam["_id"]), admin)
        self.assertEqual(result["created_by_user_id"], str(owner.id))

    def test_exam_docx_export_contains_questions_and_answers(self):
        docx_bytes = render_exam_docx(
            {
                "school_name": "CTU",
                "faculty_name": "College of ICT",
                "exam_name": "Midterm",
                "subject_name": "Data Structures",
                "duration_minutes": 90,
            },
            "A01",
            [
                {
                    "order": 1,
                    "content_snapshot": {
                        "content": "Queue uses which policy?",
                        "question_data": {
                            "options": {"A": "FIFO", "B": "LIFO"},
                            "correct_answer": "A",
                        },
                    },
                }
            ],
            "de_dapan",
        )

        document = Document(BytesIO(docx_bytes))
        paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertGreater(len(docx_bytes), 1000)
        self.assertIn("Midterm", paragraph_text)
        self.assertIn("Queue uses which policy?", paragraph_text)
        self.assertIn("FIFO", paragraph_text)
        self.assertIn("A", table_text)

    def test_exam_lifecycle_requires_exact_current_approved_questions(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        question_id = ObjectId()
        version_id = ObjectId()
        question = {
            "_id": question_id,
            "question_code": "Q-1",
            "current_version": 1,
            "current_version_id": version_id,
            "approved_version_id": version_id,
            "lifecycle_status": "ACTIVE",
            "evaluation_status": "PASSED",
            "review_status": "APPROVED",
            "publication_status": "NOT_PUBLISHED",
            "quality_summary": {},
            "review_assignment": {"status": "UNASSIGNED"},
            "latest_review_id": None,
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc),
        }
        version = {
            "_id": version_id,
            "question_id": question_id,
            "document_id": None,
            "content": "Queue uses FIFO",
            "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
            "classification": {"subject": {"id": exam["subject_id"]}, "assessment_type": "TRAC_NGHIEM"},
            "clos": [],
            "sources": [],
            "content_hash": "hash",
        }
        service = ExamService(
            FakeExamRepository([exam]),
            FakeExamQuestionRepository([(question, version)]),
        )

        with self.assertRaises(ValueError):
            service.update_status(str(exam["_id"]), ExamStatusUpdateRequest(status="READY"), owner)

        exam["questions"] = [
            {
                "question_id": question_id,
                "version_id": version_id,
                "content_snapshot": {"content": "Queue uses FIFO"},
            }
        ]
        ready = service.update_status(
            str(exam["_id"]),
            ExamStatusUpdateRequest(status="READY"),
            owner,
        )
        result = service.update_status(
            str(exam["_id"]),
            ExamStatusUpdateRequest(status="FINALIZED"),
            owner,
        )

        self.assertEqual(ready["status"], "READY")
        self.assertEqual(result["status"], "FINALIZED")
        self.assertIn("finalized_snapshot", exam)

    def test_exam_lifecycle_requires_approved_version_and_subject(self):
        owner = _current_user("Teacher")

        def make_pair(*, approved_version_id=True, subject_id=None):
            question_id = ObjectId()
            version_id = ObjectId()
            question = {
                "_id": question_id,
                "question_code": "Q-1",
                "current_version": 1,
                "current_version_id": version_id,
                "approved_version_id": version_id if approved_version_id is True else approved_version_id,
                "lifecycle_status": "ACTIVE",
                "evaluation_status": "PASSED",
                "review_status": "APPROVED",
                "publication_status": "NOT_PUBLISHED",
                "quality_summary": {},
                "review_assignment": {"status": "UNASSIGNED"},
                "latest_review_id": None,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
            subject = {} if subject_id is None else {"id": subject_id}
            version = {
                "_id": version_id,
                "question_id": question_id,
                "document_id": None,
                "content": "Queue uses FIFO",
                "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
                "classification": {"subject": subject, "assessment_type": "TRAC_NGHIEM"},
                "clos": [],
                "sources": [],
                "content_hash": "hash",
            }
            return question, version

        for question, version in (
            make_pair(approved_version_id=None, subject_id=ObjectId()),
            make_pair(approved_version_id=True, subject_id=None),
        ):
            exam = _exam_doc(owner.id)
            if (version.get("classification") or {}).get("subject"):
                version["classification"]["subject"]["id"] = exam["subject_id"]
            exam["questions"] = [
                {
                    "question_id": question["_id"],
                    "version_id": version["_id"],
                    "content_snapshot": {"content": version["content"]},
                }
            ]
            service = ExamService(
                FakeExamRepository([exam]),
                FakeExamQuestionRepository([(question, version)]),
            )

            with self.assertRaises(ValueError):
                service.update_status(str(exam["_id"]), ExamStatusUpdateRequest(status="READY"), owner)

    def test_matrix_availability_counts_only_usable_approved_questions(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        exam["matrix"] = [
            {
                "chapter_id": None,
                "cognitive_level": "nhan_biet",
                "difficulty": "de",
                "count": 2,
            }
        ]

        def make_pair(approved_version_id):
            question_id = ObjectId()
            version_id = ObjectId()
            question = {
                "_id": question_id,
                "question_code": f"Q-{question_id}",
                "current_version": 1,
                "current_version_id": version_id,
                "approved_version_id": version_id if approved_version_id else ObjectId(),
                "lifecycle_status": "ACTIVE",
                "evaluation_status": "PASSED",
                "review_status": "APPROVED",
                "publication_status": "NOT_PUBLISHED",
                "quality_summary": {},
                "review_assignment": {"status": "UNASSIGNED"},
                "latest_review_id": None,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
            version = {
                "_id": version_id,
                "question_id": question_id,
                "document_id": None,
                "content": "Queue uses FIFO",
                "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
                "classification": {"subject": {"id": exam["subject_id"]}},
                "clos": [],
                "sources": [],
                "content_hash": "hash",
            }
            return question, version

        service = ExamService(
            FakeExamRepository([exam]),
            FakeExamQuestionRepository([
                make_pair(True),
                make_pair(False),
            ]),
        )

        availability = service.matrix_availability(str(exam["_id"]), owner)

        self.assertEqual(availability[0]["available"], 1)
        self.assertFalse(availability[0]["sufficient"])

    def test_finalized_exam_cannot_be_hard_deleted_or_edited(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        question_id = ObjectId()
        version_id = ObjectId()
        snapshot = {
            "id": str(question_id),
            "question_code": "Q-1",
            "content": "Queue uses FIFO",
            "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
            "classification": {"subject": {"id": str(exam["subject_id"])}},
        }
        exam.update(
            {
                "status": "FINALIZED",
                "question_count": 1,
                "matrix": [],
                "questions": [
                    {
                        "question_id": question_id,
                        "version_id": version_id,
                        "content_snapshot": snapshot,
                    }
                ],
                "finalized_snapshot": {
                    "subject_id": exam["subject_id"],
                    "question_count": 1,
                    "matrix": [],
                    "questions": [
                        {
                            "question_id": question_id,
                            "version_id": version_id,
                            "content_snapshot": snapshot,
                        }
                    ],
                },
            }
        )
        repository = FakeExamRepository([exam])
        service = ExamService(repository, FakeExamQuestionRepository())

        with self.assertRaises(ValueError):
            service.update_exam(str(exam["_id"]), ExamUpdateRequest(name="Edited"), owner)
        with self.assertRaises(ValueError):
            service.save_matrix(str(exam["_id"]), ExamMatrixRequest(cells=[]), owner)
        with self.assertRaises(ValueError):
            service.auto_generate_pool(str(exam["_id"]), owner)
        with self.assertRaises(ValueError):
            service.add_questions_manual(
                str(exam["_id"]),
                AddQuestionsManualRequest(question_ids=[str(ObjectId())]),
                owner,
            )
        with self.assertRaises(ValueError):
            service.remove_question(str(exam["_id"]), str(question_id), owner)
        with self.assertRaises(ValueError):
            service.delete_exam(str(exam["_id"]), owner)

        self.assertIsNotNone(repository.find(exam["_id"]))

    def test_duplicate_exam_creates_editable_draft_without_variants(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        question_id = ObjectId()
        version_id = ObjectId()
        snapshot = {
            "id": str(question_id),
            "question_code": "Q-1",
            "content": "Queue uses FIFO",
            "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
            "classification": {"subject": {"id": str(exam["subject_id"])}},
        }
        exam.update(
            {
                "status": "FINALIZED",
                "question_count": 1,
                "matrix": [
                    {
                        "chapter_id": None,
                        "cognitive_level": "nhan_biet",
                        "difficulty": "de",
                        "count": 1,
                    }
                ],
                "questions": [
                    {
                        "question_id": question_id,
                        "version_id": version_id,
                        "content_snapshot": snapshot,
                    }
                ],
                "finalized_snapshot": {
                    "subject_id": exam["subject_id"],
                    "question_count": 1,
                    "header": {"school": "CTU", "duration_minutes": 90},
                    "matrix": [
                        {
                            "chapter_id": None,
                            "cognitive_level": "nhan_biet",
                            "difficulty": "de",
                            "count": 1,
                        }
                    ],
                    "questions": [
                        {
                            "question_id": question_id,
                            "version_id": version_id,
                            "content_snapshot": snapshot,
                        }
                    ],
                },
            }
        )
        repository = FakeExamRepository(
            [exam],
            variants=[{"_id": ObjectId(), "exam_id": exam["_id"]}],
        )
        service = ExamService(repository, FakeExamQuestionRepository())

        duplicated = service.duplicate_exam(str(exam["_id"]), owner)
        stored_duplicate = repository.find(duplicated["id"])

        self.assertNotEqual(duplicated["id"], str(exam["_id"]))
        self.assertEqual(duplicated["status"], "DRAFT")
        self.assertEqual(duplicated["variant_count"], 0)
        self.assertEqual(duplicated["questions"][0]["question_id"], str(question_id))
        self.assertEqual(duplicated["matrix"][0]["count"], 1)
        self.assertEqual(stored_duplicate["created_by_user_id"], owner.id)
        self.assertNotIn("finalized_snapshot", stored_duplicate)
        self.assertEqual(exam["status"], "FINALIZED")

    def test_exam_question_pool_uses_server_side_filters(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        chapter_id = ObjectId()

        class RecordingQuestionRepository:
            def __init__(self):
                self.calls = []

            def list(self, page, page_size, review_status, search, **filters):
                self.calls.append(
                    {
                        "page": page,
                        "page_size": page_size,
                        "review_status": review_status,
                        "search": search,
                        "filters": filters,
                    }
                )
                return [], 0

        question_repository = RecordingQuestionRepository()
        service = ExamService(FakeExamRepository([exam]), question_repository)

        result = service.question_pool(
            str(exam["_id"]),
            owner,
            page=3,
            page_size=15,
            search="queue",
            question_type="trac_nghiem",
            bloom_level=2,
            chapter_id=str(chapter_id),
            difficulty="de",
        )

        self.assertEqual(result["page"], 3)
        self.assertEqual(result["page_size"], 15)
        self.assertEqual(result["total"], 0)
        call = question_repository.calls[0]
        self.assertEqual(call["review_status"], "APPROVED")
        self.assertEqual(call["search"], "queue")
        self.assertEqual(call["filters"]["question_type"], "trac_nghiem")
        self.assertEqual(call["filters"]["bloom_level"], 2)
        self.assertEqual(call["filters"]["subject_id"], str(exam["subject_id"]))
        self.assertEqual(call["filters"]["chapter_id"], str(chapter_id))
        self.assertEqual(call["filters"]["difficulty"], "de")
        self.assertEqual(call["filters"]["owner_user_id"], owner.id)
        self.assertTrue(call["filters"]["approved_current_only"])

    def test_exam_question_pool_hides_stale_approved_versions(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)

        def make_pair(current_is_approved):
            question_id = ObjectId()
            version_id = ObjectId()
            question = {
                "_id": question_id,
                "question_code": f"Q-{question_id}",
                "current_version": 1,
                "current_version_id": version_id,
                "approved_version_id": version_id if current_is_approved else ObjectId(),
                "lifecycle_status": "ACTIVE",
                "evaluation_status": "PASSED",
                "review_status": "APPROVED",
                "publication_status": "NOT_PUBLISHED",
                "quality_summary": {},
                "review_assignment": {"status": "UNASSIGNED"},
                "latest_review_id": None,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc),
            }
            version = {
                "_id": version_id,
                "question_id": question_id,
                "document_id": None,
                "content": "Queue uses FIFO",
                "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
                "classification": {"subject": {"id": exam["subject_id"]}},
                "clos": [],
                "sources": [],
                "content_hash": "hash",
            }
            return question, version

        usable, _usable_version = make_pair(True)
        stale, _stale_version = make_pair(False)
        service = ExamService(
            FakeExamRepository([exam]),
            FakeExamQuestionRepository([
                (usable, _usable_version),
                (stale, _stale_version),
            ]),
        )

        result = service.question_pool(str(exam["_id"]), owner)

        self.assertEqual(result["total"], 1)
        self.assertEqual([item["id"] for item in result["items"]], [str(usable["_id"])])

    def test_variant_creation_requires_finalized_exam(self):
        owner = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        question_id = ObjectId()
        version_id = ObjectId()
        exam["questions"] = [
            {
                "question_id": question_id,
                "version_id": version_id,
                "content_snapshot": {
                    "content": "Queue uses FIFO",
                    "question_data": {"options": {"A": "FIFO"}, "correct_answer": "A"},
                },
            }
        ]
        exams = FakeExamRepository([exam])
        service = ExamVariantService(exams, FakeExamVariantRepository())

        with self.assertRaises(ValueError):
            service.create_variant(
                str(exam["_id"]),
                ExamVariantCreateRequest(exam_code="A01"),
                owner,
            )

        exam["status"] = "FINALIZED"
        exam["finalized_snapshot"] = {"questions": exam["questions"]}
        result = service.create_variant(
            str(exam["_id"]),
            ExamVariantCreateRequest(exam_code="A01"),
            owner,
        )
        self.assertEqual(result["exam_code"], "A01")

    def test_variant_preview_enforces_exam_ownership(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        variant = {
            "_id": ObjectId(),
            "exam_id": exam["_id"],
            "exam_code": "A01",
            "questions": [],
            "answer_key": {},
            "created_at": datetime.now(timezone.utc),
        }
        service = ExamVariantService(
            FakeExamRepository([exam]),
            FakeExamVariantRepository([variant]),
        )

        with self.assertRaises(PermissionError):
            service.build_preview(str(exam["_id"]), str(variant["_id"]), other_teacher)

    def test_exam_ownership_is_enforced_across_mutation_and_variant_methods(self):
        owner = _current_user("Teacher")
        other_teacher = _current_user("Teacher")
        exam = _exam_doc(owner.id)
        question_id = ObjectId()
        variant = {
            "_id": ObjectId(),
            "exam_id": exam["_id"],
            "exam_code": "A01",
            "questions": [],
            "answer_key": {},
            "created_at": datetime.now(timezone.utc),
        }
        exam_repository = FakeExamRepository([exam], variants=[variant])
        service = ExamService(exam_repository, FakeExamQuestionRepository())
        variant_service = ExamVariantService(exam_repository, FakeExamVariantRepository([variant]))

        exam_method_calls = [
            lambda: service.update_status(
                str(exam["_id"]),
                ExamStatusUpdateRequest(status="READY"),
                other_teacher,
            ),
            lambda: service.update_exam(
                str(exam["_id"]),
                ExamUpdateRequest(name="Edited"),
                other_teacher,
            ),
            lambda: service.delete_exam(str(exam["_id"]), other_teacher),
            lambda: service.save_matrix(
                str(exam["_id"]),
                ExamMatrixRequest(cells=[]),
                other_teacher,
            ),
            lambda: service.matrix_availability(str(exam["_id"]), other_teacher),
            lambda: service.question_pool(str(exam["_id"]), other_teacher),
            lambda: service.auto_generate_pool(str(exam["_id"]), other_teacher),
            lambda: service.add_questions_manual(
                str(exam["_id"]),
                AddQuestionsManualRequest(question_ids=[str(question_id)]),
                other_teacher,
            ),
            lambda: service.remove_question(str(exam["_id"]), str(question_id), other_teacher),
        ]
        variant_method_calls = [
            lambda: variant_service.create_variant(
                str(exam["_id"]),
                ExamVariantCreateRequest(exam_code="B01"),
                other_teacher,
            ),
            lambda: variant_service.list_variants(str(exam["_id"]), other_teacher),
            lambda: variant_service.get_variant(str(exam["_id"]), str(variant["_id"]), other_teacher),
            lambda: variant_service.delete_variant(str(exam["_id"]), str(variant["_id"]), other_teacher),
            lambda: variant_service.build_preview(str(exam["_id"]), str(variant["_id"]), other_teacher),
        ]

        for call in [*exam_method_calls, *variant_method_calls]:
            with self.assertRaises(PermissionError):
                call()

    def test_cannot_deactivate_last_active_admin(self):
        admin_doc = _user_doc("Admin", True)
        service = UserService(FakeUserRepository([admin_doc]), FakeIdentityGateway(), FakeSessions())
        actor = _current_user("Admin", admin_doc["_id"])

        with self.assertRaises(ValueError):
            service.deactivate(str(admin_doc["_id"]), actor)

    def test_cannot_downgrade_last_active_admin(self):
        admin_doc = _user_doc("Admin", True)
        identity = FakeIdentityGateway()
        service = UserService(FakeUserRepository([admin_doc]), identity, FakeSessions())
        actor = _current_user("Admin", admin_doc["_id"])

        with self.assertRaises(ValueError):
            service.update_admin(
                str(admin_doc["_id"]),
                UserAdminUpdateRequest(role="Teacher"),
                actor,
            )
        self.assertEqual(identity.disabled_calls, [])

    def test_user_service_lists_teacher_options_for_review_filters(self):
        teacher = _user_doc("Teacher", True)
        inactive_teacher = _user_doc("Teacher", False)
        reviewer = _user_doc("Reviewer", True)
        service = UserService(
            FakeUserRepository([teacher, inactive_teacher, reviewer]),
            FakeIdentityGateway(),
            FakeSessions(),
        )

        result = service.list_teacher_options()

        self.assertEqual(result["total"], 2)
        self.assertEqual(
            {item["id"] for item in result["items"]},
            {str(teacher["_id"]), str(inactive_teacher["_id"])},
        )
        self.assertNotIn(str(reviewer["_id"]), {item["id"] for item in result["items"]})

    def test_user_service_lists_active_reviewer_options_for_assignment(self):
        teacher = _user_doc("Teacher", True)
        reviewer = _user_doc("Reviewer", True)
        inactive_reviewer = _user_doc("Reviewer", False)
        admin = _user_doc("Admin", True)
        service = UserService(
            FakeUserRepository([teacher, reviewer, inactive_reviewer, admin]),
            FakeIdentityGateway(),
            FakeSessions(),
        )

        result = service.list_reviewer_options()

        self.assertEqual(result["total"], 2)
        self.assertEqual(
            {item["id"] for item in result["items"]},
            {str(reviewer["_id"]), str(admin["_id"])},
        )
        self.assertNotIn(str(teacher["_id"]), {item["id"] for item in result["items"]})
        self.assertNotIn(str(inactive_reviewer["_id"]), {item["id"] for item in result["items"]})

    def test_catalog_lifecycle_counts_usage_and_blocks_duplicate_codes(self):
        subject_id = ObjectId()
        other_subject_id = ObjectId()
        chapter_id = ObjectId()
        clo_id = ObjectId()
        version_id = ObjectId()
        subject = {
            "_id": subject_id,
            "subject_code": "CTDL",
            "subject_name": "Cấu trúc dữ liệu",
            "description": "",
            "is_active": True,
            "chapters": [
                {
                    "_id": chapter_id,
                    "chapter_code": "CH01",
                    "chapter_name": "Stack",
                    "sequence_no": 1,
                    "is_active": True,
                }
            ],
            "learning_outcomes": [
                {
                    "_id": clo_id,
                    "clo_code": "CLO1",
                    "description": "Hiểu cấu trúc dữ liệu tuyến tính",
                    "target_weight": 0.5,
                    "is_active": True,
                }
            ],
        }
        db = FakeCatalogDatabase(
            subjects=[
                subject,
                {
                    "_id": other_subject_id,
                    "subject_code": "MMT",
                    "subject_name": "Mạng máy tính",
                    "chapters": [],
                    "learning_outcomes": [],
                    "is_active": True,
                },
            ],
            documents=[{"subject_id": subject_id, "chapter_id": chapter_id, "archived_at": None}],
            question_versions=[
                {
                    "_id": version_id,
                    "classification": {
                        "subject": {"id": subject_id},
                        "chapter": {"id": chapter_id},
                    },
                    "clos": [{"id": clo_id}],
                }
            ],
            questions=[
                {
                    "schema_version": SCHEMA_VERSION,
                    "lifecycle_status": "ACTIVE",
                    "current_version_id": version_id,
                }
            ],
            exams=[{"subject_id": subject_id, "matrix": [{"chapter_id": chapter_id}]}],
        )
        service = CatalogService(db)

        result = service.list_subjects()[0]
        self.assertEqual(result["usage_counts"]["documents"], 1)
        self.assertEqual(result["usage_counts"]["questions"], 1)
        self.assertEqual(result["usage_counts"]["exams"], 1)
        self.assertEqual(result["chapters"][0]["usage_counts"]["documents"], 1)
        self.assertEqual(result["learning_outcomes"][0]["usage_counts"]["questions"], 1)

        with self.assertRaises(ValueError):
            service.add_chapter(
                str(subject_id),
                ChapterPayload(chapter_code="CH01", chapter_name="Duplicate"),
            )
        with self.assertRaises(ValueError):
            service.add_learning_outcome(
                str(subject_id),
                LearningOutcomePayload(clo_code="CLO1", description="Duplicate"),
            )
        with self.assertRaises(ValueError):
            service.update_subject(
                str(subject_id),
                SubjectUpdatePayload(subject_code="MMT"),
            )

        updated = service.update_chapter(
            str(subject_id),
            str(chapter_id),
            ChapterUpdatePayload(chapter_code="CH02", is_active=False),
        )
        self.assertEqual(updated["chapters"][0]["chapter_code"], "CH02")
        self.assertFalse(updated["chapters"][0]["is_active"])

        updated = service.update_learning_outcome(
            str(subject_id),
            str(clo_id),
            LearningOutcomeUpdatePayload(target_weight=0.75, is_active=False),
        )
        self.assertEqual(updated["learning_outcomes"][0]["target_weight"], 0.75)
        self.assertFalse(updated["learning_outcomes"][0]["is_active"])

    def test_catalog_runtime_controls_prompt_policy_and_model_state(self):
        db = FakeCatalogDatabase(
            ai_models=[
                {
                    "_id": ObjectId(),
                    "model_code": "qwen",
                    "model_name": "Qwen local",
                    "runtime": "OLLAMA",
                    "kind": "CHAT",
                    "revision": "local",
                    "capabilities": ["QUESTION_GENERATION"],
                    "priority": 1,
                    "is_local": True,
                    "is_active": True,
                    "config": {},
                },
                {
                    "_id": ObjectId(),
                    "model_code": "unknown-provider",
                    "model_name": "Unknown",
                    "runtime": "CUSTOM",
                    "kind": "CHAT",
                    "revision": "local",
                    "capabilities": [],
                    "priority": 2,
                    "is_local": True,
                    "is_active": True,
                    "config": {},
                },
            ],
            prompt_templates=[
                {
                    "_id": ObjectId(),
                    "template_key": "system",
                    "version": 1,
                    "kind": "SYSTEM",
                    "name": "System v1",
                    "prompt_body": "v1",
                    "is_active": False,
                },
                {
                    "_id": ObjectId(),
                    "template_key": "system",
                    "version": 2,
                    "kind": "SYSTEM",
                    "name": "System v2",
                    "prompt_body": "v2",
                    "is_active": True,
                },
            ],
            evaluation_policies=[
                {
                    "_id": ObjectId(),
                    "policy_name": "Default",
                    "version": 1,
                    "weights": {"faithfulness": 1},
                    "thresholds": {"pass_min": 0.7},
                    "is_active": False,
                },
                {
                    "_id": ObjectId(),
                    "policy_name": "Default",
                    "version": 2,
                    "weights": {"faithfulness": 1},
                    "thresholds": {"pass_min": 0.8},
                    "is_active": True,
                },
            ],
        )
        service = CatalogService(db)

        models = service.list_ai_models()
        self.assertTrue(models[0]["factory_status"]["supported"])
        self.assertFalse(models[1]["factory_status"]["supported"])

        model = service.set_ai_model_active(
            AiModelActivationPayload(model_code="qwen", is_active=False)
        )
        self.assertFalse(model["is_active"])

        prompt = service.activate_prompt_template(
            PromptTemplateActivationPayload(template_key="system", version=1)
        )
        self.assertTrue(prompt["is_active"])
        self.assertFalse(db.prompt_templates.find_one({"template_key": "system", "version": 2})["is_active"])

        policy = service.activate_evaluation_policy(
            EvaluationPolicyActivationPayload(policy_name="Default", version=1)
        )
        self.assertTrue(policy["is_active"])
        self.assertFalse(db.evaluation_policies.find_one({"policy_name": "Default", "version": 2})["is_active"])

        runtime = service.runtime_config()
        self.assertIn("prompt_source", runtime)
        self.assertIn("supported_provider_patterns", runtime)
        if settings.model_provider == "qwen":
            self.assertTrue(
                any("Model sinh c" in warning and "inactive" in warning for warning in runtime["warnings"])
            )

    def test_admin_job_summary_tracks_operational_states(self):
        jobs = [
            {"status": "queued", "is_long_running": False},
            {"status": "PROCESSING", "is_long_running": True},
            {"status": "ERROR", "is_long_running": False},
            {"status": "STALE", "is_long_running": False},
            {"status": "COMPLETED", "is_long_running": False},
        ]

        summary = AdminJobService._summary(jobs)

        self.assertEqual(summary["total"], 5)
        self.assertEqual(summary["active"], 2)
        self.assertEqual(summary["failed"], 2)
        self.assertEqual(summary["long_running"], 1)

    def test_admin_job_search_matches_entity_and_error(self):
        job = {
            "id": "job-1",
            "kind": "evaluation",
            "type": "Evaluation",
            "status": "ERROR",
            "entity": {"id": "question-1", "label": "Binary tree"},
            "error_message": "Model timeout",
        }

        self.assertTrue(AdminJobService._matches_search(job, "binary"))
        self.assertTrue(AdminJobService._matches_search(job, "timeout"))
        self.assertFalse(AdminJobService._matches_search(job, "moodle"))

    def test_admin_job_rejects_invalid_object_id(self):
        with self.assertRaises(ValueError):
            _parse_object_id("not-an-object-id", "job_id")

    def test_admin_job_status_filters_support_rollups(self):
        self.assertEqual(_generation_status_filter("active"), {"$in": ["queued", "processing"]})
        self.assertEqual(_generation_status_filter("retryable"), {"$in": ["failed"]})
        self.assertEqual(_uppercase_status_filter("active"), {"$in": ["QUEUED", "PROCESSING"]})
        self.assertEqual(_uppercase_status_filter("retryable"), {"$in": ["FAILED", "ERROR", "STALE"]})

    def test_admin_audit_list_filters_legacy_and_nested_records(self):
        actor_id = ObjectId()
        entity_id = ObjectId()
        version_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeAuditDatabase:
            def __init__(self):
                self.audit_logs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "action": "admin.job_cancel",
                            "actor_user_id": str(actor_id),
                            "actor_role": "Admin",
                            "entity_type": "generation",
                            "entity_id": str(entity_id),
                            "before": {"status": "processing"},
                            "after": {"status": "failed"},
                            "metadata": {"reason": "manual cancel"},
                            "created_at": now,
                        },
                        {
                            "_id": ObjectId(),
                            "action": "user.update",
                            "actor": {"user_id": actor_id, "role": "Admin"},
                            "entity": {"type": "user", "id": entity_id, "version_id": version_id},
                            "before": {"role": "Teacher"},
                            "after": {"role": "Reviewer"},
                            "changes": [{"field": "role"}],
                            "created_at": now - timedelta(minutes=5),
                        },
                        {
                            "_id": ObjectId(),
                            "action": "auth.demo_login",
                            "actor_user_id": str(ObjectId()),
                            "actor_role": "Teacher",
                            "entity_type": "auth",
                            "entity_id": "demo",
                            "created_at": now - timedelta(days=2),
                        },
                    ]
                )

        service = AdminAuditService(FakeAuditDatabase())

        actor_result = service.list(page=1, page_size=10, actor_user_id=str(actor_id))
        entity_result = service.list(
            page=1,
            page_size=10,
            entity_type="user",
            entity_id=str(entity_id),
        )
        search_result = service.list(page=1, page_size=10, search="manual cancel")
        date_result = service.list(
            page=1,
            page_size=10,
            date_from=now - timedelta(hours=1),
            date_to=now + timedelta(seconds=1),
        )
        paged = service.list(page=2, page_size=1)

        self.assertEqual(actor_result["total"], 2)
        self.assertEqual([item["action"] for item in actor_result["items"]], ["admin.job_cancel", "user.update"])
        self.assertEqual(entity_result["total"], 1)
        self.assertEqual(entity_result["items"][0]["entity"]["version_id"], str(version_id))
        self.assertEqual(search_result["total"], 1)
        self.assertEqual(search_result["items"][0]["metadata"]["reason"], "manual cancel")
        self.assertEqual(date_result["total"], 2)
        self.assertEqual(paged["items"][0]["action"], "user.update")

    def test_admin_job_list_filters_user_status_and_stale_jobs(self):
        owner_id = ObjectId()
        other_id = ObjectId()
        owner_document_id = ObjectId()
        other_document_id = ObjectId()
        now = datetime.now(timezone.utc)
        old = now - timedelta(minutes=settings.job_recovery_timeout_minutes + 10)

        class FakeJobDatabase:
            def __init__(self):
                self.generation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "processing",
                            "requested_by_user_id": owner_id,
                            "request": {"document_id": str(owner_document_id)},
                            "created_at": old,
                            "updated_at": old,
                        },
                        {
                            "_id": ObjectId(),
                            "status": "failed",
                            "requested_by_user_id": other_id,
                            "request": {"document_id": str(other_document_id)},
                            "error_message": "Model timeout",
                            "created_at": now,
                            "updated_at": now,
                        },
                    ]
                )
                self.evaluation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "ERROR",
                            "requested_by_user_id": owner_id,
                            "question_id": ObjectId(),
                            "question_version": 2,
                            "error": {"message": "Evaluator failed"},
                            "queued_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.documents = InMemoryCollection(
                    [
                        {"_id": owner_document_id, "uploaded_by_user_id": owner_id, "title": "Owner PDF"},
                        {"_id": other_document_id, "uploaded_by_user_id": other_id, "title": "Other PDF"},
                    ]
                )
                self.document_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "document_id": owner_document_id,
                            "job_type": "OCR",
                            "status": "PROCESSING",
                            "queued_at": old,
                            "updated_at": old,
                        },
                        {
                            "_id": ObjectId(),
                            "document_id": other_document_id,
                            "job_type": "OCR",
                            "status": "FAILED",
                            "queued_at": now,
                            "updated_at": now,
                        },
                    ]
                )

        service = AdminJobService(FakeJobDatabase())

        active = service.list_jobs(
            page=1,
            page_size=10,
            status="active",
            user_id=str(owner_id),
        )
        stale = service.list_jobs(page=1, page_size=10, stale_only=True)
        retryable = service.list_jobs(page=1, page_size=10, status="retryable")
        recent = service.list_jobs(
            page=1,
            page_size=10,
            date_from=now - timedelta(minutes=1),
            date_to=now + timedelta(seconds=1),
        )

        self.assertEqual(active["summary"]["total"], 2)
        self.assertEqual(active["summary"]["long_running"], 2)
        self.assertEqual({item["kind"] for item in active["items"]}, {"generation", "document"})
        self.assertEqual(stale["summary"]["long_running"], 2)
        self.assertEqual(retryable["summary"]["failed"], 3)
        self.assertEqual(recent["summary"]["total"], 3)
        self.assertEqual({item["kind"] for item in recent["items"]}, {"generation", "evaluation", "document"})
        with self.assertRaises(ValueError):
            service.list_jobs(
                page=1,
                page_size=10,
                date_from=now + timedelta(days=1),
                date_to=now,
            )

    def test_admin_job_cancel_updates_jobs_questions_and_audit(self):
        admin = _current_user("Admin")
        generation_job_id = ObjectId()
        evaluation_job_id = ObjectId()
        question_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeJobDatabase:
            def __init__(self):
                self.generation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": generation_job_id,
                            "status": "processing",
                            "requested_by_user_id": ObjectId(),
                            "request": {},
                            "created_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.evaluation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": evaluation_job_id,
                            "status": "PROCESSING",
                            "question_id": question_id,
                            "requested_by_user_id": ObjectId(),
                            "queued_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "evaluation_status": "PROCESSING",
                            "quality_summary": {"latest_evaluation_job_id": evaluation_job_id},
                            "updated_at": now,
                        }
                    ]
                )

        db = FakeJobDatabase()
        audit_events = []
        original_audit = admin_jobs_module.record_audit_event
        try:
            admin_jobs_module.record_audit_event = lambda **kwargs: audit_events.append(kwargs)
            service = AdminJobService(db)
            generation_result = service.cancel_job("generation", str(generation_job_id), admin)
            evaluation_result = service.cancel_job("evaluation", str(evaluation_job_id), admin)
        finally:
            admin_jobs_module.record_audit_event = original_audit

        self.assertEqual(generation_result["job"]["status"], "failed")
        self.assertIn("Cancelled by admin", generation_result["job"]["error_message"])
        self.assertEqual(evaluation_result["job"]["status"], "STALE")
        question = db.questions.find_one({"_id": question_id})
        self.assertEqual(question["evaluation_status"], "STALE")
        self.assertIn("Cancelled by admin", question["quality_summary"]["error"]["message"])
        self.assertEqual([event["action"] for event in audit_events], ["admin.job_cancel", "admin.job_cancel"])
        self.assertEqual({event["entity_type"] for event in audit_events}, {"generation", "evaluation"})

    def test_admin_job_retry_evaluation_dispatches_background_and_audit(self):
        admin = _current_user("Admin")
        requester_id = ObjectId()
        job_id = ObjectId()
        question_id = ObjectId()
        queued_job_id = ObjectId()
        now = datetime.now(timezone.utc)

        class FakeBackgroundTasks:
            def __init__(self):
                self.tasks = []

            def add_task(self, fn, *args, **kwargs):
                self.tasks.append((fn, args, kwargs))

        class FakeWorkflowService:
            calls = []

            def __init__(self, database):
                self.db = database

            def enqueue_auto_evaluation(self, question_id_arg, **kwargs):
                self.__class__.calls.append((question_id_arg, kwargs))
                return {"_id": queued_job_id, "status": "QUEUED"}

        class FakeJobDatabase:
            def __init__(self):
                self.evaluation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": job_id,
                            "status": "ERROR",
                            "question_id": question_id,
                            "requested_by_user_id": requester_id,
                            "evaluator_model_code": "qwen",
                            "queued_at": now,
                            "updated_at": now,
                        }
                    ]
                )
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "current_version": 3,
                        }
                    ]
                )

        db = FakeJobDatabase()
        background_tasks = FakeBackgroundTasks()
        audit_events = []
        original_workflow_service = admin_jobs_module.QuestionWorkflowService
        original_audit = admin_jobs_module.record_audit_event
        try:
            admin_jobs_module.QuestionWorkflowService = FakeWorkflowService
            admin_jobs_module.record_audit_event = lambda **kwargs: audit_events.append(kwargs)
            result = AdminJobService(db).retry_job(
                "evaluation",
                str(job_id),
                background_tasks,
                admin,
            )
        finally:
            admin_jobs_module.QuestionWorkflowService = original_workflow_service
            admin_jobs_module.record_audit_event = original_audit

        self.assertEqual(result["job"]["_id"], str(queued_job_id))
        self.assertEqual(FakeWorkflowService.calls[0][0], str(question_id))
        self.assertEqual(FakeWorkflowService.calls[0][1]["expected_version"], 3)
        self.assertEqual(FakeWorkflowService.calls[0][1]["requested_by_user_id"], requester_id)
        self.assertEqual(FakeWorkflowService.calls[0][1]["evaluator_model_code"], "qwen")
        self.assertEqual(FakeWorkflowService.calls[0][1]["trigger"], "ADMIN_RETRY")
        self.assertEqual(background_tasks.tasks[0][1], (queued_job_id,))
        self.assertEqual(audit_events[0]["action"], "admin.job_retry")
        self.assertEqual(audit_events[0]["entity_type"], "evaluation")
        self.assertEqual(audit_events[0]["metadata"]["new_job_id"], queued_job_id)

    def test_admin_overview_summarizes_operational_state(self):
        now = datetime.now(timezone.utc)
        old = now - timedelta(minutes=settings.job_recovery_timeout_minutes + 5)
        processing_document_id = ObjectId()
        failed_document_id = ObjectId()

        class FakeOverviewDatabase:
            def __init__(self):
                self.users = InMemoryCollection(
                    [
                        {"_id": ObjectId(), "role": "Admin", "is_active": True},
                        {"_id": ObjectId(), "role": "Teacher", "is_active": True},
                        {"_id": ObjectId(), "role": "Reviewer", "is_active": True},
                        {"_id": ObjectId(), "role": "Teacher", "is_active": False},
                    ]
                )
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "lifecycle_status": "ACTIVE",
                            "review_status": "PENDING",
                            "publication_status": "NOT_PUBLISHED",
                            "quality_summary": {"color": "GREEN"},
                        },
                        {
                            "_id": ObjectId(),
                            "lifecycle_status": "ACTIVE",
                            "review_status": "APPROVED",
                            "publication_status": "PUBLISHED",
                            "quality_summary": {"color": "YELLOW"},
                        },
                        {
                            "_id": ObjectId(),
                            "lifecycle_status": "ACTIVE",
                            "review_status": "REJECTED",
                            "publication_status": "FAILED",
                            "quality_summary": {"color": "RED"},
                        },
                        {
                            "_id": ObjectId(),
                            "lifecycle_status": "ARCHIVED",
                            "review_status": "PENDING",
                            "publication_status": "NOT_PUBLISHED",
                        },
                    ]
                )
                self.documents = InMemoryCollection(
                    [
                        {"_id": processing_document_id, "archived_at": None, "status": "PROCESSING"},
                        {"_id": failed_document_id, "archived_at": None, "status": "FAILED"},
                        {"_id": ObjectId(), "archived_at": now, "status": "FAILED"},
                    ]
                )
                self.generation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "failed",
                            "request": {"document_id": "doc-1"},
                            "error_message": "Model timeout",
                            "created_at": old,
                            "updated_at": old,
                        },
                        {
                            "_id": ObjectId(),
                            "status": "processing",
                            "request": {"document_id": "doc-2"},
                            "error_message": None,
                            "created_at": old,
                            "updated_at": old,
                        },
                    ]
                )
                self.evaluation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "COMPLETED",
                            "evaluator_model_code": "qwen",
                            "duration_ms": 1200,
                            "usage": {
                                "prompt_tokens": 100,
                                "completion_tokens": 40,
                                "total_tokens": 140,
                                "cost_usd": 0.012,
                            },
                            "queued_at": now,
                            "started_at": now,
                            "finished_at": now,
                            "updated_at": now,
                        },
                        {
                            "_id": ObjectId(),
                            "status": "ERROR",
                            "evaluator_model_code": "qwen",
                            "duration_ms": 800,
                            "token_usage": {
                                "prompt_tokens": 50,
                                "completion_tokens": 10,
                                "total_tokens": 60,
                            },
                            "billing": {"cost_usd": 0.004},
                            "queued_at": now,
                            "started_at": now,
                            "finished_at": now,
                            "updated_at": now,
                            "error": {"message": "Evaluator failed"},
                        },
                    ]
                )
                self.document_jobs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "document_id": processing_document_id,
                            "job_type": "OCR",
                            "status": "PROCESSING",
                            "queued_at": old,
                            "started_at": old,
                            "updated_at": old,
                        },
                        {
                            "_id": ObjectId(),
                            "document_id": failed_document_id,
                            "job_type": "CHUNK",
                            "status": "FAILED",
                            "queued_at": now,
                            "started_at": now,
                            "finished_at": now,
                            "updated_at": now,
                            "error": {"message": "Chunk failed"},
                        },
                    ]
                )
                self.generation_runs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "COMPLETED",
                            "model": {"model_code": "deepseek"},
                            "execution": {
                                "latency_ms": 2000,
                                "usage": {
                                    "prompt_tokens": 300,
                                    "completion_tokens": 120,
                                    "total_tokens": 420,
                                    "cost_usd": 0.03,
                                },
                            },
                            "created_at": now,
                            "finished_at": now,
                            "updated_at": now,
                        },
                        {
                            "_id": ObjectId(),
                            "status": "FAILED",
                            "model": {
                                "provider": "deepseek",
                                "config": {"input_cost_per_1k": 0.05, "output_cost_per_1k": 0.3},
                            },
                            "execution": {
                                "latency_ms": 4000,
                                "prompt_tokens": 80,
                                "completion_tokens": 20,
                            },
                            "created_at": now,
                            "finished_at": now,
                            "updated_at": now,
                            "error": {"message": "Generation failed"},
                        },
                    ]
                )
                self.moodle_targets = InMemoryCollection(
                    [
                        {"_id": ObjectId(), "site_key": "demo", "is_active": True},
                        {"_id": ObjectId(), "site_key": "old", "is_active": False},
                    ]
                )
                self.moodle_publications = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "status": "PUBLISHED",
                            "publication_mode": "MOCK",
                            "request_payload": {"mock": True},
                        },
                        {"_id": ObjectId(), "status": "FAILED", "request_payload": {"mock": True}},
                    ]
                )
                self.audit_logs = InMemoryCollection(
                    [
                        {
                            "_id": ObjectId(),
                            "action": "admin.job_retry",
                            "actor_user_id": str(ObjectId()),
                            "entity_type": "job",
                            "entity_id": "job-1",
                            "created_at": now,
                        }
                    ]
                )

        overview = AdminOverviewService(FakeOverviewDatabase()).overview()

        self.assertEqual(overview["users"]["active"], 3)
        self.assertEqual(overview["questions"]["pending"], 1)
        self.assertEqual(overview["questions"]["published"], 1)
        self.assertEqual(overview["questions"]["quality"]["green"], 1)
        self.assertEqual(overview["questions"]["quality"]["yellow"], 1)
        self.assertEqual(overview["questions"]["quality"]["red"], 1)
        self.assertEqual(overview["documents"]["failed"], 1)
        self.assertEqual(overview["jobs"]["failed"], 3)
        self.assertEqual(overview["jobs"]["long_running"], 2)
        job_breakdown = {item["key"]: item for item in overview["jobs"]["breakdown"]}
        self.assertEqual(job_breakdown["generation"]["failed"], 1)
        self.assertEqual(job_breakdown["evaluation"]["failed"], 1)
        self.assertEqual(job_breakdown["ocr"]["active"], 1)
        self.assertEqual(job_breakdown["chunk"]["failed"], 1)
        self.assertEqual(overview["moodle"]["active_targets"], 1)
        self.assertEqual(overview["moodle"]["publications"]["simulated"], 2)
        self.assertEqual(overview["moodle"]["publications"]["failed"], 1)
        model_rows = {item["key"]: item for item in overview["model_performance"]}
        self.assertEqual(model_rows["evaluation:qwen"]["error_rate"], 0.5)
        self.assertEqual(model_rows["evaluation:qwen"]["avg_latency_ms"], 1000)
        self.assertEqual(model_rows["evaluation:qwen"]["total_tokens"], 200)
        self.assertEqual(model_rows["evaluation:qwen"]["prompt_tokens"], 150)
        self.assertEqual(model_rows["evaluation:qwen"]["completion_tokens"], 50)
        self.assertEqual(model_rows["evaluation:qwen"]["cost_usd"], 0.016)
        self.assertEqual(model_rows["generation:deepseek"]["error_rate"], 0.5)
        self.assertEqual(model_rows["generation:deepseek"]["avg_latency_ms"], 3000)
        self.assertEqual(model_rows["generation:deepseek"]["total_tokens"], 520)
        self.assertEqual(model_rows["generation:deepseek"]["cost_usd"], 0.04)
        self.assertEqual(overview["model_usage_summary"]["total_requests"], 4)
        self.assertEqual(overview["model_usage_summary"]["total_tokens"], 720)
        self.assertEqual(overview["model_usage_summary"]["cost_usd"], 0.056)
        self.assertEqual(overview["model_usage_summary"]["avg_latency_ms"], 2000)
        self.assertIn("Model timeout", [item["error_message"] for item in overview["recent_jobs"]])
        self.assertEqual(overview["recent_audit"][0]["action"], "admin.job_retry")
        attention = {item["key"]: item for item in overview["attention"]}
        self.assertEqual(attention["retryable_jobs"]["severity"], "danger")
        self.assertEqual(attention["failed_documents"]["count"], 1)

    def test_job_recovery_marks_only_stale_active_jobs(self):
        now = datetime.now(timezone.utc)
        old = now - timedelta(minutes=180)
        fresh = now - timedelta(minutes=5)
        question_id = ObjectId()
        evaluation_job_id = ObjectId()
        document_id = ObjectId()
        chunk_job_id = ObjectId()
        chunk_set_id = ObjectId()
        fresh_generation_id = ObjectId()

        class FakeRecoveryDatabase:
            def __init__(self):
                self.generation_jobs = InMemoryCollection(
                    [
                        {"_id": ObjectId(), "status": "queued", "updated_at": old},
                        {"_id": fresh_generation_id, "status": "processing", "updated_at": fresh},
                    ]
                )
                self.evaluation_jobs = InMemoryCollection(
                    [
                        {
                            "_id": evaluation_job_id,
                            "status": "PROCESSING",
                            "question_id": question_id,
                            "updated_at": old,
                        }
                    ]
                )
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "evaluation_status": "PROCESSING",
                            "quality_summary": {"latest_evaluation_job_id": evaluation_job_id},
                        }
                    ]
                )
                self.document_jobs = InMemoryCollection(
                    [
                        {
                            "_id": chunk_job_id,
                            "document_id": document_id,
                            "job_type": "CHUNK",
                            "status": "PROCESSING",
                            "updated_at": old,
                        }
                    ]
                )
                self.documents = InMemoryCollection(
                    [
                        {
                            "_id": document_id,
                            "archived_at": None,
                            "status": "PROCESSING",
                            "pipeline_summary": {"chunk_status": "PROCESSING"},
                        }
                    ]
                )
                self.chunk_sets = InMemoryCollection(
                    [{"_id": chunk_set_id, "chunk_job_id": chunk_job_id, "status": "PROCESSING"}]
                )
                self.chunk_embeddings = InMemoryCollection(
                    [{"_id": ObjectId(), "chunk_set_id": chunk_set_id, "status": "PENDING"}]
                )

        db = FakeRecoveryDatabase()
        original_get_database = job_recovery.get_database
        try:
            job_recovery.get_database = lambda: db
            result = job_recovery.recover_stale_jobs(timeout_minutes=60)
        finally:
            job_recovery.get_database = original_get_database

        self.assertEqual(
            result,
            {"generation_failed": 1, "evaluation_stale": 1, "document_failed": 1},
        )
        self.assertEqual(db.generation_jobs.find_one({"_id": fresh_generation_id})["status"], "processing")
        self.assertEqual(db.evaluation_jobs.find_one({"_id": evaluation_job_id})["status"], "STALE")
        question = db.questions.find_one({"_id": question_id})
        self.assertEqual(question["evaluation_status"], "STALE")
        self.assertIn("exceeded 60 minute recovery timeout", question["quality_summary"]["error"]["message"])
        document = db.documents.find_one({"_id": document_id})
        self.assertEqual(document["status"], "FAILED")
        self.assertEqual(document["pipeline_summary"]["chunk_status"], "FAILED")
        self.assertEqual(db.chunk_sets.find_one({"_id": chunk_set_id})["status"], "FAILED")
        self.assertEqual(db.chunk_embeddings.find_one({"chunk_set_id": chunk_set_id})["status"], "FAILED")

    def test_admin_audit_service_normalizes_flat_and_nested_records(self):
        class FakeCursor(list):
            def sort(self, *_args):
                return self

            def skip(self, *_args):
                return self

            def limit(self, *_args):
                return self

        class FakeAuditLogs:
            def __init__(self, records):
                self.records = records
                self.match = None

            def count_documents(self, match):
                self.match = match
                return len(self.records)

            def find(self, match):
                self.match = match
                return FakeCursor(self.records)

        class FakeDatabase:
            def __init__(self, records):
                self.audit_logs = FakeAuditLogs(records)

        now = datetime.now(timezone.utc)
        actor_id = ObjectId()
        question_id = ObjectId()
        review_id = ObjectId()
        records = [
            {
                "_id": ObjectId(),
                "action": "user.admin_update",
                "actor_user_id": str(actor_id),
                "actor_role": "Admin",
                "entity_type": "user",
                "entity_id": "user-1",
                "before": {"role": "Teacher"},
                "after": {"role": "Reviewer"},
                "created_at": now,
            },
            {
                "_id": ObjectId(),
                "action": "QUESTION_APPROVED",
                "actor": {"type": "USER", "user_id": actor_id},
                "entity": {"type": "QUESTION", "id": question_id, "version_id": ObjectId()},
                "changes": [{"path": "review_status", "old_value": "PENDING", "new_value": "APPROVED"}],
                "metadata": {"review_id": review_id},
                "created_at": now,
            },
        ]
        db = FakeDatabase(records)
        service = AdminAuditService(db)

        result = service.list(
            page=1,
            page_size=10,
            actor_user_id=str(actor_id),
            entity_type="question",
            entity_id=str(question_id),
            action="QUESTION_APPROVED",
            search="approved",
        )

        self.assertEqual(result["total"], 2)
        self.assertIn("$and", db.audit_logs.match)
        self.assertEqual(result["items"][0]["actor"]["role"], "Admin")
        self.assertEqual(result["items"][0]["before"], {"role": "Teacher"})
        self.assertEqual(result["items"][1]["actor"]["user_id"], str(actor_id))
        self.assertEqual(result["items"][1]["entity"]["id"], str(question_id))
        self.assertEqual(result["items"][1]["metadata"]["review_id"], str(review_id))

    def test_reviewer_dashboard_summarizes_workload_and_reviews(self):
        def get_path(record, path):
            current = record
            for part in path.split("."):
                if not isinstance(current, dict) or part not in current:
                    return None
                current = current[part]
            return current

        def matches(record, match):
            if "$and" in match:
                return all(matches(record, clause) for clause in match["$and"])
            if "$or" in match:
                return any(matches(record, clause) for clause in match["$or"])
            for key, expected in match.items():
                actual = get_path(record, key)
                if isinstance(expected, dict):
                    if "$exists" in expected and (actual is not None) != expected["$exists"]:
                        return False
                    if "$in" in expected and actual not in expected["$in"]:
                        return False
                    if "$gte" in expected and (actual is None or actual < expected["$gte"]):
                        return False
                    if "$lte" in expected and (actual is None or actual > expected["$lte"]):
                        return False
                    continue
                if actual != expected:
                    return False
            return True

        class FakeCursor(list):
            def sort(self, *_args):
                return self

            def limit(self, count):
                return FakeCursor(self[:count])

        class FakeCollection:
            def __init__(self, records):
                self.records = records

            def count_documents(self, match):
                return sum(1 for record in self.records if matches(record, match))

            def find(self, match=None, *_args):
                match = match or {}
                return FakeCursor([record for record in self.records if matches(record, match)])

        class FakeDatabase:
            def __init__(self, **collections):
                for name, records in collections.items():
                    setattr(self, name, FakeCollection(records))

        now = datetime.now(timezone.utc)
        reviewer_id = ObjectId()
        other_reviewer_id = ObjectId()
        subject_id = ObjectId()
        version_ids = [ObjectId(), ObjectId()]
        base_question = {
            "schema_version": SCHEMA_VERSION,
            "lifecycle_status": "ACTIVE",
            "review_status": "PENDING",
        }
        db = FakeDatabase(
            questions=[
                {**base_question, "_id": ObjectId(), "review_assignment": {"status": "UNASSIGNED"}},
                {
                    **base_question,
                    "_id": ObjectId(),
                    "review_assignment": {"status": "ASSIGNED", "reviewer_user_id": reviewer_id},
                },
                {
                    **base_question,
                    "_id": ObjectId(),
                    "review_assignment": {
                        "status": "IN_REVIEW",
                        "reviewer_user_id": reviewer_id,
                        "lock_expires_at": now - timedelta(minutes=5),
                    },
                },
                {
                    **base_question,
                    "_id": ObjectId(),
                    "review_assignment": {
                        "status": "IN_REVIEW",
                        "reviewer_user_id": other_reviewer_id,
                        "lock_expires_at": now + timedelta(minutes=30),
                    },
                },
            ],
            question_reviews=[
                {
                    "_id": ObjectId(),
                    "question_version_id": version_ids[0],
                    "reviewer_user_id": reviewer_id,
                    "decision": "APPROVED",
                    "override": {"applied": True},
                    "revision_issues": [],
                    "reviewed_at": now - timedelta(days=1),
                },
                {
                    "_id": ObjectId(),
                    "question_version_id": version_ids[1],
                    "reviewer_user_id": reviewer_id,
                    "decision": "NEEDS_REVISION",
                    "override": {"applied": False},
                    "revision_issues": [{"title": "Sửa đáp án"}],
                    "reviewed_at": now - timedelta(days=2),
                },
            ],
            audit_logs=[
                {
                    "action": "QUESTION_APPROVED",
                    "actor": {"user_id": reviewer_id},
                    "metadata": {"review_assignment": {"claimed_at": now - timedelta(hours=2)}},
                    "created_at": now,
                },
                {
                    "action": "QUESTION_NEEDS_REVISION",
                    "actor": {"user_id": reviewer_id},
                    "metadata": {"review_assignment": {"assigned_at": now - timedelta(hours=4)}},
                    "created_at": now,
                },
            ],
            question_versions=[
                {"_id": version_ids[0], "classification": {"subject": {"id": subject_id}}},
                {"_id": version_ids[1], "classification": {"subject": {"id": subject_id}}},
            ],
            subjects=[
                {"_id": subject_id, "subject_code": "CTDL", "subject_name": "Cấu trúc dữ liệu"},
            ],
        )
        dashboard = QuestionWorkflowService(db).review_dashboard(_current_user("Reviewer", reviewer_id))

        self.assertEqual(dashboard["workload"]["pending"], 4)
        self.assertEqual(dashboard["workload"]["unassigned"], 1)
        self.assertEqual(dashboard["workload"]["assigned"], 1)
        self.assertEqual(dashboard["workload"]["in_review"], 2)
        self.assertEqual(dashboard["workload"]["lock_expired"], 1)
        self.assertEqual(dashboard["workload"]["mine"], 2)
        self.assertEqual(dashboard["performance"]["reviews_30d"], 2)
        self.assertEqual(dashboard["performance"]["approval_rate"], 0.5)
        self.assertEqual(dashboard["performance"]["override_count"], 1)
        self.assertEqual(dashboard["performance"]["revision_issues"], 1)
        self.assertEqual(dashboard["performance"]["average_review_hours"], 3.0)
        self.assertEqual(dashboard["subjects"][0]["label"], "CTDL")
        self.assertEqual(dashboard["subjects"][0]["reviewed"], 2)

    def test_notification_service_creates_and_marks_read(self):
        class FakeCursor(list):
            def sort(self, *_args):
                return self

            def skip(self, count):
                return FakeCursor(self[count:])

            def limit(self, count):
                return FakeCursor(self[:count])

        class FakeUpdateResult:
            def __init__(self, modified_count):
                self.modified_count = modified_count

        class FakeNotifications:
            def __init__(self):
                self.records = []

            @staticmethod
            def _matches(record, query):
                return all(record.get(key) == value for key, value in query.items())

            def insert_one(self, record):
                self.records.append(record)

            def count_documents(self, query):
                return sum(1 for record in self.records if self._matches(record, query))

            def find(self, query):
                return FakeCursor([record for record in self.records if self._matches(record, query)])

            def find_one_and_update(self, query, update, return_document=None):
                for record in self.records:
                    if self._matches(record, query):
                        record.update(update.get("$set", {}))
                        return record
                return None

            def update_many(self, query, update):
                modified = 0
                for record in self.records:
                    if self._matches(record, query):
                        record.update(update.get("$set", {}))
                        modified += 1
                return FakeUpdateResult(modified)

        class FakeDatabase:
            def __init__(self):
                self.notifications = FakeNotifications()

        db = FakeDatabase()
        service = NotificationService(db)
        teacher = _current_user("Teacher")
        reviewer = _current_user("Reviewer")
        question_id = ObjectId()
        version_id = ObjectId()
        notification = service.notify_review_decision(
            question={
                "_id": question_id,
                "question_code": "Q-1",
                "created_by_user_id": teacher.id,
            },
            version={"_id": version_id, "created_by_user_id": teacher.id},
            review={"decision": "NEEDS_REVISION", "note": "Cần sửa đáp án"},
            actor_user_id=reviewer.id,
        )

        self.assertIsNotNone(notification)
        self.assertEqual(notification["type"], "QUESTION_NEEDS_REVISION")
        self.assertEqual(notification["entity"]["id"], str(question_id))
        self.assertEqual(service.unread_count(teacher), 1)
        page = service.list(teacher, 1, 10)
        self.assertEqual(page["total"], 1)
        self.assertFalse(page["items"][0]["is_read"])

        marked = service.mark_read(notification["id"], teacher)
        self.assertTrue(marked["is_read"])
        self.assertEqual(service.unread_count(teacher), 0)

    def test_notification_service_notifies_reviewer_when_question_resubmitted(self):
        teacher = _current_user("Teacher")
        reviewer = _current_user("Reviewer")
        question_id = ObjectId()
        version_id = ObjectId()
        review_id = ObjectId()

        class FakeDatabase:
            def __init__(self):
                self.notifications = InMemoryCollection()
                self.questions = InMemoryCollection(
                    [
                        {
                            "_id": question_id,
                            "schema_version": SCHEMA_VERSION,
                            "lifecycle_status": "ACTIVE",
                            "current_version_id": version_id,
                            "latest_review_id": review_id,
                            "question_code": "Q-2",
                            "created_by_user_id": teacher.id,
                        }
                    ]
                )
                self.question_versions = InMemoryCollection(
                    [{"_id": version_id, "created_by_user_id": teacher.id}]
                )
                self.question_reviews = InMemoryCollection(
                    [
                        {
                            "_id": review_id,
                            "question_id": question_id,
                            "reviewer_user_id": reviewer.id,
                            "reviewed_at": datetime.now(timezone.utc),
                        }
                    ]
                )

        db = FakeDatabase()
        service = NotificationService(db)

        ignored = service.notify_question_resubmitted(
            question_id=question_id,
            previous_review_status="DRAFT",
            actor_user_id=teacher.id,
        )
        notifications = service.notify_question_resubmitted(
            question_id=question_id,
            previous_review_status="NEEDS_REVISION",
            actor_user_id=teacher.id,
        )

        self.assertEqual(ignored, [])
        self.assertEqual(len(notifications), 1)
        self.assertEqual(notifications[0]["type"], "QUESTION_RESUBMITTED")
        self.assertEqual(notifications[0]["link"], f"/kiem-duyet?questionId={question_id}")
        self.assertEqual(notifications[0]["entity"]["version_id"], str(version_id))
        self.assertEqual(service.unread_count(reviewer), 1)

    def test_question_hash_is_order_independent(self):
        self.assertEqual(
            stable_hash({"content": "Q", "data": {"a": 1, "b": 2}}),
            stable_hash({"data": {"b": 2, "a": 1}, "content": "Q"}),
        )

    def test_document_has_archive_state(self):
        self.assertIn(DocumentStatus.ARCHIVED, set(DocumentStatus))

    def test_services_depend_on_abstractions_via_constructor(self):
        self.assertIn("repository", inspect.signature(DocumentService).parameters)
        self.assertIn("repository", inspect.signature(QuestionService).parameters)
        self.assertIn("references", inspect.signature(QuestionService).parameters)

    def test_auth_and_rag_databases_are_separate(self):
        self.assertEqual(settings.auth_db_name, "NCKH")
        self.assertEqual(settings.rag_db_name, "rag_database")
        self.assertNotEqual(settings.auth_db_name, settings.rag_db_name)

    def test_generation_request_accepts_question_plan(self):
        req = QuestionGenerateRequest(
            document_id="507f1f77bcf86cd799439011",
            bloom_level=BloomLevel.HIEU,
            question_plan=[
                {"question_type": QuestionType.TRAC_NGHIEM, "bloom_level": BloomLevel.HIEU, "num_questions": 3},
                {"question_type": QuestionType.DUNG_SAI, "bloom_level": BloomLevel.PHAN_TICH, "num_questions": 2},
            ],
            instruction="Tập trung vào cây nhị phân tìm kiếm.",
        )

        self.assertEqual(sum(item.num_questions for item in req.effective_plan()), 5)
        self.assertEqual(req.effective_plan()[1].bloom_level, BloomLevel.PHAN_TICH)
        self.assertEqual(req.instruction, "Tập trung vào cây nhị phân tìm kiếm.")

    def test_generation_request_rejects_oversized_plan(self):
        with self.assertRaises(ValidationError):
            QuestionGenerateRequest(
                document_id="507f1f77bcf86cd799439011",
                bloom_level=BloomLevel.HIEU,
                question_plan=[
                    {"question_type": QuestionType.TRAC_NGHIEM, "num_questions": 10},
                    {"question_type": QuestionType.DUNG_SAI, "num_questions": 10},
                    {"question_type": QuestionType.DIEN_KHUYET, "num_questions": 1},
                ],
            )

    def test_generated_question_can_return_persisted_metadata(self):
        question = GeneratedQuestion(
            question="Cấu trúc dữ liệu nào dùng nguyên tắc FIFO?",
            options={"A": "Stack", "B": "Queue", "C": "Tree", "D": "Graph"},
            correct_answer="B",
            explanation="Queue xử lý phần tử vào trước ra trước.",
            question_type="trac_nghiem",
            bloom_level="hieu",
            source_context="Queue là hàng đợi FIFO.",
            question_id="507f1f77bcf86cd799439011",
            question_code="Q-507F1F77BCF86CD799439011",
            current_version=1,
            current_version_id="507f1f77bcf86cd799439012",
        )

        self.assertEqual(question.question_id, "507f1f77bcf86cd799439011")
        self.assertEqual(question.current_version, 1)

    def test_generation_plan_summary_reports_shortfall(self):
        summary = GenerationPlanSummary(
            plan_index=2,
            question_type="trac_nghiem",
            bloom_level="phan_tich",
            requested_count=5,
            parsed_count=5,
            valid_count=4,
            duplicate_count=1,
            saved_count=3,
            skipped_count=2,
            warnings=["Bỏ 1 câu trùng nội dung.", "Lưu thiếu 2 câu so với yêu cầu."],
        )

        self.assertEqual(summary.plan_index, 2)
        self.assertEqual(summary.skipped_count, 2)
        self.assertIn("Lưu thiếu 2 câu so với yêu cầu.", summary.warnings)

    def test_key_validators_require_schema_version(self):
        for collection in ("users", "documents", "questions", "question_versions", "evaluation_jobs", "moodle_targets"):
            required = VALIDATORS[collection]["$jsonSchema"]["required"]
            self.assertIn("schema_version", required)

    def test_user_validator_accepts_reviewer_role(self):
        role_enum = set(VALIDATORS["users"]["$jsonSchema"]["properties"]["role"]["enum"])
        self.assertEqual(role_enum, {role.value for role in RoleEnum})

    def test_question_validator_accepts_ai_draft_status(self):
        review_status_enum = set(
            VALIDATORS["questions"]["$jsonSchema"]["properties"]["review_status"]["enum"]
        )
        self.assertIn("DRAFT", review_status_enum)

    def test_question_validator_accepts_evaluation_queue_statuses(self):
        evaluation_status_enum = set(
            VALIDATORS["questions"]["$jsonSchema"]["properties"]["evaluation_status"]["enum"]
        )
        self.assertIn("QUEUED", evaluation_status_enum)
        self.assertIn("ERROR", evaluation_status_enum)
        self.assertIn("STALE", evaluation_status_enum)

    def test_evaluation_job_validator_tracks_worker_statuses(self):
        status_enum = set(VALIDATORS["evaluation_jobs"]["$jsonSchema"]["properties"]["status"]["enum"])
        self.assertTrue({"QUEUED", "PROCESSING", "COMPLETED", "ERROR", "STALE"}.issubset(status_enum))

    def test_question_validator_allows_aggregate_owner(self):
        question_properties = VALIDATORS["questions"]["$jsonSchema"]["properties"]
        self.assertEqual(question_properties["created_by_user_id"]["bsonType"], ["objectId", "null"])

    def test_question_validator_allows_review_assignment_state(self):
        question_properties = VALIDATORS["questions"]["$jsonSchema"]["properties"]
        self.assertEqual(question_properties["review_assignment"]["bsonType"], "object")

    def test_auth_user_is_minimal_uid_and_token_link(self):
        schema = VALIDATORS["User"]["$jsonSchema"]
        self.assertEqual(set(schema["required"]), {"uid", "token"})
        self.assertEqual(set(schema["properties"]), {"_id", "uid", "token"})
        self.assertFalse(schema["additionalProperties"])

    def test_question_sources_cannot_mix_documents(self):
        first_document_id = ObjectId()
        second_document_id = ObjectId()
        chunks = {
            ObjectId(): {
                "_id": ObjectId(),
                "document_id": first_document_id,
                "chunk_set_id": ObjectId(),
                "content": "A",
                "content_hash": "a",
            },
            ObjectId(): {
                "_id": ObjectId(),
                "document_id": second_document_id,
                "chunk_set_id": ObjectId(),
                "content": "B",
                "content_hash": "b",
            },
        }

        class References:
            def find_chunk(self, chunk_id):
                return chunks.get(chunk_id)

            def find_document(self, _document_id):
                return None

            def find_subject(self, _subject_id):
                return None

        service = QuestionService(repository=None, references=References())
        with self.assertRaises(ValueError):
            service._sources([str(chunk_id) for chunk_id in chunks])

    def test_rag_heading_match_normalizes_vietnamese_text(self):
        normalized_target = _normalize_heading_text("do thi")
        self.assertTrue(
            _heading_matches_target({"heading": "Đồ thị"}, normalized_target)
        )

        normalized_target = _normalize_heading_text("cay nhi phan")
        self.assertTrue(
            _heading_matches_target(
                {"heading_path_text": "Cấu trúc dữ liệu > Cây nhị phân tìm kiếm"},
                normalized_target,
            )
        )
        self.assertFalse(
            _heading_matches_target({"heading": "Hàng đợi"}, normalized_target)
        )

    def test_output_format_keeps_question_type_option_shapes(self):
        prompt_path = Path(__file__).resolve().parents[1] / "prompts" / "output_format.txt"
        output_format = prompt_path.read_text(encoding="utf-8")

        self.assertIn("trac_nghiem", output_format)
        self.assertIn('"A", "B", "C", "D"', output_format)
        self.assertIn("dung_sai", output_format)
        self.assertIn('"A": "Đúng", "B": "Sai"', output_format)

    def test_question_rule_is_loaded_into_generation_prompt(self):
        prompt = PromptBuilder().build(
            context="Stack hoạt động theo nguyên tắc LIFO.",
            bloom_level="hieu",
            question_type="trac_nghiem",
            num_questions=1,
        )

        self.assertIn("FORBIDDEN QUESTION RULES", prompt)
        self.assertIn("Do not create source-referencing questions", prompt)
        self.assertIn("If any rule is violated, reject the item and generate a replacement", prompt)

    def test_mcq_validation_rejects_two_option_shape(self):
        error = _check_type_format(
            {
                "question": "Câu hỏi mẫu?",
                "options": {"A": "Một", "B": "Hai"},
                "correct_answer": "A",
            },
            "trac_nghiem",
        )
        self.assertIn("4 lựa chọn", error)

        self.assertIsNone(
            _check_type_format(
                {
                    "question": "Câu hỏi mẫu?",
                    "options": {"A": "Một", "B": "Hai", "C": "Ba", "D": "Bốn"},
                    "correct_answer": "A",
                },
                "trac_nghiem",
            )
        )

    def test_retry_prompt_reinforces_mcq_option_shape(self):
        prompt = _build_retry_prompt(
            original_prompt="ORIGINAL",
            question_type="trac_nghiem",
            bloom_level="hieu",
            missing_count=2,
            validation_errors=["trac_nghiem phải có đúng 4 lựa chọn A/B/C/D"],
            avoid_questions=["Câu đã nhận"],
        )
        self.assertIn('exactly "A", "B", "C", "D"', prompt)
        self.assertIn("Generate exactly 2 additional questions", prompt)

    def test_multi_answer_validation_allows_five_options(self):
        self.assertIsNone(
            _check_type_format(
                {
                    "question": "Câu hỏi mẫu?",
                    "options": {
                        "A": "Một",
                        "B": "Hai",
                        "C": "Ba",
                        "D": "Bốn",
                        "E": "Năm",
                    },
                    "correct_answer": "A, C, E",
                },
                "nhieu_lua_chon",
            )
        )

    def test_multi_answer_validation_rejects_all_options_correct(self):
        error = _check_type_format(
            {
                "question": "Câu hỏi mẫu?",
                "options": {"A": "Một", "B": "Hai", "C": "Ba", "D": "Bốn", "E": "Năm"},
                "correct_answer": "A, B, C, D, E",
            },
            "nhieu_lua_chon",
        )
        self.assertIn("không được chọn tất cả", error)

    def test_retry_prompt_allows_five_option_multi_answer_shape(self):
        prompt = _build_retry_prompt(
            original_prompt="ORIGINAL",
            question_type="nhieu_lua_chon",
            bloom_level="phan_tich",
            missing_count=3,
            validation_errors=["nhieu_lua_chon phải có 4 hoặc 5 lựa chọn"],
            avoid_questions=[],
        )
        self.assertIn('"A", "B", "C", "D", "E"', prompt)
        self.assertIn("Generate exactly 3 additional questions", prompt)

    def test_generation_preset_payload_limits_plan_rows(self):
        payload = GenerationPresetPayload(
            name="Ôn tập demo",
            planItems=[
                {"questionTypeId": "mcq", "bloomId": "understand", "count": 2},
            ],
            instruction="Tập trung vào cây nhị phân.",
        )
        self.assertEqual(payload.planItems[0].count, 2)

        with self.assertRaises(ValidationError):
            GenerationPresetPayload(
                name="Sai số lượng",
                planItems=[
                    {"questionTypeId": "mcq", "bloomId": "understand", "count": 11},
                ],
            )

    def test_user_service_caps_generation_presets_per_user(self):
        class FakeUserRepository:
            def __init__(self):
                self.user = {
                    "_id": ObjectId(),
                    "generation_presets": [],
                }

            def find_by_id(self, _user_id):
                return self.user

            def update(self, _user_id, fields):
                self.user.update(fields)
                return self.user

        repository = FakeUserRepository()
        service = UserService(repository=repository, identity=None, sessions=None)
        user_id = str(repository.user["_id"])

        for index in range(13):
            service.save_generation_preset(
                user_id,
                GenerationPresetPayload(
                    name=f"Preset {index}",
                    planItems=[
                        {"questionTypeId": "mcq", "bloomId": "understand", "count": 1},
                    ],
                ),
            )

        presets = service.list_generation_presets(user_id)["items"]
        self.assertEqual(len(presets), 12)
        self.assertEqual(presets[0]["name"], "Preset 12")
        self.assertNotIn("Preset 0", {preset["name"] for preset in presets})

        self.assertTrue(service.delete_generation_preset(user_id, presets[0]["id"]))
        self.assertEqual(len(service.list_generation_presets(user_id)["items"]), 11)

    def test_question_repository_list_pushes_reviewer_filters_to_mongo(self):
        class FakeQuestionsCollection:
            def __init__(self):
                self.pipeline = None

            def aggregate(self, pipeline):
                self.pipeline = pipeline
                return [{"items": [], "count": [{"total": 0}]}]

        class FakeDatabase:
            def __init__(self):
                self.questions = FakeQuestionsCollection()

        fake_db = FakeDatabase()
        repo = MongoQuestionRepository(fake_db)
        document_id = ObjectId()
        subject_id = ObjectId()
        chapter_id = ObjectId()
        clo_id = ObjectId()
        reviewer_id = ObjectId()
        creator_id = ObjectId()
        waiting_since = datetime.now(timezone.utc) - timedelta(hours=24)
        overdue_at = datetime.now(timezone.utc)

        pairs, total = repo.list(
            2,
            10,
            "PENDING",
            "queue",
            question_type="trac_nghiem",
            bloom_level=3,
            document_id=str(document_id),
            subject_id=str(subject_id),
            chapter_id=str(chapter_id),
            clo_id=str(clo_id),
            difficulty="kho",
            quality_color="green",
            min_score=0.8,
            publication_status="NOT_PUBLISHED",
            evaluation_status="PASSED",
            assignment_status="IN_REVIEW",
            assigned_reviewer_user_id=reviewer_id,
            creator_user_id=creator_id,
            waiting_since=waiting_since,
            overdue_at=overdue_at,
        )

        self.assertEqual(pairs, [])
        self.assertEqual(total, 0)
        pipeline = fake_db.questions.pipeline
        self.assertIsNotNone(pipeline)
        question_match = pipeline[0]["$match"]
        self.assertEqual(question_match["review_status"], "PENDING")
        self.assertEqual(question_match["publication_status"], "NOT_PUBLISHED")
        self.assertEqual(question_match["evaluation_status"], "PASSED")
        self.assertEqual(question_match["review_assignment.status"], "IN_REVIEW")
        self.assertEqual(question_match["review_assignment.reviewer_user_id"], reviewer_id)
        self.assertEqual(question_match["quality_summary.color"], "GREEN")
        self.assertEqual(question_match["quality_summary.overall_score"], {"$gte": 0.8})
        self.assertEqual(question_match["updated_at"], {"$lte": waiting_since})
        self.assertEqual(question_match["review_assignment.lock_expires_at"], {"$lte": overdue_at})

        creator_match = pipeline[3]["$match"]
        self.assertIn({"created_by_user_id": creator_id}, creator_match["$or"])
        self.assertIn({"version.created_by_user_id": creator_id}, creator_match["$or"])

        version_match = pipeline[4]["$match"]
        self.assertEqual(version_match["version.classification.assessment_type"], "TRAC_NGHIEM")
        self.assertEqual(version_match["version.classification.bloom.level"], 3)
        self.assertEqual(version_match["version.document_id"], document_id)
        self.assertEqual(version_match["version.classification.subject.id"], subject_id)
        self.assertEqual(version_match["version.classification.chapter.id"], chapter_id)
        self.assertEqual(version_match["version.clos.id"], clo_id)
        self.assertEqual(version_match["version.classification.difficulty"], "kho")

        search_match = pipeline[5]["$match"]
        self.assertIn({"question_code": {"$regex": "queue", "$options": "i"}}, search_match["$or"])
        self.assertIn({"version.content": {"$regex": "queue", "$options": "i"}}, search_match["$or"])
        facet = pipeline[-1]["$facet"]
        self.assertEqual(facet["items"][0], {"$skip": 10})
        self.assertEqual(facet["items"][1], {"$limit": 10})

    def test_question_repository_can_filter_approved_current_version_in_mongo(self):
        class FakeQuestionsCollection:
            def __init__(self):
                self.pipeline = None

            def aggregate(self, pipeline):
                self.pipeline = pipeline
                return [{"items": [], "count": [{"total": 0}]}]

        class FakeDatabase:
            def __init__(self):
                self.questions = FakeQuestionsCollection()

        fake_db = FakeDatabase()
        repo = MongoQuestionRepository(fake_db)

        repo.list(
            1,
            20,
            "APPROVED",
            None,
            approved_current_only=True,
        )

        self.assertIn(
            {"$match": {"$expr": {"$eq": ["$approved_version_id", "$current_version_id"]}}},
            fake_db.questions.pipeline,
        )


if __name__ == "__main__":
    unittest.main()
