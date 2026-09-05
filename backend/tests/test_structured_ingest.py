from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

from docx import Document

from modules.documents.ingest.models import ParseContext
from modules.documents.ingest.parsers.docx import DocxParser
from modules.documents.ingest.parsers.markdown import MarkdownParser
from modules.documents.ingest.parsers.pdf import PdfParser, _blocks_from_text, _has_table_signature
from modules.documents.ingest.parsers.text import TextParser
from modules.documents.ingest.quality import validate_chunks, validate_parsed_document
from modules.documents.ingest.registry import build_default_registry
from modules.documents.repository import compact_raw_extraction, json_safe
from modules.rag.chunking import _merge_small_structured_chunks, _split_protected_block, _structured_chunks
from modules.rag.search import _hybrid_score


def _context(name: str, document_type: str) -> ParseContext:
    return ParseContext(
        document_id="document-1",
        source_file_name=name,
        source_uri=f"uploads/{name}",
        mime_type="application/octet-stream",
        document_type=document_type,
    )


def test_registry_and_text_parser_keep_nonpaged_provenance(tmp_path: Path):
    source = tmp_path / "lesson.txt"
    source_text = "Dữ liệu tiếng Việt\nDòng thứ hai"
    source.write_text(source_text, encoding="utf-8")

    registry = build_default_registry(pdf_ocr_enabled=False)
    parsed = registry.resolve(source, "text/plain").parse(source, _context(source.name, "txt"))

    assert isinstance(registry.resolve(source), TextParser)
    assert parsed.units[0].page_number is None
    assert parsed.units[0].source_location == {"character_start": 0, "character_end": len(source_text)}
    assert parsed.units[0].raw_text == source_text
    assert parsed.units[0].content_blocks[0].provenance.source_file_name == source.name
    assert validate_parsed_document(parsed).passed


def test_markdown_parser_preserves_protected_blocks_and_image_reference(tmp_path: Path):
    source = tmp_path / "lesson.md"
    source.write_text(
        "# Tiêu đề\n\n"
        "| Cột A | Cột B |\n| --- | --- |\n| một | hai |\n\n"
        "```python\nif x:\n    print(x)\n```\n\n"
        "$$\na^2 + b^2 = c^2\n$$\n\n"
        "![Sơ đồ](diagram.png)\n",
        encoding="utf-8",
    )

    parsed = MarkdownParser().parse(source, _context(source.name, "md"))
    block_types = [block.block_type for block in parsed.units[0].content_blocks]

    assert block_types == ["heading", "table", "code", "formula", "image"]
    assert "    print(x)" in parsed.units[0].content_blocks[2].content
    assert parsed.units[0].content_blocks[1].structured_content["rows"][2] == ["một", "hai"]
    assert parsed.assets[0].storage_uri == "diagram.png"
    assert parsed.assets[0].status == "reference_only"
    assert validate_parsed_document(parsed).passed


def test_docx_parser_keeps_order_table_and_source_indexes(tmp_path: Path):
    source = tmp_path / "lesson.docx"
    document = Document()
    document.add_heading("Chương một", level=1)
    document.add_paragraph("Nội dung")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Tên"
    table.cell(0, 1).text = "Giá trị"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"
    document.save(source)

    parsed = DocxParser().parse(source, _context(source.name, "docx"))

    assert [block.block_type for block in parsed.units[0].content_blocks] == ["heading", "prose", "table"]
    assert parsed.units[0].page_number is None
    assert parsed.units[0].content_blocks[0].provenance.source_location["paragraph_index"] == 1
    assert parsed.units[0].content_blocks[2].structured_content["rows"] == [["Tên", "Giá trị"], ["A", "1"]]
    assert validate_parsed_document(parsed).passed


class _FakePage:
    def __init__(self, text: str, *, image: bool = False):
        self.text = text
        self.resources = (
            {"/XObject": {"/Im0": {"/Subtype": "/Image", "/Width": 200, "/Height": 100}}}
            if image
            else {}
        )

    def extract_text(self, extraction_mode="plain", space_width=None):
        return self.text

    def get(self, name, default=None):
        if name == "/Resources":
            return self.resources
        return default


class _FakeReader:
    def __init__(self, pages):
        self.pages = pages
        self.is_encrypted = False


def test_pdf_parser_routes_only_risky_pages_to_docling_and_keeps_good_text(tmp_path: Path):
    source = tmp_path / "mixed.pdf"
    source.write_bytes(b"fake")
    pages = [
        _FakePage("Đây là trang văn bản đủ dài để dùng trực tiếp từ text layer. " * 4),
        _FakePage(""),
        _FakePage("Trang có sơ đồ nhưng text layer vẫn đầy đủ. " * 5, image=True),
    ]
    calls = []

    def extract(_path, page_numbers, _context):
        calls.append(page_numbers)
        return {
            2: {
                "text": "Nội dung OCR của trang scan " * 8,
                "structured_blocks": [
                    {
                        "block_type": "table",
                        "content": "A | B\n1 | 2",
                        "bbox": [10, 20, 300, 180],
                        "structured_content": {"rows": [["A", "B"], ["1", "2"]]},
                    }
                ],
            },
            3: {
                "text": "Docling layout",
                "structured_blocks": [
                    {"block_type": "caption", "content": "Hình 1: Sơ đồ", "structured_content": None}
                ],
            },
        }

    with patch("modules.documents.ingest.parsers.pdf.PdfReader", return_value=_FakeReader(pages)):
        parsed = PdfParser(ocr_page_extractor=extract).parse(source, _context(source.name, "pdf"))

    assert calls == [[2, 3]]
    assert parsed.units[0].quality["selected_method"] == "pypdf_plain"
    assert parsed.units[1].quality["selected_method"] == "docling_page_selective_ocr"
    hard_ocr_table = next(block for block in parsed.units[1].content_blocks if block.block_type == "table")
    assert hard_ocr_table.structured_content["rows"] == [["A", "B"], ["1", "2"]]
    assert hard_ocr_table.provenance.bbox == [10.0, 20.0, 300.0, 180.0]
    assert parsed.units[2].raw_text.startswith("Trang có sơ đồ")
    assert parsed.units[2].content_blocks[-1].block_type == "caption"
    assert parsed.units[2].asset_ids
    assert parsed.assets[0].asset_type == "diagram"
    assert parsed.assets[0].source_caption
    assert parsed.stats == {
        "source_format": "pdf",
        "page_count": 3,
        "asset_count": 1,
        "ocr_page_count": 1,
        "layout_page_count": 1,
        "docling_page_count": 2,
        "text_layer_page_count": 2,
    }
    assert validate_parsed_document(parsed).passed


def test_pdf_parser_accepts_short_clean_text_page_without_unnecessary_ocr(tmp_path: Path):
    source = tmp_path / "short-clean.pdf"
    source.write_bytes(b"fake")
    calls = []

    def extract(_path, page_numbers, _context):
        calls.append(page_numbers)
        return {}

    with patch(
        "modules.documents.ingest.parsers.pdf.PdfReader",
        return_value=_FakeReader([_FakePage("Tiếng Việt có dấu và nội dung ngắn nhưng hợp lệ.")]),
    ):
        parsed = PdfParser(ocr_page_extractor=extract).parse(source, _context(source.name, "pdf"))

    assert calls == []
    assert parsed.units[0].quality["requires_ocr"] is False
    assert validate_parsed_document(parsed).passed


def test_pdf_table_signature_recognizes_repeated_two_column_layout_rows():
    layout = "Thao tác                    Thời gian\n\nPush                        O(1)\n\nPop                         O(1)"
    assert _has_table_signature(layout)


def test_pdf_aligned_table_header_is_not_misclassified_as_heading():
    blocks = _blocks_from_text(
        "COLUMN A                    VALUE\nroot                        10",
        _context("table.pdf", "pdf"),
        page_number=1,
        extractor="pypdf",
        extraction_method="layout",
        confidence=1.0,
    )

    assert [block.block_type for block in blocks] == ["table"]
    assert blocks[0].structured_content["rows"] == [["COLUMN A", "VALUE"], ["root", "10"]]


def test_protected_split_is_lossless_and_uses_explicit_continuation():
    content = "line one has enough source characters 01\nline two has enough source characters 02\n"
    content += "line three has enough source characters 03\nline four has enough source characters 04\n"
    parts = _split_protected_block(
        {"block_id": "code-parent", "block_type": "code", "content": content},
        chunk_size=100,
        max_code_lines=2,
    )

    assert "".join(part["content"] for part in parts) == content
    assert [part["part_index"] for part in parts] == [1, 2]
    assert all(part["continuation_of"] == "code-parent" for part in parts)
    assert not any("continued" in part["content"].lower() for part in parts)


def test_tiny_fragments_merge_only_with_same_page_context():
    base_source = {"source_file_name": "lesson.pdf", "source_uri": "uploads/lesson.pdf", "document_id": "d"}
    chunks = [
        {"chunk_id": "1", "content": "Giải thích", "metadata": {**base_source, "page_marks": [4], "content_type": "text", "block_ids": ["a"]}},
        {"chunk_id": "2", "content": "{", "metadata": {**base_source, "page_marks": [4], "content_type": "code", "block_ids": ["b"]}},
        {"chunk_id": "3", "content": "Trang sau", "metadata": {**base_source, "page_marks": [5], "content_type": "text", "block_ids": ["c"]}},
    ]

    with patch("modules.rag.chunking.embedding_token_lengths", side_effect=lambda values: [len(v.split()) for v in values]):
        merged = _merge_small_structured_chunks(chunks, "d", 500, [])

    assert [chunk["content"] for chunk in merged] == ["Giải thích\n\n{", "Trang sau"]
    assert merged[0]["metadata"]["content_type"] == "mixed"
    assert merged[0]["metadata"]["block_ids"] == ["a", "b"]


def test_structured_chunks_have_source_and_exact_page_provenance():
    document = {
        "_id": "document-1",
        "title": "Bài học",
        "original_filename": "lesson.pdf",
        "current_version": 2,
        "artifacts": [
            {
                "_id": "artifact-1",
                "type": "ORIGINAL_PDF",
                "is_current": True,
                "storage": {"uri": "uploads/lesson.pdf"},
            }
        ],
    }
    pages = [
        {
            "page_number": 7,
            "content_blocks": [
                {
                    "block_id": "heading-1",
                    "block_type": "heading",
                    "content": "Cấu trúc hàng đợi",
                    "provenance": {"page_number": 7, "source_location": {"page_number": 7, "line_start": 1}},
                },
                {
                    "block_id": "prose-1",
                    "block_type": "prose",
                    "content": "FIFO lấy phần tử vào trước ra trước.",
                    "provenance": {"page_number": 7, "source_location": {"page_number": 7, "line_start": 2}},
                },
            ],
        }
    ]

    with patch("modules.rag.chunking.embedding_token_lengths", side_effect=lambda values: [len(v.split()) for v in values]):
        chunks = _structured_chunks(document, pages, "document-1", 500, 50, 20, [], 512)

    assert len(chunks) == 1
    metadata = chunks[0]["metadata"]
    assert metadata["page_marks"] == [7]
    assert metadata["source_file_name"] == "lesson.pdf"
    assert metadata["source_uri"] == "uploads/lesson.pdf"
    assert metadata["block_ids"] == ["heading-1", "prose-1"]
    assert metadata["document_version"] == 2
    assert validate_chunks(chunks).passed


def test_chunk_quality_gate_rejects_missing_provenance():
    report = validate_chunks([{"chunk_id": "bad", "content": "text", "metadata": {"document_id": "d"}}])

    assert not report.passed
    assert any("source_file_name" in error for error in report.errors)
    assert any("source location" in error for error in report.errors)


def test_mongo_projection_keeps_raw_audit_bson_safe_without_duplicate_text():
    raw = {
        "candidates": [{"method": "plain", "text": "source text", "engine_id": 2**80}],
        "docling": {
            "text": "duplicate text",
            "raw_document": {"large": True},
            "structured_blocks": [{"block_type": "table"}],
            "diagnostics": {"request_id": 2**80},
        },
    }

    compact = compact_raw_extraction(raw)

    assert compact["candidates"][0]["text_length"] == len("source text")
    assert "text" not in compact["candidates"][0]
    assert compact["candidates"][0]["engine_id"] == str(2**80)
    assert compact["docling"]["structured_block_count"] == 1
    assert "raw_document" not in compact["docling"]
    assert json_safe({"value": 2**80}) == {"value": str(2**80)}


def test_hybrid_ranking_allows_exact_identifier_to_beat_irrelevant_vector_top_one():
    irrelevant = _hybrid_score("cây nhị phân", {}, {"lifo"}, 0)
    exact = _hybrid_score("Ngăn xếp tuân theo LIFO", {}, {"lifo"}, 20)

    assert exact > irrelevant
