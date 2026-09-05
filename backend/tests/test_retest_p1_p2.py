from __future__ import annotations

import contextlib
import gzip
import hashlib
import json
import shutil
import subprocess
from datetime import datetime, timedelta, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock

import pytest
from bson import ObjectId
from docx import Document
from PIL import Image

from modules.documents.ingest.base import DocumentConversionError, UnsupportedDocumentError
from modules.documents.ingest.benchmark import (
    character_error_rate,
    evaluate_golden_case,
    evaluate_retrieval_cases,
    load_golden_case,
    word_error_rate,
)
from modules.documents.ingest.models import ContentBlock, ParseContext, SourceProvenance
from modules.documents.ingest.parsers.docx import LegacyDocParser
from modules.documents.ingest.quality import validate_parsed_document
from modules.documents.retention import (
    RetentionPolicy,
    build_retention_plan,
    deduplicate_artifact_file,
    protected_artifact_ids,
)
from modules.ocr import mongodb as ocr_mongodb
from modules.ocr.pipeline import run_document_pipeline
from modules.rag import lineage
from modules.rag.lineage import CandidateLineage, LineagePromotionService, LineageValidator
from modules.rag.mongodb import _completion_pointer_update
from modules.rag.search import _requires_lexical_fallback
from scripts.reprocess_document_lineage import (
    ReprocessRequest,
    execute_reprocess,
    inspect_reprocess_plan,
    parse_args as parse_reprocess_args,
)


def _context() -> ParseContext:
    return ParseContext(
        document_id="document-1",
        source_file_name="legacy.doc",
        source_uri="golden://legacy.doc",
        mime_type="application/msword",
        document_type="doc",
    )


def test_golden_corpus_has_independent_versioned_truth_and_twelve_pages():
    root = Path(__file__).parent / "golden_corpus" / "v1"
    manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))
    truth = json.loads((root / manifest["ground_truth"]).read_text(encoding="utf-8"))

    assert manifest["corpus_version"] == truth["corpus_version"]
    assert len(truth["pages"]) == 12
    assert {case["profile"] for case in manifest["cases"]} >= {
        "born_digital", "standard_scan", "degraded_scan", "structured_document"
    }
    case, pages = load_golden_case(root, "mixed-pdf-vi")
    assert case["pages"] == [7, 8]
    assert [page["page_number"] for page in pages] == [1, 2]


def test_golden_metrics_enforce_cer_page_and_protected_structure_thresholds():
    expected = [{
        "page_number": 1,
        "text": "Dữ liệu tiếng Việt",
        "code_lines": ["    print(x)"],
        "tables": [{"rows": [["A", "B"]]}],
        "formulae": ["a+b=c"],
    }]
    actual = [{
        "page_number": 1,
        "text": "Dữ liệu tiếng Việt",
        "content_blocks": [
            {"block_type": "code", "content": "    print(x)"},
            {"block_type": "table", "structured_content": {"rows": [["A", "B"]]}},
            {"block_type": "formula", "content": "a+b=c", "structured_content": {"raw": "a+b=c"}},
        ],
    }]
    report = evaluate_golden_case(expected, actual, "born_digital")

    assert character_error_rate("đúng", "đung") == pytest.approx(0.25)
    assert word_error_rate("một hai", "một ba") == pytest.approx(0.5)
    assert report["status"] == "passed"
    assert report["metrics"]["table_cell_accuracy"] == 1.0
    assert evaluate_golden_case(expected, [], "born_digital")["status"] == "failed"


def test_golden_text_metric_ignores_only_layout_whitespace_and_table_delimiters():
    expected = [{
        "page_number": 1,
        "text": "| Thao tác | Thời gian |\n| --- | --- |\n| Push | O(1) |",
        "tables": [{"rows": [["Thao tác", "Thời gian"], ["Push", "O(1)"]]}],
    }]
    actual = [{
        "page_number": 1,
        "text": "Thao tác                  Thời gian\nPush                      O(1)",
        "content_blocks": [
            {"block_type": "table", "structured_content": {"rows": [["Thao tác", "Thời gian"], ["Push", "O(1)"]]}}
        ],
    }]

    report = evaluate_golden_case(expected, actual, "born_digital")
    assert report["metrics"]["cer"] == 0.0
    assert report["status"] == "passed"


def test_golden_asset_recall_is_typed_and_never_exceeds_one():
    expected = [{"page_number": 1, "text": "Sơ đồ", "assets": [{"type": "diagram"}]}]
    actual = [{
        "page_number": 1,
        "text": "Sơ đồ",
        "content_blocks": [{"block_type": "prose", "content": "Sơ đồ"}],
        "assets": [
            {"asset_type": "image", "content_sha256": "a", "storage_uri": "golden://a", "provenance": {"page_number": 1}},
            {"asset_type": "image", "content_sha256": "b", "storage_uri": "golden://b", "provenance": {"page_number": 1}},
        ],
    }]

    report = evaluate_golden_case(expected, actual, "born_digital")
    assert report["metrics"]["asset_recall"] == 0.0
    assert report["metrics"]["asset_type_precision"] == 0.0


def test_golden_metrics_include_source_captions_and_ignore_unlabeled_page_rasters():
    expected = [{
        "page_number": 1,
        "text": "Edge u connects to v.",
        "assets": [{"type": "diagram", "caption": "Graph diagram"}],
    }]
    actual = [{
        "page_number": 1,
        "text": "Edge u connects to v.\n\nGraph diagram",
        "content_blocks": [{"block_type": "caption", "content": "Graph diagram"}],
        "assets": [
            {
                "asset_type": "image",
                "metadata": {"is_page_raster": True},
                "content_sha256": "page",
                "storage_uri": "golden://page",
                "provenance": {"page_number": 1},
            },
            {
                "asset_type": "diagram",
                "source_caption": "Graph diagram",
                "content_sha256": "diagram",
                "storage_uri": "golden://diagram",
                "provenance": {"page_number": 1},
            },
        ],
    }]

    report = evaluate_golden_case(expected, actual, "standard_scan")
    assert report["metrics"]["cer"] == 0.0
    assert report["metrics"]["asset_recall"] == 1.0
    assert report["metrics"]["asset_type_precision"] == 1.0
    assert report["status"] == "passed"


def test_formula_metric_treats_trailing_sentence_punctuation_as_non_semantic():
    expected = [{"page_number": 1, "text": "T(n) = n", "formulae": ["T(n) = n"]}]
    actual = [{
        "page_number": 1,
        "text": "T(n) = n",
        "content_blocks": [
            {"block_type": "formula", "content": "T(n) = n.", "structured_content": {"raw": "T(n) = n."}}
        ],
    }]

    report = evaluate_golden_case(expected, actual, "degraded_scan")
    assert report["metrics"]["formula_exact_accuracy"] == 1.0
    assert report["metrics"]["formula_preserved_or_needs_review"] == 1.0


def test_retrieval_truth_measures_hit_mrr_provenance_and_abstention():
    cases = [
        {"id": "answerable", "answerable": True, "relevant_truth_pages": [4]},
        {"id": "unsupported", "answerable": False, "relevant_truth_pages": [], "required_abstention": True},
    ]
    results = {
        "answerable": {
            "results": [
                {"chunk_id": "wrong", "source_uri": "golden://x", "page_marks": [1]},
                {"chunk_id": "right", "source_uri": "golden://x", "page_marks": [4]},
            ]
        },
        "unsupported": {"results": [], "abstained": True},
    }
    report = evaluate_retrieval_cases(cases, results)
    assert report["hit_at_5"] == 1.0
    assert report["mrr"] == 0.5
    assert report["provenance_completeness"] == 1.0
    assert report["abstention_accuracy"] == 1.0


def test_legacy_doc_fails_closed_with_structured_error_when_runtime_missing(tmp_path, monkeypatch):
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"not-indexable")
    monkeypatch.setattr(shutil, "which", lambda _name: None)

    with pytest.raises(UnsupportedDocumentError) as exc_info:
        LegacyDocParser().parse(source, _context())

    assert exc_info.value.to_dict()["code"] == "UNSUPPORTED_DOCUMENT"


def test_legacy_doc_fails_closed_when_converter_emits_corrupt_docx(tmp_path, monkeypatch):
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    monkeypatch.setattr(shutil, "which", lambda _name: "soffice")

    def fake_run(command, **_kwargs):
        output_dir = Path(command[command.index("--outdir") + 1])
        (output_dir / "legacy.docx").write_bytes(b"corrupt")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(subprocess, "run", fake_run)
    with pytest.raises(DocumentConversionError) as exc_info:
        LegacyDocParser().parse(source, _context())
    assert exc_info.value.to_dict()["code"] == "DOCUMENT_CONVERSION_FAILED"


@pytest.mark.skipif(not (shutil.which("soffice") or shutil.which("libreoffice")), reason="LibreOffice unavailable")
def test_legacy_doc_runtime_preserves_vietnamese_and_table(tmp_path):
    source_docx = tmp_path / "legacy.docx"
    document = Document()
    document.add_paragraph("Dữ liệu tiếng Việt có dấu")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "Tên", "Giá trị"
    table.cell(1, 0).text, table.cell(1, 1).text = "Hàng đợi", "FIFO"
    image_path = tmp_path / "source-image.png"
    Image.new("RGB", (64, 32), "blue").save(image_path)
    document.add_picture(str(image_path))
    document.save(source_docx)
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    completed = subprocess.run(
        [converter, "--headless", "--convert-to", "doc:MS Word 97", "--outdir", str(tmp_path), str(source_docx)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    source_doc = tmp_path / "legacy.doc"
    assert completed.returncode == 0 and source_doc.exists()

    parsed = LegacyDocParser().parse(source_doc, _context())
    assert "Dữ liệu tiếng Việt có dấu" in parsed.units[0].rendered_text()
    assert next(block for block in parsed.units[0].content_blocks if block.block_type == "table").structured_content["rows"][1] == ["Hàng đợi", "FIFO"]
    assert parsed.assets and parsed.assets[0].content_sha256


@pytest.mark.skipif(not (shutil.which("soffice") or shutil.which("libreoffice")), reason="LibreOffice unavailable")
def test_corrupt_legacy_doc_runtime_fails_closed(tmp_path):
    source = tmp_path / "corrupt.doc"
    source.write_bytes(b"not a compound document")
    with pytest.raises(DocumentConversionError):
        LegacyDocParser().parse(source, _context())


def test_typed_validation_rejects_unstructured_table_and_invented_visual():
    provenance = SourceProvenance(
        source_file_name="a.pdf", source_uri="golden://a.pdf", document_id="d", document_type="pdf",
        page_number=1, source_location={"page_number": 1}, extractor="test", extraction_method="test"
    )
    table = ContentBlock(block_id="table", block_type="table", content="A B", provenance=provenance)
    visual = ContentBlock(block_id="image", block_type="image", content="invented description", provenance=provenance)
    document = MagicMock(units=[SimpleNamespace(quality={}, content_blocks=[table, visual])], assets=[])
    report = validate_parsed_document(document)
    assert not report.passed
    assert any("table cell grid missing" in error for error in report.errors)
    assert any("visual block missing source asset" in error for error in report.errors)


def test_retention_never_archives_active_pending_or_rollback_artifacts():
    now = datetime.now(timezone.utc)
    artifacts = [
        {"_id": "active", "job_id": "ocr-active", "is_current": True, "created_at": now - timedelta(days=500), "sha256": "same"},
        {"_id": "pending", "job_id": "ocr-pending", "created_at": now - timedelta(days=500), "sha256": "same"},
        {"_id": "cold", "job_id": "old", "created_at": now - timedelta(days=60), "sha256": "other"},
        {"_id": "expired", "job_id": "older", "created_at": now - timedelta(days=500), "sha256": "third"},
    ]
    document = {
        "artifacts": artifacts,
        "current_processing": {"ocr_job_id": "ocr-active"},
        "pending_processing": {"ocr_job_id": "ocr-pending"},
    }
    protected = protected_artifact_ids(document, [{"rollback_snapshot": {"artifact_ids": ["active"]}}])
    plan = build_retention_plan(artifacts, protected_ids=protected, policy=RetentionPolicy(30, 365), now=now)
    assert set(plan["keep_hot"]) == {"active", "pending"}
    assert plan["archive_cold"] == ["cold"]
    assert plan["delete_candidates"] == [{"artifact_id": "expired", "requires_explicit_confirmation": True}]
    assert plan["dedup_groups"][0]["duplicate_artifact_ids"] == ["pending"]


def test_artifact_compression_and_content_hash_dedup_are_lossless(tmp_path, monkeypatch):
    source = tmp_path / "source.txt"
    source.write_text("Dữ liệu tiếng Việt " * 100, encoding="utf-8")
    monkeypatch.setattr("modules.ocr.pipeline.settings.raw_artifact_compression", "gzip")
    result = run_document_pipeline(
        str(source),
        str(tmp_path / "result.md"),
        document_id="golden-dedup",
        source_uri="golden://source.txt",
    )
    raw_path = Path(result["raw_extraction_file"])
    with gzip.open(raw_path, "rt", encoding="utf-8") as source_stream:
        raw_payload = json.load(source_stream)
    first = deduplicate_artifact_file(raw_path, tmp_path / "blobs")
    duplicate = tmp_path / "duplicate.raw.json.gz"
    duplicate.write_bytes(Path(first["uri"]).read_bytes())
    second = deduplicate_artifact_file(duplicate, tmp_path / "blobs")

    assert raw_payload["document_id"] == "golden-dedup"
    assert result["stats"]["storage"]["raw_compression_ratio"] < 1
    assert second["reused"] is True
    assert second["uri"] == first["uri"]
    assert not duplicate.exists()


def test_replacement_ocr_completion_preserves_active_pipeline_summary(monkeypatch):
    document_id = ObjectId()
    active_chunk_set_id = ObjectId()
    document = {
        "_id": document_id,
        "status": "READY",
        "current_processing": {"chunk_set_id": active_chunk_set_id},
        "pipeline_summary": {
            "ocr_status": "COMPLETED",
            "chunk_status": "COMPLETED",
            "index_status": "COMPLETED",
            "total_chunks": 284,
        },
    }
    repository = MagicMock()
    repository.find_by_id.return_value = document
    database = SimpleNamespace(documents=MagicMock())
    monkeypatch.setattr(ocr_mongodb, "_repository", lambda: repository)
    monkeypatch.setattr(ocr_mongodb, "get_database", lambda: database)

    ocr_mongodb.update_document_status(
        str(document_id),
        str(ObjectId()),
        status="completed",
        stats={"total_pages": 151},
    )

    repository.update_job.assert_called_once()
    database.documents.update_one.assert_not_called()


def test_first_ocr_completion_still_marks_chunk_and_index_as_not_started(monkeypatch):
    document_id = ObjectId()
    document = {
        "_id": document_id,
        "status": "PROCESSING",
        "current_processing": {"ocr_job_id": ObjectId()},
        "pipeline_summary": {"ocr_status": "COMPLETED"},
    }
    repository = MagicMock()
    repository.find_by_id.return_value = document
    database = SimpleNamespace(documents=MagicMock())
    monkeypatch.setattr(ocr_mongodb, "_repository", lambda: repository)
    monkeypatch.setattr(ocr_mongodb, "get_database", lambda: database)

    ocr_mongodb.update_document_status(
        str(document_id),
        str(ObjectId()),
        status="completed",
    )

    update = database.documents.update_one.call_args.args[1]["$set"]
    assert update == {
        "status": "PROCESSING",
        "pipeline_summary.chunk_status": "NOT_STARTED",
        "pipeline_summary.index_status": "NOT_STARTED",
    }


def _reprocess_guard_fixture(tmp_path):
    source = tmp_path / "source.pdf"
    source.write_bytes(b"same-document-source")
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    document_id, ocr_id, chunk_id, vector_id = (ObjectId() for _ in range(4))
    document = {
        "_id": document_id,
        "status": "READY",
        "current_version": 1,
        "current_processing": {
            "ocr_job_id": ocr_id,
            "chunk_set_id": chunk_id,
            "vector_collection_id": vector_id,
        },
        "pending_processing": {},
        "pipeline_summary": {
            "ocr_status": "COMPLETED",
            "chunk_status": "COMPLETED",
            "index_status": "COMPLETED",
        },
        "artifacts": [
            {
                "type": "ORIGINAL_PDF",
                "is_current": True,
                "sha256": digest,
                "mime_type": "application/pdf",
                "storage": {"uri": str(source)},
            }
        ],
        "archived_at": None,
    }
    database = SimpleNamespace(
        document_jobs=MagicMock(),
        chunk_sets=MagicMock(),
        vector_collections=MagicMock(),
    )
    database.document_jobs.find.return_value = []
    database.document_jobs.find_one.return_value = {
        "_id": ocr_id,
        "document_id": document_id,
        "document_version": 1,
        "job_type": "OCR",
        "status": "COMPLETED",
    }
    database.chunk_sets.find_one.return_value = {
        "_id": chunk_id,
        "document_id": document_id,
        "document_version": 1,
        "source_ocr_job_id": ocr_id,
        "status": "COMPLETED",
    }

    def find_vector(query):
        if query.get("_id") == vector_id:
            return {
                "_id": vector_id,
                "provider": "CHROMA",
                "is_active": True,
                "collection_name": "active",
            }
        return None

    database.vector_collections.find_one.side_effect = find_vector
    repository = MagicMock()
    repository.find_by_id.return_value = document
    request = ReprocessRequest(
        document_id=str(document_id),
        collection_name="candidate-r6",
        expected_source_sha256=digest,
        expected_active_ocr_job_id=str(ocr_id),
        expected_active_chunk_set_id=str(chunk_id),
        expected_active_vector_collection_id=str(vector_id),
    )
    return request, document, database, repository


def test_reprocess_script_defaults_to_read_only_and_validates_complete_plan(tmp_path):
    parsed = parse_reprocess_args(
        [
            "--document-id", str(ObjectId()),
            "--collection-name", "candidate-r6",
            "--expected-source-sha256", "a" * 64,
            "--expected-active-ocr-job-id", str(ObjectId()),
            "--expected-active-chunk-set-id", str(ObjectId()),
            "--expected-active-vector-collection-id", str(ObjectId()),
        ]
    )
    assert parsed.apply is False
    with pytest.raises(ValueError, match="requires --apply"):
        execute_reprocess(parsed)

    request, _document, database, repository = _reprocess_guard_fixture(tmp_path)
    plan = inspect_reprocess_plan(
        request,
        db=database,
        repository=repository,
        chroma_collection_exists=lambda _name: False,
    )
    assert plan.public_dict()["mode"] == "read_only"
    assert plan.public_dict()["promotion_allowed"] is False
    assert plan.source_sha256 == request.expected_source_sha256


@pytest.mark.parametrize(
    "guard",
    ["pending", "active_job", "bad_summary", "mongo_collision", "chroma_collision"],
)
def test_reprocess_script_rejects_unsafe_state(tmp_path, guard):
    request, document, database, repository = _reprocess_guard_fixture(tmp_path)
    chroma_exists = lambda _name: False
    if guard == "pending":
        document["pending_processing"] = {"ocr_job_id": ObjectId()}
    elif guard == "active_job":
        database.document_jobs.find.return_value = [
            {"_id": ObjectId(), "job_type": "OCR", "status": "PROCESSING"}
        ]
    elif guard == "bad_summary":
        document["pipeline_summary"]["index_status"] = "NOT_STARTED"
    elif guard == "mongo_collision":
        original_find_vector = database.vector_collections.find_one.side_effect

        def find_vector(query):
            return {"_id": ObjectId(), "collection_name": request.collection_name} if "collection_name" in query else original_find_vector(query)

        database.vector_collections.find_one.side_effect = find_vector
    else:
        chroma_exists = lambda _name: True

    with pytest.raises(ValueError):
        inspect_reprocess_plan(
            request,
            db=database,
            repository=repository,
            chroma_collection_exists=chroma_exists,
        )


def test_reprocess_script_rejects_source_or_active_snapshot_mismatch(tmp_path):
    request, document, database, repository = _reprocess_guard_fixture(tmp_path)
    wrong_source = ReprocessRequest(
        **{
            **request.__dict__,
            "expected_source_sha256": "f" * 64,
        }
    )
    with pytest.raises(ValueError, match="source SHA-256"):
        inspect_reprocess_plan(
            wrong_source,
            db=database,
            repository=repository,
            chroma_collection_exists=lambda _name: False,
        )

    document["current_processing"]["chunk_set_id"] = ObjectId()
    with pytest.raises(ValueError, match="active lineage mismatch"):
        inspect_reprocess_plan(
            request,
            db=database,
            repository=repository,
            chroma_collection_exists=lambda _name: False,
        )


def test_lexical_fallback_only_for_identifier_or_low_vector_confidence(monkeypatch):
    monkeypatch.setattr("modules.rag.search.settings.lexical_fallback_distance_threshold", 0.55)
    assert _requires_lexical_fallback("Giải thích FIFO", [0.1])
    assert _requires_lexical_fallback("giải thích khái niệm", [0.7])
    assert not _requires_lexical_fallback("giải thích khái niệm", [0.2])


def test_reprocessing_completed_lineage_stays_pending_until_formal_promotion():
    now = datetime.now(timezone.utc)
    old_chunk, new_ocr, new_chunk, new_vector, chunk_job = (ObjectId() for _ in range(5))
    fields, unsets = _completion_pointer_update(
        {"current_processing": {"chunk_set_id": old_chunk}},
        source_ocr_job_id=new_ocr,
        chunk_set_id=new_chunk,
        vector_collection_id=new_vector,
        chunk_job_id=chunk_job,
        total_chunks=454,
        now=now,
    )
    assert fields["pending_processing.chunk_set_id"] == new_chunk
    assert fields["pending_processing.validation_status"] == "AWAITING_VALIDATION"
    assert "current_processing.chunk_set_id" not in fields
    assert unsets == {}


def test_first_completed_lineage_can_activate_without_replacement_promotion():
    now = datetime.now(timezone.utc)
    new_ocr, new_chunk, new_vector, chunk_job = (ObjectId() for _ in range(4))
    fields, unsets = _completion_pointer_update(
        {"current_processing": {"chunk_set_id": None}},
        source_ocr_job_id=new_ocr,
        chunk_set_id=new_chunk,
        vector_collection_id=new_vector,
        chunk_job_id=chunk_job,
        total_chunks=10,
        now=now,
    )
    assert fields["current_processing.chunk_set_id"] == new_chunk
    assert fields["pipeline_summary.total_chunks"] == 10
    assert unsets == {"pending_processing": ""}


class _Validator:
    def __init__(self, status="passed"):
        self.status = status

    def validate(self, _candidate, *, smoke_queries):
        return {"status": self.status, "errors": [] if self.status == "passed" else ["failed"], "warnings": [], "metrics": {}}


class _Documents:
    def __init__(self, document, modified=1):
        self.document = document
        self.modified = modified
        self.updates = []

    def find_one(self, _query):
        return self.document

    def update_one(self, query, update, session=None):
        self.updates.append((query, update, session))
        return SimpleNamespace(modified_count=self.modified)


class _Events:
    def __init__(self, event=None, fail_insert=False):
        self.event = event
        self.fail_insert = fail_insert
        self.inserted = []

    def insert_one(self, value, session=None):
        if self.fail_insert:
            raise RuntimeError("audit write failed")
        self.inserted.append(value)

    def find_one(self, _query):
        return self.event

    def update_one(self, *_args, **_kwargs):
        return SimpleNamespace(modified_count=1)


def _promotion_fixture(fail_insert=False):
    document_id, ocr_id, chunk_id, vector_id = ObjectId(), ObjectId(), ObjectId(), ObjectId()
    before = {"ocr_job_id": ObjectId(), "chunk_set_id": ObjectId(), "vector_collection_id": ObjectId()}
    pending = {"ocr_job_id": ocr_id, "chunk_set_id": chunk_id, "vector_collection_id": vector_id}
    document = {
        "_id": document_id,
        "archived_at": None,
        "current_processing": before,
        "pending_processing": pending,
    }
    db = SimpleNamespace(
        documents=_Documents(document),
        document_jobs=MagicMock(find_one=MagicMock(return_value={"status": "COMPLETED"})),
        chunk_sets=MagicMock(find_one=MagicMock(return_value={"status": "COMPLETED"})),
        vector_collections=MagicMock(find_one=MagicMock(return_value={"is_active": True})),
        pipeline_lineage_events=_Events(fail_insert=fail_insert),
    )
    candidate = CandidateLineage(str(document_id), str(ocr_id), str(chunk_id), str(vector_id))
    return db, candidate, before


def test_promotion_validation_failure_does_not_change_pointer():
    db, candidate, _before = _promotion_fixture()
    service = LineagePromotionService(db, _Validator("failed"))
    with pytest.raises(ValueError):
        service.promote(candidate, smoke_queries=["FIFO"], actor="admin", reason="test", confirmation=service.confirmation_token(candidate))
    assert db.documents.updates == []


def test_promotion_mongo_compare_and_set_failure_does_not_write_audit(monkeypatch):
    db, candidate, _before = _promotion_fixture()
    db.documents.modified = 0
    service = LineagePromotionService(db, _Validator())
    monkeypatch.setattr(lineage, "mongo_transaction", lambda: contextlib.nullcontext(MagicMock()))
    with pytest.raises(RuntimeError, match="changed concurrently"):
        service.promote(
            candidate,
            smoke_queries=["FIFO"],
            actor="admin",
            reason="validated",
            confirmation=service.confirmation_token(candidate),
        )
    assert db.pipeline_lineage_events.inserted == []


def test_promotion_success_keeps_old_snapshot_for_rollback(monkeypatch):
    db, candidate, before = _promotion_fixture()
    service = LineagePromotionService(db, _Validator())
    monkeypatch.setattr(lineage, "mongo_transaction", lambda: contextlib.nullcontext(MagicMock()))
    result = service.promote(
        candidate, smoke_queries=["FIFO"], actor="admin", reason="validated", confirmation=service.confirmation_token(candidate)
    )
    assert result["from_snapshot"] == before
    assert db.pipeline_lineage_events.inserted[0]["rollback_available"] is True


def test_promotion_compensates_pointer_when_audit_fails_without_transaction(monkeypatch):
    db, candidate, _before = _promotion_fixture(fail_insert=True)
    service = LineagePromotionService(db, _Validator())
    monkeypatch.setattr(lineage, "mongo_transaction", lambda: contextlib.nullcontext(None))
    with pytest.raises(RuntimeError, match="audit write failed"):
        service.promote(
            candidate, smoke_queries=["FIFO"], actor="admin", reason="validated", confirmation=service.confirmation_token(candidate)
        )
    assert len(db.documents.updates) == 2


def test_rollback_requires_promoted_lineage_to_still_be_active(monkeypatch):
    db, candidate, before = _promotion_fixture()
    after = {key: value for key, value in candidate.as_object_ids().items() if key != "document_id"}
    db.pipeline_lineage_events.event = {
        "document_id": candidate.as_object_ids()["document_id"],
        "from_snapshot": before,
        "to_snapshot": after,
    }
    monkeypatch.setattr(lineage, "mongo_transaction", lambda: contextlib.nullcontext(MagicMock()))
    result = LineagePromotionService(db, _Validator()).rollback("promotion-1", actor="admin", reason="regression")
    assert result["restored_snapshot"] == before


def test_lineage_validator_fails_fast_for_cross_document_candidate():
    document_id, other_document_id = ObjectId(), ObjectId()
    ocr_id, chunk_set_id, vector_id = ObjectId(), ObjectId(), ObjectId()
    db = SimpleNamespace(
        documents=MagicMock(find_one=MagicMock(return_value={"_id": document_id})),
        document_jobs=MagicMock(find_one=MagicMock(return_value={
            "_id": ocr_id, "document_id": other_document_id, "job_type": "OCR", "status": "COMPLETED"
        })),
        chunk_sets=MagicMock(find_one=MagicMock(return_value={
            "_id": chunk_set_id, "document_id": other_document_id, "source_ocr_job_id": ocr_id,
            "status": "COMPLETED", "total_chunks": 1,
        })),
        vector_collections=MagicMock(find_one=MagicMock(return_value={
            "_id": vector_id, "collection_name": "candidate", "is_active": True
        })),
        document_pages=MagicMock(),
    )
    candidate = CandidateLineage(str(document_id), str(ocr_id), str(chunk_set_id), str(vector_id))

    report = LineageValidator(db).validate(candidate, smoke_queries=["FIFO"])

    assert report["status"] == "failed"
    assert "OCR job belongs to another document" in report["errors"]
    assert "chunk set belongs to another document" in report["errors"]
    db.document_pages.find.assert_not_called()


def test_lineage_validator_detects_chroma_mongo_partial_mismatch():
    ids = {name: ObjectId() for name in ("document", "ocr", "chunk_set", "vector", "chunk", "embedding")}
    db = SimpleNamespace(
        documents=MagicMock(find_one=MagicMock(return_value={"_id": ids["document"]})),
        document_jobs=MagicMock(find_one=MagicMock(return_value={
            "_id": ids["ocr"], "document_id": ids["document"], "job_type": "OCR", "status": "COMPLETED"
        })),
        chunk_sets=MagicMock(find_one=MagicMock(return_value={
            "_id": ids["chunk_set"], "document_id": ids["document"], "source_ocr_job_id": ids["ocr"],
            "status": "COMPLETED", "total_chunks": 1
        })),
        vector_collections=MagicMock(find_one=MagicMock(return_value={
            "_id": ids["vector"], "collection_name": "candidate", "is_active": True
        })),
        document_pages=MagicMock(find=MagicMock(return_value=[{
            "page_number": 1, "cleaned_text": "FIFO", "quality": {"status": "passed"},
            "content_blocks": [{"block_id": "b", "block_type": "prose", "provenance": {"document_id": str(ids["document"])}}]
        }])),
        document_chunks=MagicMock(find=MagicMock(return_value=[{
            "_id": ids["chunk"], "content_hash": "hash", "source": {"source_uri": "golden://x", "block_ids": ["b"]},
            "page_range": {"pages": [1]}
        }])),
        chunk_embeddings=MagicMock(find=MagicMock(return_value=[{
            "_id": ids["embedding"], "chunk_id": ids["chunk"], "status": "INDEXED", "embedding_content_hash": "hash",
            "external_vector_id": "expected-vector"
        }])),
    )
    chroma = MagicMock()
    chroma.get.return_value = {"ids": ["wrong-vector"], "metadatas": [], "documents": []}
    chroma.query.return_value = {"metadatas": [[{"document_id": str(ids["document"]), "chunk_set_id": str(ids["chunk_set"])}]]}
    candidate = CandidateLineage(str(ids["document"]), str(ids["ocr"]), str(ids["chunk_set"]), str(ids["vector"]))
    report = LineageValidator(db, lambda _name: chroma).validate(candidate, smoke_queries=["FIFO"])
    assert report["status"] == "failed"
    assert "Chroma IDs do not match Mongo embedding IDs" in report["errors"]
